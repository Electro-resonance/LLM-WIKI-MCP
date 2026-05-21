#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# =============================================================================
# Created By  : Martin Timms
# Created Date: 21st May 2026
# License: MIT License
# Project: https://github.com/Electro-resonance/LLM-WIKI-MCP
# Description: Local-first Markdown wiki, CLI, and MCP server for LLM context
# retrieval, reflective notes, provenance-aware ingestion, and agentic context.
# =============================================================================

"""
Engine-facing Python API for LLM Wiki MCP.

This module lets an application or agent engine use the LLM Wiki as a context
provider in the same role that a RAG/vector store would normally occupy:

    from llm_wiki_mcp.context_api import LLMWikiContextEngine
    wiki = LLMWikiContextEngine("./wiki_vault")
    context = wiki.context_for_llm("project constraints", max_chars=6000)

The knowledge source remains Markdown; SQLite is used as a fast index/cache.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence

from .server import WikiStore, WikiConfig, set_vault, safe_slug, plain_words, stable_key, search_cache, expand_search_terms, search_query_diagnostics, wiki_links


@dataclass
class ContextDocument:
    """A single context item selected from the wiki."""
    title: str
    body: str
    path: str
    score: float = 0.0


class LLMWikiContextEngine:
    """High-level wrapper for engines that need LLM-ready context.

    Design goals:
    - behave like a simple retrieval/context store;
    - preserve the LLM Wiki principle that Markdown is the source of truth;
    - expose enough metadata for agents to cite, inspect, and maintain pages;
    - avoid forcing an engine to speak MCP when it is running in-process.
    """

    def __init__(self, vault: str | Path = "./wiki_vault", auto_init: bool = True):
        """Internal helper for init."""
        self.vault = Path(vault).expanduser().resolve()
        set_vault(self.vault)
        self.store = WikiStore(WikiConfig(self.vault))
        if auto_init and not any(self.store.cfg.wiki_dir.glob("*.md")):
            self.store.init_seed()

    def ingest_directory(self, directory: str | Path, patterns: Sequence[str] | None = None,
                         recursive: bool = True, limit: int = 1000) -> Dict[str, Any]:
        """Ingest Markdown/TXT-like files from a directory.

        For richer formats such as RTF, PDF, and DOCX, prefer the CLI or
        provenance-aware ingest helpers, which include optional/fallback
        extractors and single-file ingest handling.
        """
        patterns = list(patterns or ["*.md", "*.txt"])
        all_pages: List[str] = []
        reports: List[Dict[str, Any]] = []
        for pattern in patterns:
            report = self.store.ingest_directory(str(directory), pattern=pattern, recursive=recursive, limit=limit)
            reports.append(report)
            all_pages.extend(report.get("pages", []))
        self.store.reindex()
        return {"directory": str(Path(directory).expanduser().resolve()), "patterns": patterns,
                "ingested_count": len(all_pages), "pages": all_pages, "reports": reports}

    def search(self, query: str, limit: int = 8) -> Dict[str, Any]:
        """Search the wiki index and return page snippets."""
        return self.store.search(query, limit=limit)

    def read(self, title: str) -> Dict[str, Any]:
        """Read a page by title/handle."""
        resolved = self.store.resolve(title)
        return self.store.read_page(resolved["canonical_title"])

    def retrieve_documents(self, query: str, limit: int = 6) -> List[ContextDocument]:
        """Return structured documents for downstream reranking or formatting."""
        search = self.store.search(query, limit=limit)
        docs: List[ContextDocument] = []
        for result in search.get("results", []):
            try:
                page = self.store.read_page(result["title"])
            except FileNotFoundError:
                continue
            docs.append(ContextDocument(
                title=page["title"], body=page["body"], path=page["path"], score=float(result.get("score", 0.0))
            ))
        return docs


    def retrieve_articles(self, prompt: str, top_k: int = 5, max_chars_per_article: int = 4000,
                          max_total_chars: int = 12000, include_metadata: bool = True) -> Dict[str, Any]:
        """Prompt -> search -> retrieve top-k maintained wiki articles.

        This is the clearest drop-in call for an engine that previously used a
        RAG store to fetch chunks. It returns article metadata, article bodies,
        and a ready-to-send Markdown context string.
        """
        return self.store.retrieve_articles(
            prompt,
            top_k=top_k,
            max_chars_per_article=max_chars_per_article,
            max_total_chars=max_total_chars,
            include_metadata=include_metadata,
        )


    def ask_ollama(self, question: str, top_k: int | None = None, model: str | None = None,
                   host: str | None = None, dry_run: bool = False, **kwargs: Any) -> Dict[str, Any]:
        """Retrieve wiki context and ask a local Ollama model.

        This is the in-process equivalent of the CLI `ask` command and MCP
        `wiki_ask_ollama` tool. It uses `wiki_vault/llm_wiki_config.json` by
        default, with optional per-call overrides.
        """
        return self.store.ask_ollama(question, top_k=top_k, model=model, host=host, dry_run=dry_run, **kwargs)

    def config(self) -> Dict[str, Any]:
        """Create/read the wiki config, including Ollama host and model defaults."""
        return self.store.config()

    def context_for_llm(self, query: str, limit: int = 6, max_chars: int = 8000,
                        include_instructions: bool = True, include_citations: bool = True) -> str:
        """Build an LLM-ready context string.

        This is the main drop-in replacement for a simple RAG context builder.
        It returns deterministic, readable Markdown rather than anonymous chunks.
        """
        docs = self.retrieve_documents(query, limit=limit)
        parts: List[str] = []
        if include_instructions:
            parts.append(
                "# Retrieved LLM Wiki Context\n\n"
                "Use this maintained wiki context as project memory. Prefer page titles "
                "and explicit notes over unsupported inference. If context is missing or "
                "conflicting, say so and suggest a wiki maintenance action.\n"
            )
            parts.append(f"Query: `{query}`\n")
        used = sum(len(p) for p in parts)
        for idx, doc in enumerate(docs, 1):
            header = f"\n---\n\n## Source {idx}: [[{doc.title}]]\n"
            if include_citations:
                header += f"Path: `{doc.path}`\nScore: `{doc.score}`\n"
            chunk = header + "\n" + doc.body.strip() + "\n"
            if used + len(chunk) > max_chars:
                remaining = max_chars - used
                if remaining <= 0:
                    break
                chunk = chunk[:remaining]
            parts.append(chunk)
            used += len(chunk)
            if used >= max_chars:
                break
        return "".join(parts).strip() + "\n"

    def context_pack(self, query: str, limit: int = 6, max_chars: int = 8000) -> Dict[str, Any]:
        """Return both text context and metadata for an engine."""
        text = self.context_for_llm(query, limit=limit, max_chars=max_chars)
        docs = self.retrieve_documents(query, limit=limit)
        return {
            "query": query,
            "context": text,
            "chars": len(text),
            "sources": [{"title": d.title, "path": d.path, "score": d.score} for d in docs],
            "vault": str(self.vault),
        }

    def add_note(self, title: str, body: str, tags: Optional[List[str]] = None,
                 overwrite: bool = False) -> Dict[str, Any]:
        """Create a maintained wiki note from application code."""
        return self.store.create_page(title, body, tags=tags or ["engine-note"], source="context_api", overwrite=overwrite)

    def stats(self) -> Dict[str, Any]:
        """Return high-level wiki status and health."""
        status = self.store.status()
        lint = self.store.lint()
        health = self.store.health_report()
        return {"status": status, "lint": lint, "health": health}

    def maintenance_report(self) -> Dict[str, Any]:
        """Return maintenance information suitable for dashboards or CI."""
        return {
            "status": self.store.status(),
            "schema": self.store.schema(),
            "graph": self.store.graph_links(),
            "lint": self.store.lint(),
            "health": self.store.health_report(),
        }


# ---------------------------------------------------------------------------
# Release helper functions: tool self-questioning and usage summaries
# ---------------------------------------------------------------------------

def _release_safe_read_text(path):
    """Internal helper for safe read text."""
    from pathlib import Path
    p = Path(path)
    try:
        return p.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return ""

def tool_self_question_context(vault_path="wiki_vault", question="What can this tool tell me about itself?"):
    """Build a Markdown context block for asking the wiki/tool about itself.

    This is intentionally dependency-light: it inspects local documentation,
    config, and the wiki vault, then returns a self-describing context block
    that can be passed to an LLM or written into the wiki.
    """
    from pathlib import Path
    vault = Path(vault_path)
    root = vault.parent if vault.name == "wiki_vault" else Path(".")
    config_path = vault / "llm_wiki_config.json"
    docs_dir = root / "docs"
    candidate_docs = [
        docs_dir / "FUNCTION_LIST.md",
        docs_dir / "ARCHITECTURE.md",
        docs_dir / "02_USER_GUIDE.md",
        docs_dir / "03_FUNCTIONS_AND_TOOLS.md",
        docs_dir / "TOOL_CALLING_ASK_WORKFLOW.md",
        docs_dir / "TOOL_SELF_QUESTIONING.md",
    ]
    parts = [
        "# Tool Self-Question Context",
        f"Question: {question}",
        "",
        "## Active files",
        f"- Vault: `{vault}`",
        f"- Config: `{config_path}`",
        "",
    ]
    if config_path.exists():
        parts += ["## Current config", "```json", _release_safe_read_text(config_path)[:4000], "```", ""]
    for doc in candidate_docs:
        if doc.exists():
            parts += [f"## {doc.name}", _release_safe_read_text(doc)[:5000], ""]
    return {
        "question": question,
        "vault": str(vault),
        "context": "\n".join(parts),
        "sources": [str(p) for p in candidate_docs if p.exists()] + ([str(config_path)] if config_path.exists() else []),
    }

def estimate_tokens_from_text(text):
    """Approximate token count without external tokenizer dependencies."""
    return max(1, int(len(text) / 4))

def wiki_usage_snapshot(vault_path="wiki_vault"):
    """Return a lightweight query/wiki usage snapshot with approximate tokens."""
    from pathlib import Path
    vault = Path(vault_path)
    wiki_dir = vault / "wiki"
    pages = list(wiki_dir.glob("*.md")) if wiki_dir.exists() else []
    total_chars = 0
    total_words = 0
    page_rows = []
    for p in pages:
        txt = _release_safe_read_text(p)
        words = len(txt.split())
        chars = len(txt)
        total_chars += chars
        total_words += words
        page_rows.append({"file": p.name, "words": words, "chars": chars, "tokens_estimate": estimate_tokens_from_text(txt)})
    return {
        "vault": str(vault),
        "pages": len(pages),
        "words": total_words,
        "chars": total_chars,
        "tokens_estimate": estimate_tokens_from_text("x" * total_chars),
        "page_rows": sorted(page_rows, key=lambda r: r["tokens_estimate"], reverse=True),
    }


# ---------------------------------------------------------------------------
# Release recursive ask/history/config helpers
# ---------------------------------------------------------------------------
import json as _release_json
import time as _release_time
from pathlib import Path as _ProjectPath

DEFAULT_CONFIG = {
    "ollama": {
        "host": "http://localhost:11434",
        "model": "llama3.2:3b",
        "timeout_seconds": 120,
        "temperature": 0.2
    },
    "ask": {
        "recursive_enabled": True,
        "compress_history": True,
        "max_tool_calls": 5,
        "max_history_tokens": 4000,
        "history_enabled": True,
        "history_file": "ask_history.jsonl"
    },
    "wiki": {
        "default_top_k": 5,
        "context_char_limit": 12000
    },
    "cli": {
        "color": "auto",
        "json_output": False
    }
}

def _release_deep_merge(base, override):
    """Internal helper for deep merge."""
    out = dict(base)
    for k, v in (override or {}).items():
        if isinstance(v, dict) and isinstance(out.get(k), dict):
            out[k] = _release_deep_merge(out[k], v)
        else:
            out[k] = v
    return out

def release_config_path(vault_path="wiki_vault"):
    """Read, write, or update local wiki configuration for config path."""
    return _ProjectPath(vault_path) / "llm_wiki_config.json"

def release_save_config(vault_path, config):
    """Read, write, or update local wiki configuration for save config."""
    path = release_config_path(vault_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(_release_json.dumps(config, indent=2, sort_keys=True), encoding="utf-8")
    tmp.replace(path)
    return str(path)

def release_strip_json_comments(text):
    """Remove // and # comments from JSON-like config while preserving strings."""
    lines = []
    for line in str(text or "").splitlines():
        out = []
        in_string = False
        escape = False
        i = 0
        while i < len(line):
            ch = line[i]
            nxt = line[i + 1] if i + 1 < len(line) else ""
            if escape:
                out.append(ch)
                escape = False
            elif ch == "\\" and in_string:
                out.append(ch)
                escape = True
            elif ch == '"':
                out.append(ch)
                in_string = not in_string
            elif not in_string and ch == "#":
                break
            elif not in_string and ch == "/" and nxt == "/":
                break
            else:
                out.append(ch)
            i += 1
        stripped = "".join(out).rstrip()
        if stripped.strip():
            lines.append(stripped)
    return "\n".join(lines)

def release_load_config(vault_path="wiki_vault"):
    """Read, write, or update local wiki configuration for load config."""
    path = release_config_path(vault_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        try:
            raw = path.read_text(encoding="utf-8")
            user = _release_json.loads(release_strip_json_comments(raw))
        except Exception:
            user = {}
        return _release_deep_merge(DEFAULT_CONFIG, user)
    user = {}
    cfg = _release_deep_merge(DEFAULT_CONFIG, user)
    release_save_config(vault_path, cfg)
    return cfg

def release_set_config_value(vault_path, dotted_key, value):
    """Read, write, or update local wiki configuration for set config value."""
    cfg = release_load_config(vault_path)
    cur = cfg
    parts = dotted_key.split(".")
    for p in parts[:-1]:
        cur = cur.setdefault(p, {})
    if isinstance(value, str):
        lv = value.strip().lower()
        if lv in {"true", "on", "yes", "1"}:
            value = True
        elif lv in {"false", "off", "no", "0"}:
            value = False
        else:
            try:
                value = int(value) if "." not in value else float(value)
            except Exception:
                pass
    cur[parts[-1]] = value
    release_save_config(vault_path, cfg)
    return cfg

def release_append_history(vault_path, record):
    """Inspect runtime history, token metrics, or usage state for append history."""
    cfg = release_load_config(vault_path)
    ask = cfg.get("ask", {})
    if not ask.get("history_enabled", True):
        return None
    path = _ProjectPath(vault_path) / ask.get("history_file", "ask_history.jsonl")
    record = dict(record)
    record.setdefault("ts", _release_time.time())
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(_release_json.dumps(record, ensure_ascii=False) + "\n")
    return str(path)

def release_load_history(vault_path, max_records=20):
    """Inspect runtime history, token metrics, or usage state for load history."""
    cfg = release_load_config(vault_path)
    path = _ProjectPath(vault_path) / cfg.get("ask", {}).get("history_file", "ask_history.jsonl")
    if not path.exists():
        return []
    rows = []
    for line in path.read_text(encoding="utf-8", errors="replace").splitlines()[-max_records:]:
        try:
            rows.append(_release_json.loads(line))
        except Exception:
            pass
    return rows

def release_estimate_tokens(text):
    """Implement the estimate tokens operation for the local LLM Wiki workflow."""
    return max(1, int(len(str(text)) / 4))

def release_compress_history_text(history, max_tokens=4000):
    """Inspect runtime history, token metrics, or usage state for compress history text."""
    if not history:
        return ""
    lines = ["# Rolling Conversation Memory", ""]
    for row in history[-12:]:
        question = str(row.get("question", "")).strip()
        answer = " ".join(str(row.get("answer", "")).split())[:800]
        sources = row.get("sources") or []
        lines.append(f"- User: {question}")
        if answer:
            lines.append(f"  Answer summary: {answer}")
        if sources:
            lines.append("  Sources: " + ", ".join(map(str, sources[:8])))
    text = "\n".join(lines)
    if release_estimate_tokens(text) > max_tokens:
        text = "# Rolling Conversation Memory (trimmed)\n\n" + text[-max_tokens * 4:]
    return text

def release_recursive_ask_preview(vault_path="wiki_vault", question="", base_context="", recursive=None):
    """Build or execute the local LLM ask workflow for recursive ask preview."""
    cfg = release_load_config(vault_path)
    ask = cfg.get("ask", {})
    enabled = bool(ask.get("recursive_enabled", False)) if recursive is None else bool(recursive)
    history_text = ""
    if enabled and ask.get("compress_history", True):
        history_text = release_compress_history_text(
            release_load_history(vault_path, 50),
            ask.get("max_history_tokens", 4000)
        )
    parts = []
    if history_text:
        parts.append(history_text)
    if base_context:
        parts.append("# Retrieved Wiki Context\n\n" + str(base_context))
    parts.append("# Current User Question\n\n" + str(question))
    context = "\n\n---\n\n".join(parts)
    return {
        "recursive_enabled": enabled,
        "compress_history": bool(ask.get("compress_history", True)),
        "context": context,
        "history_tokens_estimate": release_estimate_tokens(history_text) if history_text else 0,
        "total_tokens_estimate": release_estimate_tokens(context),
        "config": ask,
    }


# ---------------------------------------------------------------------------
# Release ask history: prompt + answer + compression
# ---------------------------------------------------------------------------

def release_history_paths(vault_path="wiki_vault"):
    """Inspect runtime history, token metrics, or usage state for history paths."""
    vault = _ProjectPath(vault_path)
    cfg = release_load_config(vault_path)
    ask = cfg.get("ask", {})
    live = vault / ask.get("history_file", "ask_history.jsonl")
    compressed = vault / ask.get("compressed_history_file", "ask_history_compressed.md")
    return live, compressed

def release_read_compressed_history(vault_path="wiki_vault"):
    """Inspect runtime history, token metrics, or usage state for read compressed history."""
    _, compressed = release_history_paths(vault_path)
    if compressed.exists():
        return compressed.read_text(encoding="utf-8", errors="replace")
    return ""

def release_write_compressed_history(vault_path, text):
    """Inspect runtime history, token metrics, or usage state for write compressed history."""
    _, compressed = release_history_paths(vault_path)
    compressed.parent.mkdir(parents=True, exist_ok=True)
    compressed.write_text(text, encoding="utf-8")
    return str(compressed)

def release_record_ask_turn(vault_path, question, prompt, answer, sources=None, metadata=None):
    """Record the full question, final prompt/context, answer, sources, and token estimates.

    This is intentionally local-first and dependency-free. It keeps exact turns in
    JSONL until the configured token budget is exceeded, then writes/updates a
    compressed Markdown memory file.
    """
    cfg = release_load_config(vault_path)
    ask = cfg.get("ask", {})
    record = {
        "question": question,
        "prompt": prompt,
        "answer": answer,
        "sources": sources or [],
        "metadata": metadata or {},
        "prompt_tokens_estimate": release_estimate_tokens(prompt),
        "answer_tokens_estimate": release_estimate_tokens(answer),
        "ts": _release_time.time(),
    }
    path = release_append_history(vault_path, record)
    release_compress_history_if_needed(vault_path)
    return {"history_file": path, "record": record}

def release_history_token_total(history):
    """Inspect runtime history, token metrics, or usage state for history token total."""
    total = 0
    for row in history:
        total += release_estimate_tokens(row.get("question", ""))
        total += release_estimate_tokens(row.get("prompt", ""))
        total += release_estimate_tokens(row.get("answer", ""))
    return total

def release_compress_history_if_needed(vault_path="wiki_vault"):
    """Inspect runtime history, token metrics, or usage state for compress history if needed."""
    cfg = release_load_config(vault_path)
    ask = cfg.get("ask", {})
    max_tokens = int(ask.get("max_history_tokens", 4000))
    history = release_load_history(vault_path, max_records=500)
    total = release_history_token_total(history)
    if not ask.get("compress_history", True):
        return {"compressed": False, "reason": "compression disabled", "history_tokens": total}
    if total <= max_tokens:
        return {"compressed": False, "reason": "under budget", "history_tokens": total}
    # Keep recent turns in JSONL and move older turns into compressed Markdown.
    keep_recent = int(ask.get("keep_recent_turns", 8))
    older = history[:-keep_recent] if len(history) > keep_recent else []
    recent = history[-keep_recent:] if keep_recent else []
    existing = release_read_compressed_history(vault_path)
    summary_parts = []
    if existing:
        summary_parts.append(existing.strip())
    summary_parts.append("\n# Compressed Ask History Update\n")
    for row in older:
        q = str(row.get("question", "")).strip()
        a = " ".join(str(row.get("answer", "")).split())[:1000]
        src = row.get("sources") or []
        summary_parts.append(f"## Turn: {q[:120]}")
        if a:
            summary_parts.append(a)
        if src:
            summary_parts.append("Sources: " + ", ".join(map(str, src[:10])))
        summary_parts.append("")
    compressed_text = "\n".join(summary_parts).strip() + "\n"
    compressed_path = release_write_compressed_history(vault_path, compressed_text)
    # Rewrite JSONL with only recent exact turns.
    live, _ = release_history_paths(vault_path)
    live.parent.mkdir(parents=True, exist_ok=True)
    with live.open("w", encoding="utf-8") as f:
        for row in recent:
            f.write(_release_json.dumps(row, ensure_ascii=False) + "\n")
    return {
        "compressed": True,
        "history_tokens": total,
        "compressed_path": compressed_path,
        "kept_recent_turns": len(recent),
        "compressed_turns": len(older),
    }

def release_recursive_context_with_memory(vault_path, question, base_context=""):
    """Implement the recursive context with memory operation for the local LLM Wiki workflow."""
    cfg = release_load_config(vault_path)
    ask = cfg.get("ask", {})
    compressed = release_read_compressed_history(vault_path)
    recent = release_compress_history_text(release_load_history(vault_path, int(ask.get("keep_recent_turns", 8))), ask.get("max_history_tokens", 4000))
    parts = []
    if ask.get("recursive_enabled", True):
        if compressed:
            parts.append("# Compressed Long-Term Ask Memory\n\n" + compressed)
        if recent:
            parts.append("# Recent Ask Memory\n\n" + recent)
    if base_context:
        parts.append("# Retrieved Wiki Context\n\n" + str(base_context))
    parts.append("# Current User Question\n\n" + str(question))
    context = "\n\n---\n\n".join(parts)
    return {
        "context": context,
        "recursive_enabled": bool(ask.get("recursive_enabled", True)),
        "history_tokens_estimate": release_estimate_tokens((compressed or "") + (recent or "")),
        "total_tokens_estimate": release_estimate_tokens(context),
        "compressed_history_present": bool(compressed),
    }


# ---------------------------------------------------------------------------
# Release metrics, token/s, and self-access helpers
# ---------------------------------------------------------------------------

def release_metrics_path(vault_path="wiki_vault"):
    """Inspect runtime history, token metrics, or usage state for metrics path."""
    return _ProjectPath(vault_path) / "ask_metrics.jsonl"

def release_record_llm_metrics(vault_path, question="", prompt="", answer="", elapsed_seconds=None, sources=None, metadata=None):
    """Record per-ask prompt/answer token estimates, elapsed seconds, and token/s.

    The token counts are dependency-free estimates. If the Ollama response contains
    exact token counts in future, metadata can carry those exact values.
    """
    prompt_tokens = release_estimate_tokens(prompt)
    answer_tokens = release_estimate_tokens(answer)
    total_tokens = prompt_tokens + answer_tokens
    elapsed = float(elapsed_seconds or 0.0)
    output_tokens_per_second = (answer_tokens / elapsed) if elapsed > 0 else None
    total_tokens_per_second = (total_tokens / elapsed) if elapsed > 0 else None
    rec = {
        "ts": _release_time.time(),
        "question": question,
        "prompt_tokens_estimate": prompt_tokens,
        "answer_tokens_estimate": answer_tokens,
        "total_tokens_estimate": total_tokens,
        "elapsed_seconds": elapsed,
        "output_tokens_per_second": output_tokens_per_second,
        "total_tokens_per_second": total_tokens_per_second,
        "sources": sources or [],
        "metadata": metadata or {},
    }
    path = release_metrics_path(vault_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(_release_json.dumps(rec, ensure_ascii=False) + "\n")
    return rec

def release_load_metrics(vault_path="wiki_vault", limit=100):
    """Inspect runtime history, token metrics, or usage state for load metrics."""
    path = release_metrics_path(vault_path)
    if not path.exists():
        return []
    rows = []
    for line in path.read_text(encoding="utf-8", errors="replace").splitlines()[-int(limit):]:
        try:
            rows.append(_release_json.loads(line))
        except Exception:
            pass
    return rows

def release_metrics_summary(vault_path="wiki_vault", limit=100):
    """Inspect runtime history, token metrics, or usage state for metrics summary."""
    rows = release_load_metrics(vault_path, limit)
    if not rows:
        return {
            "count": 0,
            "avg_elapsed_seconds": None,
            "avg_output_tokens_per_second": None,
            "avg_total_tokens_per_second": None,
            "total_prompt_tokens_estimate": 0,
            "total_answer_tokens_estimate": 0,
            "recent": [],
        }
    def avg(vals):
        """Implement the avg operation for the local LLM Wiki workflow."""
        vals = [v for v in vals if isinstance(v, (int, float))]
        return sum(vals) / len(vals) if vals else None
    return {
        "count": len(rows),
        "avg_elapsed_seconds": avg([r.get("elapsed_seconds") for r in rows]),
        "avg_output_tokens_per_second": avg([r.get("output_tokens_per_second") for r in rows]),
        "avg_total_tokens_per_second": avg([r.get("total_tokens_per_second") for r in rows]),
        "total_prompt_tokens_estimate": sum(int(r.get("prompt_tokens_estimate") or 0) for r in rows),
        "total_answer_tokens_estimate": sum(int(r.get("answer_tokens_estimate") or 0) for r in rows),
        "recent": rows[-10:],
    }

def release_wiki_usage_snapshot(vault_path="wiki_vault"):
    """Self-access usage snapshot for the wiki, history, and metrics."""
    vault = _ProjectPath(vault_path)
    wiki_dir = vault / "wiki"
    pages = sorted(wiki_dir.glob("*.md")) if wiki_dir.exists() else []
    page_rows = []
    total_chars = 0
    total_words = 0
    for p in pages:
        txt = p.read_text(encoding="utf-8", errors="replace")
        words = len(txt.split())
        chars = len(txt)
        total_chars += chars
        total_words += words
        page_rows.append({
            "file": p.name,
            "words": words,
            "chars": chars,
            "tokens_estimate": release_estimate_tokens(txt),
        })
    metrics = release_metrics_summary(vault_path)
    history = release_load_history(vault_path, 500)
    compressed = release_read_compressed_history(vault_path) if "release_read_compressed_history" in globals() else ""
    return {
        "vault": str(vault),
        "pages": len(pages),
        "words": total_words,
        "chars": total_chars,
        "tokens_estimate": release_estimate_tokens("x" * total_chars),
        "largest_pages": sorted(page_rows, key=lambda r: r["tokens_estimate"], reverse=True)[:20],
        "ask_history_turns": len(history),
        "ask_history_tokens_estimate": release_history_token_total(history) if "release_history_token_total" in globals() else 0,
        "compressed_history_tokens_estimate": release_estimate_tokens(compressed) if compressed else 0,
        "metrics": metrics,
    }

def release_self_access_snapshot(vault_path="wiki_vault"):
    """Single structured snapshot so the LLM/tool can access its own state."""
    cfg = release_load_config(vault_path)
    usage = release_wiki_usage_snapshot(vault_path)
    compression = release_compress_history_if_needed(vault_path) if "release_compress_history_if_needed" in globals() else {}
    return {
        "config": cfg,
        "usage": usage,
        "compression": compression,
        "metrics": usage.get("metrics", {}),
        "history_recent": release_load_history(vault_path, 10),
    }

def release_record_ask_turn_with_metrics(vault_path, question, prompt, answer, sources=None, elapsed_seconds=None, metadata=None):
    """Build or execute the local LLM ask workflow for record ask turn with metrics."""
    hist = release_record_ask_turn(vault_path, question, prompt, answer, sources or [], metadata or {})
    met = release_record_llm_metrics(vault_path, question, prompt, answer, elapsed_seconds, sources or [], metadata or {})
    return {"history": hist, "metrics": met}


# ---------------------------------------------------------------------------
# Release reflective runtime helpers
# ---------------------------------------------------------------------------

def release_runtime_journal_path(vault_path="wiki_vault"):
    """Implement the runtime journal path operation for the local LLM Wiki workflow."""
    return _ProjectPath(vault_path) / "runtime_journal.jsonl"

def release_record_runtime_event(vault_path="wiki_vault", command="", query="", tools_used=None, success=True, latency_seconds=None, token_metrics=None, metadata=None):
    """Implement the record runtime event operation for the local LLM Wiki workflow."""
    rec = {
        "ts": _release_time.time(),
        "command": command,
        "query": query,
        "tools_used": tools_used or [],
        "success": bool(success),
        "latency_seconds": latency_seconds,
        "token_metrics": token_metrics or {},
        "metadata": metadata or {},
    }
    path = release_runtime_journal_path(vault_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(_release_json.dumps(rec, ensure_ascii=False) + "\n")
    return rec

def release_load_runtime_journal(vault_path="wiki_vault", limit=50):
    """Implement the load runtime journal operation for the local LLM Wiki workflow."""
    path = release_runtime_journal_path(vault_path)
    if not path.exists():
        return []
    rows = []
    for line in path.read_text(encoding="utf-8", errors="replace").splitlines()[-int(limit):]:
        try:
            rows.append(_release_json.loads(line))
        except Exception:
            pass
    return rows

def release_runtime_capabilities(vault_path="wiki_vault"):
    """Implement the runtime capabilities operation for the local LLM Wiki workflow."""
    cfg = release_load_config(vault_path)
    ask = cfg.get("ask", {})
    ollama = cfg.get("ollama", {})
    return {
        "llm": {
            "provider": "ollama",
            "host": ollama.get("host"),
            "model": ollama.get("model"),
            "streaming": bool(ollama.get("stream", False)),
            "recursive_ask": bool(ask.get("recursive_enabled", True)),
            "history_compression": bool(ask.get("compress_history", True)),
        },
        "retrieval": {
            "sqlite_fts": True,
            "markdown_source_of_truth": True,
            "singular_plural_expansion": True,
            "graph_retrieval": "available",
        },
        "runtime": {
            "self_introspection": True,
            "runtime_journaling": True,
            "token_metrics": True,
            "self_access_snapshots": True,
        },
        "maintenance": {
            "lint": True,
            "repair": True,
            "backlinks": True,
            "graph_export": True,
            "reconcile_docs": True,
        },
        "files": {
            "config": str(release_config_path(vault_path)),
            "runtime_journal": str(release_runtime_journal_path(vault_path)),
            "metrics": str(release_metrics_path(vault_path)) if "release_metrics_path" in globals() else None,
        }
    }

def release_command_tool_map():
    """Implement the command tool map operation for the local LLM Wiki workflow."""
    return {
        "ask": ["wiki_retrieve_articles", "wiki_recursive_context", "ollama_generate", "wiki_record_ask_turn_with_metrics", "wiki_record_runtime_event"],
        "search": ["wiki_search", "wiki_search_diagnostics"],
        "retrieve": ["wiki_retrieve_articles"],
        "stats": ["wiki_status", "wiki_health_report"],
        "usage": ["wiki_usage_snapshot"],
        "tokens": ["wiki_llm_metrics"],
        "self": ["wiki_tool_self_question", "wiki_self_access_snapshot"],
        "self-stats": ["wiki_self_access_snapshot"],
        "self-usage": ["wiki_usage_snapshot"],
        "history-status": ["wiki_compress_history", "wiki_ask_history"],
        "lint": ["wiki_lint"],
        "repair": ["wiki_repair_lint"],
        "mermaid": ["wiki_mermaid_graph"],
        "map": ["wiki_mermaid_neighbourhood"],
        "config": ["wiki_config_load", "wiki_config_set"],
        "capabilities": ["wiki_capabilities"],
        "runtime-journal": ["wiki_runtime_journal"],
        "reconcile": ["wiki_reconcile_docs"],
    }

def release_reconcile_docs(vault_path="wiki_vault", project_root=None):
    """Implement the reconcile docs operation for the local LLM Wiki workflow."""
    root = _ProjectPath(project_root) if project_root else _ProjectPath(vault_path).parent
    docs = root / "docs"
    function_list = root / "FUNCTION_LIST.md"
    readme = root / "README.md"
    expected_docs = [
        "RUNTIME_EXECUTION_JOURNAL.md",
        "CAPABILITIES_COMMAND.md",
        "ASK_PIPELINE_INSTRUMENTATION.md",
        "RUNTIME_CONTEXT_INJECTION.md",
        "SELF_HEALING_DOCUMENTATION.md",
        "COMMAND_TO_TOOL_MAPPING.md",
        "RUNTIME_GRAPH_VISUALISATION.md",
        "RECURSIVE_COGNITION_LAYER.md",
        "TOOL_ONTOLOGY.md",
        "REFLECTIVE_OPERATIONAL_MEMORY_ARCHITECTURE.md",
    ]
    missing_docs = [d for d in expected_docs if not (docs / d).exists()]
    commands = list(release_command_tool_map().keys())
    fn_text = function_list.read_text(encoding="utf-8", errors="replace") if function_list.exists() else ""
    undocumented = [c for c in commands if c not in fn_text]
    return {
        "docs_dir": str(docs),
        "missing_docs": missing_docs,
        "commands_checked": commands,
        "undocumented_commands_in_function_list": undocumented,
        "readme_present": readme.exists(),
        "function_list_present": function_list.exists(),
        "ok": not missing_docs and not undocumented,
    }

def release_runtime_graph_mermaid():
    """Create graph-oriented output for runtime graph mermaid."""
    return """```mermaid
flowchart TD
    User[User / Agent] --> CLI[CLI Command]
    CLI --> Capabilities[Capabilities Surface]
    CLI --> Retrieval[Search / Retrieve]
    Retrieval --> SQLite[SQLite FTS Index]
    Retrieval --> Markdown[Markdown Wiki Pages]
    CLI --> Prompt[Prompt Builder]
    Prompt --> RecMem[Recursive Memory]
    RecMem --> History[ask_history.jsonl]
    RecMem --> Compressed[ask_history_compressed.md]
    Prompt --> Ollama[Ollama LLM]
    Ollama --> Answer[Answer]
    Answer --> Metrics[ask_metrics.jsonl]
    Answer --> RuntimeJournal[runtime_journal.jsonl]
    RuntimeJournal --> SelfAccess[Self Access Snapshot]
    SelfAccess --> LLM[LLM Runtime Context]
```"""


# ---------------------------------------------------------------------------
# Natural-language safe self-access dispatch
# ---------------------------------------------------------------------------

def release_detect_safe_self_command(text):
    """Map natural-language ask requests to safe read-only tool commands.

    This fixes the UX gap where `ask run self-stats` was answered by the LLM
    instead of executing the self-access tool.
    """
    q = " ".join(str(text or "").lower().replace("-", " ").split())
    mappings = [
        ("self-stats", ["self stats", "selfstats", "run self stats", "tool stats", "access the tool stats", "read the stats for this tool"]),
        ("self-usage", ["self usage", "own usage", "ask your own usage", "tool usage", "usage snapshot"]),
        ("tokens", ["token/s", "tokens per second", "token speed", "tokens s", "token metrics"]),
        ("stats", ["run stats", "show stats", "wiki stats"]),
        ("usage", ["run usage", "show usage", "wiki usage"]),
        ("capabilities", ["capabilities", "what can you do", "available capabilities"]),
        ("runtime-journal", ["runtime journal", "execution journal", "recent runtime events"]),
        ("command-map", ["command map", "command to tool", "tool mapping"]),
        ("history-status", ["history status", "recursive memory status", "compression status"]),
        ("reconcile", ["reconcile docs", "check docs", "self healing documentation"]),
    ]
    for command, phrases in mappings:
        if command.replace("-", " ") in q:
            return command
        for phrase in phrases:
            if phrase in q:
                return command
    # "run the self-stats" may normalise to "run the self stats"
    if "run" in q and "self" in q and "stat" in q:
        return "self-stats"
    if "can you" in q and "use" in q and "self" in q and "stat" in q:
        return "self-stats"
    return None

def release_safe_self_command_result(vault_path, command):
    """Execute safe read-only self-access commands and return structured data."""
    if command == "self-stats":
        return {"command": command, "result": release_self_access_snapshot(vault_path)}
    if command == "self-usage":
        return {"command": command, "result": release_wiki_usage_snapshot(vault_path)}
    if command == "tokens":
        return {"command": command, "result": release_metrics_summary(vault_path, 100)}
    if command == "stats":
        # lightweight status via usage snapshot; full CLI stats is store-backed
        return {"command": command, "result": release_wiki_usage_snapshot(vault_path)}
    if command == "usage":
        return {"command": command, "result": release_wiki_usage_snapshot(vault_path)}
    if command == "capabilities":
        return {"command": command, "result": release_runtime_capabilities(vault_path)}
    if command == "runtime-journal":
        return {"command": command, "result": {"events": release_load_runtime_journal(vault_path, 20)}}
    if command == "command-map":
        return {"command": command, "result": release_command_tool_map()}
    if command == "history-status":
        return {"command": command, "result": {
            "compression": release_compress_history_if_needed(vault_path),
            "history": release_load_history(vault_path, 10),
            "compressed_history_present": bool(release_read_compressed_history(vault_path)),
        }}
    if command == "reconcile":
        return {"command": command, "result": release_reconcile_docs(vault_path)}
    return {"command": command, "error": "Unknown safe self-access command"}


# ---------------------------------------------------------------------------
# Release AI page notes, link discovery and self-extending wiki helpers
# ---------------------------------------------------------------------------
import re as _release_re
from collections import Counter as _WordCounter

_COMMON_STOPWORDS = {
    "the","and","for","with","that","this","from","into","your","you","are","can","has","have","not","but","was","were",
    "wiki","page","source","content","file","docs","doc","llm","mcp","tool","tools","command","commands","system",
    "using","used","use","about","what","how","why","when","where","then","than","also","will","would","should",
    "there","their","they","them","its","itself","these","those","each","more","less","been","being","over","under"
}

def release_notes_dir(vault_path="wiki_vault"):
    """Read, create, or summarise AI sidecar notes for notes dir."""
    return _ProjectPath(vault_path) / "notes"

def release_wiki_dir(vault_path="wiki_vault"):
    """Implement the wiki dir operation for the local LLM Wiki workflow."""
    return _ProjectPath(vault_path) / "wiki"

def release_slug(title):
    """Implement the slug operation for the local LLM Wiki workflow."""
    text = str(title).strip().replace("/", "_").replace("\\", "_").replace(":", " -")
    text = _release_re.sub(r"\s+", " ", text)
    return text if text.endswith(".md") else text + ".md"

def release_page_path(vault_path, title):
    """Implement the page path operation for the local LLM Wiki workflow."""
    title = str(title)
    wiki_dir = release_wiki_dir(vault_path)
    direct = wiki_dir / release_slug(title)
    if direct.exists(): return direct
    wanted = title.lower().replace(".md", "")
    for p in wiki_dir.glob("*.md"):
        stem = p.stem.lower()
        if stem == wanted or stem.endswith(wanted) or wanted in stem:
            return p
    return direct

def release_note_path(vault_path, title):
    """Read, create, or summarise AI sidecar notes for note path."""
    p = release_page_path(vault_path, title)
    notes = release_notes_dir(vault_path); notes.mkdir(parents=True, exist_ok=True)
    return notes / p.name

def release_read_page(vault_path, title):
    """Implement the read page operation for the local LLM Wiki workflow."""
    p = release_page_path(vault_path, title)
    return p.read_text(encoding="utf-8", errors="replace") if p.exists() else ""

def release_read_note(vault_path, title):
    """Read, create, or summarise AI sidecar notes for read note."""
    p = release_note_path(vault_path, title)
    return p.read_text(encoding="utf-8", errors="replace") if p.exists() else ""

def release_extract_terms(text, max_terms=20):
    """Implement the extract terms operation for the local LLM Wiki workflow."""
    words = _release_re.findall(r"[A-Za-z][A-Za-z0-9_\-]{2,}", text or "")
    words = [w.lower().strip("-_") for w in words]
    words = [w for w in words if len(w) > 2 and w not in _COMMON_STOPWORDS]
    counts = _WordCounter(words)
    return [w for w, _ in counts.most_common(max_terms)]

def release_find_candidate_links(vault_path, title, limit=12):
    """Implement the find candidate links operation for the local LLM Wiki workflow."""
    current_path = release_page_path(vault_path, title)
    current_title = current_path.stem
    text = release_read_page(vault_path, title) + "\n" + release_read_note(vault_path, title)
    terms = set(release_extract_terms(text, 40))
    rows=[]
    wiki = release_wiki_dir(vault_path)
    if not wiki.exists(): return rows
    for p in wiki.glob("*.md"):
        if p.name == current_path.name: continue
        body = p.read_text(encoding="utf-8", errors="replace")
        blob = (p.stem + " " + body[:3000]).lower()
        score=0; matched=[]
        for term in terms:
            if term and term in blob:
                score += 1
                if len(matched) < 8: matched.append(term)
        if current_title.lower() in body.lower():
            score += 5; matched.append("backlink-title")
        if score: rows.append({"title":p.stem,"score":score,"matched_terms":matched})
    return sorted(rows, key=lambda r:(-r["score"], r["title"]))[:int(limit)]

def release_build_page_note(vault_path, title, mode="iterate", max_links=12):
    """Read, create, or summarise AI sidecar notes for build page note."""
    page_text = release_read_page(vault_path, title)
    if not page_text: return {"ok":False,"error":f"Page not found: {title}"}
    old_note = release_read_note(vault_path, title)
    terms = release_extract_terms(page_text + "\n" + old_note, 25)
    links = release_find_candidate_links(vault_path, title, max_links)
    source_path = release_page_path(vault_path, title)
    link_lines = [f"- [[{r['title']}]] — score {r['score']}; matched: {', '.join(r['matched_terms'][:8])}" for r in links]
    if not link_lines: link_lines = ["- No strong candidate links found yet."]
    terms_lines = "\n".join("- " + t for t in terms[:20]) if terms else "- No extracted terms"
    note = f"""# AI Notes — {source_path.stem}

Source page: [[{source_path.stem}]]
Mode: {mode}

## Working Summary

This page appears to discuss: {', '.join(terms[:12]) if terms else 'no clear extracted terms yet'}.

## Key Terms

{terms_lines}

## Candidate Links To Review

{chr(10).join(link_lines)}

## Iteration Notes

- These notes are AI-maintained sidecar notes, not the canonical source document.
- Search should include both the source page and these notes.
- Candidate links are suggestions until promoted into the source page or accepted by a human/agent.
- Re-run `notes-iterate {source_path.stem}` after related pages are added or edited.

## Suggested Next Actions

- Review candidate links and promote the most useful ones.
- Add missing bridge pages where repeated concepts have no destination.
- Use `notes-search <query>` to test whether these notes improve retrieval.
"""
    path = release_note_path(vault_path, title); path.write_text(note, encoding="utf-8")
    return {"ok":True,"title":source_path.stem,"note_path":str(path),"candidate_links":links,"terms":terms}

def release_iterate_all_notes(vault_path, limit=None):
    """Read, create, or summarise AI sidecar notes for iterate all notes."""
    wiki = release_wiki_dir(vault_path)
    pages=list(wiki.glob("*.md")) if wiki.exists() else []
    if limit: pages=pages[:int(limit)]
    return {"ok":True,"count":len(pages),"results":[release_build_page_note(vault_path,p.stem) for p in pages]}

def release_combined_page_text(vault_path, title):
    """Implement the combined page text operation for the local LLM Wiki workflow."""
    return release_read_page(vault_path, title) + "\n\n# AI Sidecar Notes\n\n" + release_read_note(vault_path, title)

def release_search_with_notes(vault_path, query, limit=8):
    """Search or retrieve wiki content for search with notes."""
    q_terms=set(release_extract_terms(query,20)); rows=[]
    wiki=release_wiki_dir(vault_path)
    if not wiki.exists(): return {"query":query,"count":0,"results":[]}
    for p in wiki.glob("*.md"):
        source=p.read_text(encoding="utf-8", errors="replace")
        np=release_note_path(vault_path,p.stem)
        note=np.read_text(encoding="utf-8", errors="replace") if np.exists() else ""
        blob=(p.stem+"\n"+source+"\n"+note).lower()
        score=0; matched=[]
        for term in q_terms:
            if term in blob:
                bonus=3 if term in p.stem.lower() else 1
                if term in note.lower(): bonus += 1
                score += bonus; matched.append(term)
        if score: rows.append({"title":p.stem,"score":score,"matched_terms":matched[:12],"has_note":bool(note),"snippet":(note or source)[:500].replace("\n"," ")})
    return {"query":query,"count":len(rows),"results":sorted(rows,key=lambda r:(-r["score"],r["title"]))[:int(limit)]}

def release_notes_status(vault_path):
    """Read, create, or summarise AI sidecar notes for notes status."""
    wiki=release_wiki_dir(vault_path); pages=list(wiki.glob("*.md")) if wiki.exists() else []
    notes=list(release_notes_dir(vault_path).glob("*.md")) if release_notes_dir(vault_path).exists() else []
    missing=[p.stem for p in pages if not release_note_path(vault_path,p.stem).exists()]
    return {"pages":len(pages),"notes":len(notes),"missing_notes":missing,"coverage_percent":round((len(pages)-len(missing))*100/len(pages),2) if pages else 0,"notes_dir":str(release_notes_dir(vault_path))}

def release_notes_graph_mermaid(vault_path, limit_pages=80):
    """Create graph-oriented output for notes graph mermaid."""
    lines=["```mermaid","flowchart TD"]; count=0
    notes_dir=release_notes_dir(vault_path)
    if notes_dir.exists():
        for note in sorted(notes_dir.glob("*.md")):
            if count >= int(limit_pages): break
            text=note.read_text(encoding="utf-8", errors="replace"); src=note.stem
            for target in _release_re.findall(r"\[\[([^\]]+)\]\]", text):
                if target != src:
                    a=_release_re.sub(r"[^A-Za-z0-9_]", "_", src); b=_release_re.sub(r"[^A-Za-z0-9_]", "_", target)
                    lines.append(f'    {a}["{src}"] --> {b}["{target}"]')
            count += 1
    lines.append("```"); return "\n".join(lines)


# ---------------------------------------------------------------------------
# Bounded agentic ask with local read-only tool-calling
# ---------------------------------------------------------------------------

def release_read_page_tool(vault_path, title):
    """Implement the read page tool operation for the local LLM Wiki workflow."""
    path = release_page_path(vault_path, title) if "release_page_path" in globals() else (_ProjectPath(vault_path) / "wiki" / (str(title) if str(title).endswith(".md") else str(title) + ".md"))
    if not path.exists():
        return {"ok": False, "title": title, "error": "page not found"}
    return {"ok": True, "title": path.stem, "path": str(path), "text": path.read_text(encoding="utf-8", errors="replace")[:8000]}

def release_list_pages_tool(vault_path, limit=200):
    """Implement the list pages tool operation for the local LLM Wiki workflow."""
    wiki = _ProjectPath(vault_path) / "wiki"
    pages = sorted([p.stem for p in wiki.glob("*.md")]) if wiki.exists() else []
    return {"ok": True, "pages": pages[:int(limit)], "count": len(pages)}

def release_agentic_tool_manifest():
    """Implement the agentic tool manifest operation for the local LLM Wiki workflow."""
    return {
        "wiki_search_with_notes": "Search source pages plus AI sidecar notes.",
        "wiki_read_page": "Read one wiki page.",
        "wiki_usage_snapshot": "Read wiki/page/history usage statistics.",
        "wiki_self_access_snapshot": "Read tool config, usage, metrics and history state.",
        "wiki_capabilities": "Read live capability surface.",
        "wiki_candidate_links": "Find candidate links for a page.",
        "wiki_notes_status": "Read AI notes coverage.",
        "wiki_list_pages": "List wiki pages.",
    }

def release_agentic_select_tools(question, step=0):
    """Implement the agentic select tools operation for the local LLM Wiki workflow."""
    q = str(question or "").lower()
    tools = []
    if any(w in q for w in ["usage","stats","token","tokens","metrics"]):
        tools += ["wiki_usage_snapshot","wiki_self_access_snapshot"]
    if any(w in q for w in ["capability","capabilities","can you","able to"]):
        tools += ["wiki_capabilities"]
    if any(w in q for w in ["link","links","related","connect","connection","notes"]):
        tools += ["wiki_search_with_notes","wiki_candidate_links","wiki_notes_status"]
    if any(w in q for w in ["tool","function","mcp","command","architecture","limitation","limitations"]):
        tools += ["wiki_search_with_notes","wiki_list_pages","wiki_capabilities"]
    if not tools:
        tools += ["wiki_search_with_notes","wiki_usage_snapshot"]
    seen=[]
    for t in tools:
        if t not in seen: seen.append(t)
    return seen[:int(6)]

def release_execute_agentic_tool(vault_path, tool, question, state=None):
    """Implement the execute agentic tool operation for the local LLM Wiki workflow."""
    state = state or {}
    if tool == "wiki_search_with_notes":
        return release_search_with_notes(vault_path, question, 8)
    if tool == "wiki_usage_snapshot":
        return release_wiki_usage_snapshot(vault_path)
    if tool == "wiki_self_access_snapshot":
        return release_self_access_snapshot(vault_path)
    if tool == "wiki_capabilities":
        return release_runtime_capabilities(vault_path)
    if tool == "wiki_candidate_links":
        title = state.get("focus_page")
        if not title:
            search = release_search_with_notes(vault_path, question, 1)
            if search.get("results"): title = search["results"][0]["title"]
        return {"title": title, "candidate_links": release_find_candidate_links(vault_path, title or "") if title else []}
    if tool == "wiki_notes_status":
        return release_notes_status(vault_path)
    if tool == "wiki_list_pages":
        return release_list_pages_tool(vault_path)
    if tool == "wiki_read_page":
        return release_read_page_tool(vault_path, state.get("focus_page") or question)
    return {"ok": False, "error": f"unknown tool: {tool}"}


def release_agentic_ask(vault_path="wiki_vault", question="", max_tool_calls=5, dry_run=False, synthesize=True, protocol_mode=False):
    """Release compatible agentic ask.

    Handles older call sites where the CLI passed synthesize=...
    but the packaged function did not accept that keyword.
    """
    original_question = question
    question = release_repair_query_text(question) if "release_repair_query_text" in globals() else str(question or "")

    # Build evidence with safe read-only tools.
    state = {}
    trace = []
    evidence = []

    selected = release_agentic_select_tools(question) if "release_agentic_select_tools" in globals() else ["wiki_search_with_notes", "wiki_capabilities"]
    if "limitation" in question and "wiki_limitations_report" not in selected:
        selected.insert(0, "wiki_limitations_report")

    for tool in selected[:int(max_tool_calls)]:
        if "release_execute_agent_tool_by_name" in globals():
            result = release_execute_agent_tool_by_name(vault_path, tool, {"query": question}, question, state)
        elif tool == "wiki_search_with_notes" and "release_search_with_notes" in globals():
            result = release_search_with_notes(vault_path, question, 8)
        elif tool == "wiki_capabilities" and "release_runtime_capabilities" in globals():
            result = release_runtime_capabilities(vault_path)
        elif tool == "wiki_usage_snapshot" and "release_wiki_usage_snapshot" in globals():
            result = release_wiki_usage_snapshot(vault_path)
        elif tool == "wiki_self_access_snapshot" and "release_self_access_snapshot" in globals():
            result = release_self_access_snapshot(vault_path)
        elif tool == "wiki_limitations_report" and "release_limitations_report" in globals():
            result = release_limitations_report(vault_path)
        else:
            result = {"ok": False, "error": f"Tool unavailable: {tool}"}

        trace.append({"tool": tool, "arguments": {"query": question}, "result_type": type(result).__name__})
        evidence.append({"tool": tool, "result": result})
        if tool == "wiki_search_with_notes" and isinstance(result, dict) and result.get("results"):
            state["focus_page"] = result["results"][0].get("title")

    manifest = release_agentic_tool_manifest() if "release_agentic_tool_manifest" in globals() else {}
    prompt_parts = [
        "# Agentic Ask Evidence Pack",
        f"Original question: {original_question}",
        f"Repaired question: {question}",
        "",
        "## Tool Manifest",
        _release_json.dumps(manifest, indent=2) if "_release_json" in globals() else str(manifest),
        "",
        "## Tool Trace",
        _release_json.dumps(trace, indent=2) if "_release_json" in globals() else str(trace),
        "",
        "## Evidence",
    ]
    for item in evidence:
        if "_release_json" in globals():
            prompt_parts.append("```json\n" + _release_json.dumps(item, indent=2, ensure_ascii=False)[:7000] + "\n```")
        else:
            prompt_parts.append(str(item)[:7000])
    prompt_parts.append("""
## Final Answer Instructions

Answer the user's question using the evidence above.
Prefer current runtime/tool state over stale docs.
Mention the tools/pages used as evidence.
If evidence is missing, say what is missing and suggest the next safe tool to call.
""")
    prompt = "\n".join(prompt_parts)

    if "release_record_runtime_event" in globals():
        runtime_event = release_record_runtime_event(
            vault_path,
            "ask-agentic",
            original_question,
            [t["tool"] for t in trace],
            True,
            None,
            {"prompt_tokens_estimate": release_estimate_tokens(prompt) if "release_estimate_tokens" in globals() else len(prompt)//4},
            {"dry_run": dry_run, "protocol_mode": protocol_mode, "synthesize": synthesize},
        )
    else:
        runtime_event = {}

    if dry_run or not synthesize:
        return {
            "ok": True,
            "question": original_question,
            "repaired_question": question,
            "dry_run": True,
            "tools_called": [t["tool"] for t in trace],
            "trace": trace,
            "evidence": evidence,
            "llm_prompt": prompt,
            "runtime_event": runtime_event,
            "answer": None,
        }

    if "release_ollama_generate" in globals():
        llm = release_ollama_generate(vault_path, prompt)
        answer = llm.get("answer") if llm.get("ok") else f"Agentic evidence was built, but Ollama synthesis failed: {llm.get('error')}"
    else:
        llm = {"ok": False, "error": "release_ollama_generate is not available"}
        answer = "Agentic evidence was built, but Ollama synthesis is not available in this package."

    if "release_record_ask_turn_with_metrics" in globals():
        metrics = release_record_ask_turn_with_metrics(
            vault_path,
            original_question,
            prompt,
            answer,
            [e.get("tool") for e in evidence],
            llm.get("elapsed_seconds"),
            {"agentic": True, "protocol_mode": protocol_mode, "ollama": llm},
        )
    else:
        metrics = {}

    return {
        "ok": bool(llm.get("ok")),
        "question": original_question,
        "repaired_question": question,
        "answer": answer,
        "tools_called": [t["tool"] for t in trace],
        "trace": trace,
        "evidence": evidence,
        "llm": llm,
        "metrics": metrics,
        "runtime_event": runtime_event,
        "llm_prompt": prompt,
    }

def release_limitations_report(vault_path="wiki_vault"):
    """Implement the limitations report operation for the local LLM Wiki workflow."""
    return {
        "known_limitations_fixed_or_reduced": [
            "Agentic ask now has bounded read-only tool-calling via ask-agentic.",
            "Self-access and runtime state can be called as tools rather than only described.",
            "AI notes can improve search and link discovery.",
            "Runtime journal can record tool traces."
        ],
        "remaining_limitations": [
            "Fully autonomous LLM JSON tool-call parsing is not enabled by default.",
            "Exact token counts are still estimated unless model metrics are supplied.",
            "Mutating actions remain blocked from natural-language dispatch.",
            "AI notes propose links but do not promote them into canonical pages automatically."
        ],
        "recommended_next_steps": [
            "Wire ask-agentic evidence pack to Ollama synthesis.",
            "Add strict JSON tool-call protocol for local models.",
            "Add review/apply flow to promote AI note links.",
            "Add file watcher reindex after edits."
        ]
    }

def release_repair_query_text(text):
    """Implement the repair query text operation for the local LLM Wiki workflow."""
    q = str(text or "").lower()
    return q.replace("limitationsof", "limitations of").replace("devleopment", "development").replace("ddo", "do").strip()

import urllib.request as _release_urllib_request
def release_ollama_generate(vault_path, prompt, model=None, host=None, temperature=None, timeout=None):
    """Implement the ollama generate operation for the local LLM Wiki workflow."""
    cfg = release_load_config(vault_path) if "release_load_config" in globals() else {}
    ollama = cfg.get("ollama", {}) if isinstance(cfg, dict) else {}
    host = host or ollama.get("host", "http://localhost:11434")
    model = model or ollama.get("model", "llama3.2:3b")
    payload = {"model": model, "prompt": prompt, "stream": False}
    data = _release_json.dumps(payload).encode("utf-8")
    req = _release_urllib_request.Request(host.rstrip("/") + "/api/generate", data=data, headers={"Content-Type": "application/json"}, method="POST")
    start = _release_time.time()
    try:
        with _release_urllib_request.urlopen(req, timeout=int(timeout or ollama.get("timeout_seconds", 120))) as resp:
            raw = resp.read().decode("utf-8", errors="replace")
        obj = _release_json.loads(raw)
        elapsed = _release_time.time() - start
        answer = obj.get("response", "")
        return {"ok": True, "answer": answer, "elapsed_seconds": elapsed, "raw": obj, "model": model, "host": host}
    except Exception as exc:
        return {"ok": False, "answer": "", "error": str(exc), "elapsed_seconds": _release_time.time() - start, "model": model, "host": host}


# ---------------------------------------------------------------------------
# Release system prompt, robust page reading and better agentic follow-up
# ---------------------------------------------------------------------------

def release_agent_system_prompt(vault_path="wiki_vault"):
    """Implement the agent system prompt operation for the local LLM Wiki workflow."""
    caps = release_runtime_capabilities(vault_path) if "release_runtime_capabilities" in globals() else {}
    manifest = release_agentic_tool_manifest() if "release_agentic_tool_manifest" in globals() else {}
    return f"""# LLM Wiki MCP Agent System Prompt

You are operating as an agent inside the local LLM Wiki MCP.

You have access to a local Markdown wiki vault, AI sidecar notes, runtime capabilities, usage metrics, and safe read-only tools. You are not limited to answering from memory. Use the evidence pack and tool outputs supplied to you.

## What you can access

- Maintained Markdown wiki pages in `wiki_vault/wiki/`
- AI sidecar notes in `wiki_vault/notes/`
- Runtime config and capabilities
- Usage and token metrics
- Runtime journal
- Candidate links and graph information
- Search-with-notes results
- Full page reads for relevant pages

## Current capabilities

```json
{_release_json.dumps(caps, indent=2, ensure_ascii=False)[:6000]}
```

## Allowed read-only agent tools

```json
{_release_json.dumps(manifest, indent=2, ensure_ascii=False)}
```

## Behaviour rules

1. Prefer current tool/runtime state over stale documentation.
2. When a search result names a relevant page, read the full page before giving a detailed answer.
3. If the user asks what pages say, quote/summarise the pages actually read.
4. Do not say "I could call a tool next" when the evidence pack already contains the necessary tool output.
5. If evidence is missing, name the missing page/tool precisely.
6. Never mutate the wiki during agentic ask unless the user explicitly uses an apply/edit command.
7. Cite evidence by page title or tool name in plain text.
"""

def release_page_lookup(vault_path, title):
    """Implement the page lookup operation for the local LLM Wiki workflow."""
    wiki = _ProjectPath(vault_path) / "wiki"
    if not wiki.exists():
        return None
    raw = str(title or "").replace(".md", "").strip().lower()
    raw_norm = raw.replace("source - ", "").replace("_", " ").replace("-", " ")
    candidates = []
    for p in wiki.glob("*.md"):
        stem = p.stem
        s = stem.lower()
        s_norm = s.replace("source - ", "").replace("_", " ").replace("-", " ")
        score = 0
        if s == raw or s_norm == raw_norm:
            score += 100
        if raw and raw in s:
            score += 40
        if raw_norm and raw_norm in s_norm:
            score += 40
        for tok in raw_norm.split():
            if len(tok) > 2 and tok in s_norm:
                score += 5
        if score:
            candidates.append((score, p))
    if candidates:
        return sorted(candidates, key=lambda x: (-x[0], x[1].name))[0][1]
    direct = wiki / (str(title) if str(title).endswith(".md") else str(title) + ".md")
    return direct if direct.exists() else None

def release_read_page_tool(vault_path, title):
    """Implement the read page tool operation for the local LLM Wiki workflow."""
    path = release_page_lookup(vault_path, title)
    if not path or not path.exists():
        return {"ok": False, "title": title, "error": "page not found"}
    text = path.read_text(encoding="utf-8", errors="replace")
    note = ""
    try:
        note = release_read_note(vault_path, path.stem)
    except Exception:
        note = ""
    return {
        "ok": True,
        "title": path.stem,
        "path": str(path),
        "text": text[:12000],
        "note": note[:6000],
        "chars": len(text),
    }

def release_list_pages_tool(vault_path, limit=300):
    """Implement the list pages tool operation for the local LLM Wiki workflow."""
    wiki = _ProjectPath(vault_path) / "wiki"
    pages = sorted([p.stem for p in wiki.glob("*.md")]) if wiki.exists() else []
    return {"ok": True, "pages": pages[:int(limit)], "count": len(pages)}

def release_find_page_for_question(vault_path, question, search_result=None):
    """Implement the find page for question operation for the local LLM Wiki workflow."""
    q = str(question or "").lower()
    # Explicit architecture/doc terms should prefer those source pages.
    priority_terms = [
        ("architecture", "Source - ARCHITECTURE"),
        ("function", "Source - FUNCTION_LIST"),
        ("methodology", "Source - METHODOLOGY"),
        ("skills", "Source - SKILLS"),
        ("limitations", "Source - 04_AGENTIC_ANSWER_QUALITY"),
        ("known issues", "Source - 04_AGENTIC_ANSWER_QUALITY"),
        ("agentic", "Source - 04_AGENTIC_ASK_AND_MEMORY"),
        ("notes", "Source - 05_NOTES_GRAPHS_AND_MAINTENANCE"),
        ("runtime", "Source - RUNTIME_EXECUTION_JOURNAL"),
        ("capabilities", "Source - CAPABILITIES_COMMAND"),
    ]
    for term, title in priority_terms:
        if term in q:
            p = release_page_lookup(vault_path, title)
            if p:
                return p.stem
    if isinstance(search_result, dict) and search_result.get("results"):
        return search_result["results"][0].get("title")
    return None

def release_execute_agent_tool_by_name(vault_path, name, args, question, state=None):
    """Implement the execute agent tool by name operation for the local LLM Wiki workflow."""
    args = args or {}
    state = state or {}
    q = args.get("query") or question
    title = args.get("title") or args.get("page") or state.get("focus_page")
    if name == "wiki_search_with_notes":
        return release_search_with_notes(vault_path, q, int(args.get("limit", 8)))
    if name == "wiki_read_page":
        return release_read_page_tool(vault_path, title or q)
    if name == "wiki_usage_snapshot":
        return release_wiki_usage_snapshot(vault_path)
    if name == "wiki_self_access_snapshot":
        return release_self_access_snapshot(vault_path)
    if name == "wiki_capabilities":
        return release_runtime_capabilities(vault_path)
    if name == "wiki_candidate_links":
        return {"title": title, "candidate_links": release_find_candidate_links(vault_path, title or q, int(args.get("limit", 12)))}
    if name == "wiki_notes_status":
        return release_notes_status(vault_path)
    if name == "wiki_list_pages":
        return release_list_pages_tool(vault_path, int(args.get("limit", 300)))
    if name == "wiki_limitations_report":
        return release_limitations_report(vault_path)
    return {"ok": False, "error": f"Tool not allowed or unknown: {name}"}

def release_agentic_ask(vault_path="wiki_vault", question="", max_tool_calls=7, dry_run=False, synthesize=True, protocol_mode=False):
    """Build or execute the local LLM ask workflow for agentic ask."""
    original_question = question
    question = release_repair_query_text(question) if "release_repair_query_text" in globals() else str(question or "")
    state, trace, evidence = {}, [], []

    selected = release_agentic_select_tools(question) if "release_agentic_select_tools" in globals() else ["wiki_search_with_notes", "wiki_capabilities"]
    if "limitation" in question and "wiki_limitations_report" not in selected:
        selected.insert(0, "wiki_limitations_report")
    if "wiki_search_with_notes" not in selected:
        selected.insert(0, "wiki_search_with_notes")
    if "wiki_capabilities" not in selected:
        selected.append("wiki_capabilities")

    # First pass.
    search_result = None
    for tool in selected[:int(max_tool_calls)]:
        result = release_execute_agent_tool_by_name(vault_path, tool, {"query": question}, question, state)
        trace.append({"tool": tool, "arguments": {"query": question}, "result_type": type(result).__name__})
        evidence.append({"tool": tool, "result": result})
        if tool == "wiki_search_with_notes":
            search_result = result
            focus = release_find_page_for_question(vault_path, question, result)
            if focus:
                state["focus_page"] = focus

    # Mandatory follow-up full page read when the question asks about pages/docs/architecture or we have a focus page.
    wants_page_detail = any(x in question.lower() for x in ["what pages say", "page say", "architecture", "document", "docs", "source", "tell me what pages"])
    if state.get("focus_page") and (wants_page_detail or "wiki_read_page" not in [t["tool"] for t in trace]):
        result = release_execute_agent_tool_by_name(vault_path, "wiki_read_page", {"title": state["focus_page"]}, question, state)
        trace.append({"tool": "wiki_read_page", "arguments": {"title": state["focus_page"]}, "result_type": type(result).__name__, "follow_up": True})
        evidence.append({"tool": "wiki_read_page", "result": result})

    manifest = release_agentic_tool_manifest() if "release_agentic_tool_manifest" in globals() else {}
    system_prompt = release_agent_system_prompt(vault_path)
    prompt_parts = [
        system_prompt,
        "",
        "# Agentic Ask Evidence Pack",
        f"Original question: {original_question}",
        f"Repaired question: {question}",
        "",
        "## Tool Trace",
        _release_json.dumps(trace, indent=2),
        "",
        "## Evidence",
    ]
    for item in evidence:
        prompt_parts.append("```json\n" + _release_json.dumps(item, indent=2, ensure_ascii=False)[:10000] + "\n```")
    prompt_parts.append("""
## Final Answer Instructions

Answer the user's question using the evidence above.
If full page text was read, explain what that page says.
Prefer current runtime/tool state over stale docs.
Do not say you need to call a tool if the tool result is already present.
Mention page/tool evidence used.
""")
    prompt = "\n".join(prompt_parts)

    runtime_event = release_record_runtime_event(
        vault_path, "ask-agentic", original_question, [t["tool"] for t in trace],
        True, None, {"prompt_tokens_estimate": release_estimate_tokens(prompt)}, {"dry_run": dry_run, "protocol_mode": protocol_mode, "synthesize": synthesize}
    )

    if dry_run or not synthesize:
        return {"ok": True, "question": original_question, "repaired_question": question, "dry_run": True, "tools_called": [t["tool"] for t in trace], "trace": trace, "evidence": evidence, "llm_prompt": prompt, "runtime_event": runtime_event, "answer": None}

    llm = release_ollama_generate(vault_path, prompt) if "release_ollama_generate" in globals() else {"ok": False, "error": "release_ollama_generate missing", "answer": ""}
    answer = llm.get("answer") if llm.get("ok") else f"Agentic evidence was built, but Ollama synthesis failed: {llm.get('error')}"
    metrics = release_record_ask_turn_with_metrics(vault_path, original_question, prompt, answer, [e.get("tool") for e in evidence], llm.get("elapsed_seconds"), {"agentic": True, "ollama": llm}) if "release_record_ask_turn_with_metrics" in globals() else {}
    return {"ok": bool(llm.get("ok")), "question": original_question, "repaired_question": question, "answer": answer, "tools_called": [t["tool"] for t in trace], "trace": trace, "evidence": evidence, "llm": llm, "metrics": metrics, "runtime_event": runtime_event, "llm_prompt": prompt}


# ---------------------------------------------------------------------------
# Release unified MCP-style tool exposure and ask mode helpers
# ---------------------------------------------------------------------------

def release_get_config_value(config, dotted_key, default=None):
    """Read, write, or update local wiki configuration for get config value."""
    cur = config
    for part in dotted_key.split("."):
        if not isinstance(cur, dict) or part not in cur:
            return default
        cur = cur[part]
    return cur

def release_current_ask_mode(vault_path="wiki_vault"):
    """Build or execute the local LLM ask workflow for current ask mode."""
    cfg = release_load_config(vault_path)
    return release_get_config_value(cfg, "ask.mode", "agentic")

def release_set_ask_mode(vault_path, mode):
    """Build or execute the local LLM ask workflow for set ask mode."""
    mode = str(mode or "").strip().lower()
    if mode in {"agentic", "tools", "tool", "on"}:
        mode = "agentic"
    elif mode in {"plain", "normal", "non-agentic", "non_agentic", "off", "simple"}:
        mode = "plain"
    else:
        raise ValueError("ask mode must be 'agentic' or 'plain'")
    cfg = release_set_config_value(vault_path, "ask.mode", mode)
    release_set_config_value(vault_path, "agentic.enabled", mode == "agentic")
    return cfg

def release_mcp_tool_catalog():
    """Implement the mcp tool catalog operation for the local LLM Wiki workflow."""
    return [
        {
            "name": "wiki_ask",
            "description": "Ask the wiki. Uses agentic tool-calling by default unless ask.mode is plain.",
            "handler": "tool_wiki_ask",
            "safety": "read",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "vault_path": {"type": "string"},
                    "question": {"type": "string"},
                    "mode": {"type": "string", "enum": ["agentic", "plain"]},
                    "max_tool_calls": {"type": "integer"},
                    "dry_run": {"type": "boolean"}
                },
                "required": ["question"]
            }
        },
        {
            "name": "wiki_ask_agentic",
            "description": "Agentic ask: use safe read-only tools, read relevant pages, then synthesize with Ollama.",
            "handler": "tool_wiki_ask_agentic",
            "safety": "read",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "vault_path": {"type": "string"},
                    "question": {"type": "string"},
                    "max_tool_calls": {"type": "integer"},
                    "dry_run": {"type": "boolean"},
                    "synthesize": {"type": "boolean"},
                    "protocol_mode": {"type": "boolean"}
                },
                "required": ["question"]
            }
        },
        {
            "name": "wiki_ask_plain",
            "description": "Plain non-agentic ask: retrieve context and synthesize without additional tool-calling.",
            "handler": "tool_wiki_ask_plain",
            "safety": "read",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "vault_path": {"type": "string"},
                    "question": {"type": "string"},
                    "top_k": {"type": "integer"},
                    "dry_run": {"type": "boolean"}
                },
                "required": ["question"]
            }
        },
        {
            "name": "wiki_set_ask_mode",
            "description": "Set CLI/default ask mode: agentic or plain.",
            "handler": "tool_wiki_set_ask_mode",
            "safety": "config",
            "inputSchema": {
                "type": "object",
                "properties": {
                    "vault_path": {"type": "string"},
                    "mode": {"type": "string", "enum": ["agentic", "plain"]}
                },
                "required": ["mode"]
            }
        },
        {
            "name": "wiki_tool_catalog",
            "description": "Return the unified MCP-style tool catalogue.",
            "handler": "tool_wiki_tool_catalog",
            "safety": "read",
            "inputSchema": {"type": "object", "properties": {}, "required": []}
        }
    ]

def release_tool_defs_from_catalog():
    """Implement the tool defs from catalog operation for the local LLM Wiki workflow."""
    return [{"name": x["name"], "description": x["description"], "inputSchema": x["inputSchema"]} for x in release_mcp_tool_catalog()]

def release_plain_ask(vault_path="wiki_vault", question="", top_k=5, dry_run=False):
    """Build or execute the local LLM ask workflow for plain ask."""
    search = release_search_with_notes(vault_path, question, int(top_k)) if "release_search_with_notes" in globals() else {"results": []}
    evidence = []
    for row in search.get("results", [])[:int(top_k)]:
        title = row.get("title")
        page_text = ""
        if "release_read_page_tool" in globals():
            page = release_read_page_tool(vault_path, title)
            page_text = page.get("text", "") if isinstance(page, dict) else ""
        evidence.append({"title": title, "score": row.get("score"), "snippet": row.get("snippet"), "text": page_text[:4000]})
    prompt = "# Plain LLM Wiki Ask\n\nQuestion: " + str(question) + "\n\n## Retrieved Evidence\n"
    prompt += _release_json.dumps(evidence, indent=2, ensure_ascii=False)[:16000]
    prompt += "\n\nAnswer using only the evidence above. If evidence is weak, say so."
    if dry_run:
        return {"ok": True, "mode": "plain", "question": question, "search": search, "llm_prompt": prompt, "answer": None}
    llm = release_ollama_generate(vault_path, prompt) if "release_ollama_generate" in globals() else {"ok": False, "error": "release_ollama_generate missing", "answer": ""}
    answer = llm.get("answer") if llm.get("ok") else "Plain evidence was built, but Ollama synthesis failed: " + str(llm.get("error"))
    return {"ok": bool(llm.get("ok")), "mode": "plain", "question": question, "search": search, "answer": answer, "llm": llm, "llm_prompt": prompt}

def release_ask(vault_path="wiki_vault", question="", mode=None, max_tool_calls=7, dry_run=False):
    """Build or execute the local LLM ask workflow for ask."""
    mode = (mode or release_current_ask_mode(vault_path) or "agentic").lower()
    if mode == "plain":
        return release_plain_ask(vault_path, question, top_k=5, dry_run=dry_run)
    return release_agentic_ask(vault_path, question, max_tool_calls=max_tool_calls, dry_run=dry_run, synthesize=not dry_run, protocol_mode=False)

def release_tool_catalog_markdown():
    """Implement the tool catalog markdown operation for the local LLM Wiki workflow."""
    lines = ["# Release MCP Tool Catalogue", ""]
    for item in release_mcp_tool_catalog():
        lines += [f"## {item['name']}", "", item["description"], "", f"- Handler: `{item['handler']}`", f"- Safety: `{item['safety']}`", "", "```json", _release_json.dumps(item["inputSchema"], indent=2), "```", ""]
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Release compatibility and audit helpers
# ---------------------------------------------------------------------------

def release_all_tools_summary():
    """Implement the all tools summary operation for the local LLM Wiki workflow."""
    return {
        "catalog_size": len(release_mcp_tool_catalog()) if "release_mcp_tool_catalog" in globals() else 0,
        "default_ask_mode": release_current_ask_mode() if "release_current_ask_mode" in globals() else "unknown",
        "legacy_aliases_enabled": True,
        "notes_enabled": True,
        "runtime_journal_enabled": True,
        "recursive_history_enabled": True,
    }

def release_legacy_tool_aliases():
    """Implement the legacy tool aliases operation for the local LLM Wiki workflow."""
    return {
        "ask-agentic": "ask",
        "ask-simple": "ask-plain",
        "catalog": "tool-catalog",
    }


# Fuller planner and stats-via-LLM path
def release_tool_groups():
    """Implement the tool groups operation for the local LLM Wiki workflow."""
    return {
        "retrieval":["wiki_search_with_notes","wiki_read_page","wiki_list_pages","wiki_candidate_links","wiki_notes_status"],
        "diagnostics":["wiki_usage_snapshot","wiki_self_access_snapshot","wiki_llm_metrics","wiki_runtime_journal","wiki_capabilities","wiki_command_tool_map","wiki_reconcile_docs","wiki_tool_catalog"],
        "memory":["wiki_ask_history","wiki_compress_history","wiki_recursive_context"],
        "visual":["wiki_runtime_graph"],
    }

def release_execute_agent_tool_by_name(vault_path,name,args,question,state=None):
    """Implement the execute agent tool by name operation for the local LLM Wiki workflow."""
    args=args or {}; state=state or {}; q=args.get("query") or question
    title=args.get("title") or args.get("page") or state.get("focus_page")
    if name=="wiki_search_with_notes": return release_search_with_notes(vault_path,q,int(args.get("limit",8)))
    if name=="wiki_read_page": return release_read_page_tool(vault_path,title or q)
    if name=="wiki_list_pages": return release_list_pages_tool(vault_path,int(args.get("limit",300)))
    if name=="wiki_candidate_links": return {"title":title,"candidate_links":release_find_candidate_links(vault_path,title or q,int(args.get("limit",12)))}
    if name=="wiki_notes_status": return release_notes_status(vault_path)
    if name=="wiki_usage_snapshot": return release_wiki_usage_snapshot(vault_path)
    if name=="wiki_self_access_snapshot": return release_self_access_snapshot(vault_path)
    if name=="wiki_llm_metrics": return release_metrics_summary(vault_path,int(args.get("limit",100)))
    if name=="wiki_runtime_journal": return {"events":release_load_runtime_journal(vault_path,int(args.get("limit",20)))}
    if name=="wiki_capabilities": return release_runtime_capabilities(vault_path)
    if name=="wiki_command_tool_map": return release_command_tool_map()
    if name=="wiki_reconcile_docs": return release_reconcile_docs(vault_path)
    if name=="wiki_tool_catalog": return {"catalog":release_mcp_tool_catalog() if "release_mcp_tool_catalog" in globals() else {}}
    if name=="wiki_ask_history": return {"history":release_load_history(vault_path,int(args.get("limit",10)))}
    if name=="wiki_compress_history": return release_compress_history_if_needed(vault_path)
    if name=="wiki_recursive_context": return release_recursive_context_with_memory(vault_path,q)
    if name=="wiki_runtime_graph": return {"mermaid":release_runtime_graph_mermaid()}
    if name=="wiki_limitations_report": return release_limitations_report(vault_path)
    return {"ok":False,"error":"unknown tool "+str(name)}

def release_plan_tools(question,max_tool_calls=10):
    """Implement the plan tools operation for the local LLM Wiki workflow."""
    q=str(question or "").lower()
    tools=["wiki_search_with_notes","wiki_capabilities"]
    if any(w in q for w in ["stats","diagnostic","diagnostics","health","usage","token","tokens","metrics","info tools","self stats"]):
        tools += ["wiki_usage_snapshot","wiki_self_access_snapshot","wiki_llm_metrics","wiki_runtime_journal"]
    if any(w in q for w in ["tool","tools","catalog","command","mapping","what can you"]):
        tools += ["wiki_command_tool_map","wiki_tool_catalog"]
    if any(w in q for w in ["history","recursive","memory","compress"]):
        tools += ["wiki_ask_history","wiki_compress_history","wiki_recursive_context"]
    if any(w in q for w in ["notes","links","linkages","connect","related"]):
        tools += ["wiki_notes_status","wiki_candidate_links"]
    if any(w in q for w in ["graph","visual"]):
        tools += ["wiki_runtime_graph"]
    if any(w in q for w in ["limitation","limitations","weakness","missing"]):
        tools += ["wiki_limitations_report","wiki_reconcile_docs"]
    if any(w in q for w in ["architecture","docs","pages","page","document","source"]):
        tools += ["wiki_list_pages"]
    seen=[]
    for t in tools:
        if t not in seen: seen.append(t)
    return seen[:int(max_tool_calls)]

def release_agentic_ask(vault_path="wiki_vault",question="",max_tool_calls=10,dry_run=False,synthesize=True,protocol_mode=False):
    """Build or execute the local LLM ask workflow for agentic ask."""
    original=question
    question=release_repair_query_text(question) if "release_repair_query_text" in globals() else str(question or "")
    state={}; trace=[]; evidence=[]
    for tool in release_plan_tools(question,max_tool_calls):
        res=release_execute_agent_tool_by_name(vault_path,tool,{"query":question},question,state)
        trace.append({"tool":tool,"result_type":type(res).__name__})
        evidence.append({"tool":tool,"result":res})
        if tool=="wiki_search_with_notes" and isinstance(res,dict) and res.get("results"):
            state["focus_page"]=res["results"][0].get("title")
    if state.get("focus_page") and any(x in question.lower() for x in ["architecture","docs","pages","what pages","source"]):
        res=release_execute_agent_tool_by_name(vault_path,"wiki_read_page",{"title":state["focus_page"]},question,state)
        trace.append({"tool":"wiki_read_page","result_type":type(res).__name__,"follow_up":True})
        evidence.append({"tool":"wiki_read_page","result":res})
    prompt = (release_agent_system_prompt(vault_path) if "release_agent_system_prompt" in globals() else "# LLM Wiki MCP Agent") + "\n\n"
    prompt += "# Release Agentic Ask\n\n"
    prompt += "Answer using the tool evidence. For stats/diagnostics, explain the evidence in natural language rather than dumping raw output.\n\n"
    prompt += "Original question: "+str(original)+"\nRepaired question: "+str(question)+"\n\n## Tool Trace\n```json\n"+_release_json.dumps(trace,indent=2,ensure_ascii=False)+"\n```\n\n## Evidence\n"
    for item in evidence:
        prompt += "\n### "+item["tool"]+"\n```json\n"+_release_json.dumps(item["result"],indent=2,ensure_ascii=False)[:8000]+"\n```\n"
    if dry_run or not synthesize:
        return {"ok":True,"question":original,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm_prompt":prompt,"answer":None}
    llm=release_ollama_generate(vault_path,prompt) if "release_ollama_generate" in globals() else {"ok":False,"error":"missing ollama","answer":""}
    ans=llm.get("answer") if llm.get("ok") else "Agentic evidence was built, but Ollama synthesis failed: "+str(llm.get("error"))
    return {"ok":bool(llm.get("ok")),"question":original,"answer":ans,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm":llm,"llm_prompt":prompt}

def release_ask(vault_path="wiki_vault",question="",mode=None,max_tool_calls=10,dry_run=False):
    """Build or execute the local LLM ask workflow for ask."""
    mode=(mode or release_current_ask_mode(vault_path) or "agentic").lower()
    if mode=="plain": return release_plain_ask(vault_path,question,top_k=5,dry_run=dry_run)
    return release_agentic_ask(vault_path,question,max_tool_calls=max_tool_calls,dry_run=dry_run,synthesize=not dry_run)


# ---------------------------------------------------------------------------
# Guaranteed answer fallback for ask / ask-agentic
# ---------------------------------------------------------------------------

def release_summarise_evidence_fallback(question, trace=None, evidence=None, llm_error=None):
    """Create a useful answer even if Ollama returns nothing/fails.

    This prevents the user-facing "No answer produced" dead end.
    """
    trace = trace or []
    evidence = evidence or []
    tools = [t.get("tool") for t in trace if isinstance(t, dict) and t.get("tool")]
    lines = []
    lines.append("I could not get a final synthesized response from the LLM, but I did gather evidence and can still give a useful fallback answer.")
    lines.append("")
    if llm_error:
        lines.append(f"LLM synthesis issue: `{llm_error}`")
        lines.append("")
    if tools:
        lines.append("Tools/evidence gathered:")
        for t in tools:
            lines.append(f"- `{t}`")
        lines.append("")
    # Extract useful high-level facts from common tool outputs.
    for item in evidence:
        tool = item.get("tool") if isinstance(item, dict) else None
        result = item.get("result") if isinstance(item, dict) else None
        if not isinstance(result, dict):
            continue
        if tool == "wiki_search_with_notes":
            results = result.get("results", [])
            lines.append(f"Search found {len(results)} likely page(s).")
            for r in results[:5]:
                title = r.get("title")
                score = r.get("score")
                matched = ", ".join(r.get("matched_terms", [])[:6])
                lines.append(f"- {title} (score {score}; matched: {matched})")
            lines.append("")
        elif tool == "wiki_capabilities":
            llm = result.get("llm", {})
            retrieval = result.get("retrieval", {})
            runtime = result.get("runtime", {})
            lines.append("Current capability snapshot:")
            if llm:
                lines.append(f"- LLM provider/model: {llm.get('provider')} / {llm.get('model')}")
                lines.append(f"- Ollama host: {llm.get('host')}")
                lines.append(f"- Recursive ask: {llm.get('recursive_ask')}")
            if retrieval:
                lines.append(f"- SQLite FTS enabled: {retrieval.get('sqlite_fts')}")
                lines.append(f"- Markdown source of truth: {retrieval.get('markdown_source_of_truth')}")
            if runtime:
                lines.append(f"- Runtime journaling: {runtime.get('runtime_journaling')}")
                lines.append(f"- Token metrics: {runtime.get('token_metrics')}")
            lines.append("")
        elif tool == "wiki_usage_snapshot":
            lines.append("Usage snapshot:")
            for k in ["pages","words","tokens_estimate","ask_history_turns","ask_history_tokens_estimate"]:
                if k in result:
                    lines.append(f"- {k}: {result.get(k)}")
            lines.append("")
        elif tool == "wiki_llm_metrics":
            lines.append("LLM/token metrics:")
            for k in ["count","total_prompt_tokens_estimate","total_answer_tokens_estimate","avg_output_tokens_per_second","avg_total_tokens_per_second"]:
                if k in result:
                    lines.append(f"- {k}: {result.get(k)}")
            lines.append("")
        elif tool == "wiki_read_page":
            title = result.get("title")
            text = result.get("text","")
            lines.append(f"Read page: `{title}`")
            if text:
                snippet = text.strip().replace("\\n", " ")[:800]
                lines.append(f"- Page excerpt: {snippet}")
            lines.append("")
    lines.append("Why this happens sometimes:")
    lines.append("- Ollama may be unreachable, slow, or using the wrong host.")
    lines.append("- The model may return an empty response.")
    lines.append("- A query may produce weak search results.")
    lines.append("- Earlier versions printed `No answer produced` instead of falling back to evidence.")
    lines.append("")
    lines.append("Release changes this behaviour: `ask` should always return either an LLM answer or this evidence-based fallback.")
    return "\\n".join(lines)

def release_agentic_ask(vault_path="wiki_vault",question="",max_tool_calls=10,dry_run=False,synthesize=True,protocol_mode=False):
    """Build or execute the local LLM ask workflow for agentic ask."""
    original=question
    question=release_repair_query_text(question) if "release_repair_query_text" in globals() else str(question or "")
    state={}; trace=[]; evidence=[]
    for tool in release_plan_tools(question,max_tool_calls) if "release_plan_tools" in globals() else ["wiki_search_with_notes","wiki_capabilities"]:
        res=release_execute_agent_tool_by_name(vault_path,tool,{"query":question},question,state)
        trace.append({"tool":tool,"result_type":type(res).__name__})
        evidence.append({"tool":tool,"result":res})
        if tool=="wiki_search_with_notes" and isinstance(res,dict) and res.get("results"):
            state["focus_page"]=res["results"][0].get("title")
    if state.get("focus_page") and any(x in question.lower() for x in ["architecture","docs","pages","what pages","source"]):
        res=release_execute_agent_tool_by_name(vault_path,"wiki_read_page",{"title":state["focus_page"]},question,state)
        trace.append({"tool":"wiki_read_page","result_type":type(res).__name__,"follow_up":True})
        evidence.append({"tool":"wiki_read_page","result":res})
    prompt=(release_agent_system_prompt(vault_path) if "release_agent_system_prompt" in globals() else "# LLM Wiki MCP Agent") + "\\n\\n"
    prompt+="# Release Agentic Ask\\n\\nAlways answer from available evidence. If evidence is weak, explain what is weak and still provide a best-effort answer.\\n\\n"
    prompt+="Original question: "+str(original)+"\\nRepaired question: "+str(question)+"\\n\\n## Tool Trace\\n```json\\n"+_release_json.dumps(trace,indent=2,ensure_ascii=False)+"\\n```\\n\\n## Evidence\\n"
    for item in evidence:
        prompt+="\\n### "+item["tool"]+"\\n```json\\n"+_release_json.dumps(item["result"],indent=2,ensure_ascii=False)[:8000]+"\\n```\\n"
    if dry_run or not synthesize:
        return {"ok":True,"question":original,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm_prompt":prompt,"answer":None}
    llm=release_ollama_generate(vault_path,prompt) if "release_ollama_generate" in globals() else {"ok":False,"error":"missing ollama","answer":""}
    answer=(llm.get("answer") or "").strip()
    if not answer:
        answer=release_summarise_evidence_fallback(original,trace,evidence,llm.get("error") or "empty LLM response")
    return {"ok":bool(llm.get("ok")),"question":original,"answer":answer,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm":llm,"llm_prompt":prompt}

def release_ask(vault_path="wiki_vault",question="",mode=None,max_tool_calls=10,dry_run=False):
    """Build or execute the local LLM ask workflow for ask."""
    mode=(mode or release_current_ask_mode(vault_path) or "agentic").lower()
    if mode=="plain":
        out=release_plain_ask(vault_path,question,top_k=5,dry_run=dry_run)
        if not dry_run and not (out.get("answer") or "").strip():
            out["answer"]=release_summarise_evidence_fallback(question,[],[{"tool":"plain_search","result":out.get("search",{})}],out.get("llm",{}).get("error"))
        return out
    return release_agentic_ask(vault_path,question,max_tool_calls=max_tool_calls,dry_run=dry_run,synthesize=not dry_run)


# ---------------------------------------------------------------------------
# Full MCP catalogue and agent-visible tool surface
# ---------------------------------------------------------------------------

def release_full_tool_catalog():
    """Full read-oriented tool catalogue used for MCP exposure and agent prompts.

    Earlier versions exposed only a small front-door catalogue, while many real
    CLI/internal functions were available. Release makes the catalogue explicit.
    """
    tools = [
        ("wiki_ask", "Ask the wiki using current ask mode, agentic by default.", "read"),
        ("wiki_ask_agentic", "Agentic ask with safe tools, evidence pack, Ollama synthesis and fallback.", "read"),
        ("wiki_ask_plain", "Plain retrieve-and-synthesize ask without tool loop.", "read"),
        ("wiki_set_ask_mode", "Set default ask mode: agentic or plain.", "config"),
        ("wiki_tool_catalog", "Return the full Release tool catalogue.", "read"),
        ("wiki_planner_tools", "Return planner tool groups.", "read"),

        ("wiki_search_with_notes", "Search source pages plus AI sidecar notes with hit-centred snippets.", "read"),
        ("wiki_search_following_pages", "Search long books and return the matched page plus N following pages/excerpt.", "read"),
        ("wiki_retrieve_articles", "Retrieve top matching articles/context for a prompt.", "read"),
        ("wiki_combined_page_context", "Return source page plus sidecar notes.", "read"),
        ("wiki_read_page", "Read a full wiki page.", "read"),
        ("wiki_list_pages", "List wiki pages.", "read"),
        ("wiki_candidate_links", "Find candidate links for a page.", "read"),
        ("wiki_notes_status", "Show AI note coverage.", "read"),
        ("wiki_notes_graph", "Render Mermaid graph from notes.", "read"),
        ("wiki_page_notes_iterate", "Create/update AI sidecar notes for a page.", "write_notes"),
        ("wiki_notes_iterate_all", "Iterate AI sidecar notes for many/all pages.", "write_notes"),

        ("wiki_usage_snapshot", "Return wiki/page/history usage stats.", "read"),
        ("wiki_self_access_snapshot", "Return config, usage, metrics and history state.", "read"),
        ("wiki_capabilities", "Return live runtime capabilities.", "read"),
        ("wiki_llm_metrics", "Return token/s and ask metrics.", "read"),
        ("wiki_runtime_journal", "Return recent runtime journal events.", "read"),
        ("wiki_record_runtime_event", "Record runtime event.", "write_log"),
        ("wiki_command_tool_map", "Return command-to-tool mapping.", "read"),
        ("wiki_reconcile_docs", "Check docs/tool surface consistency.", "read"),
        ("wiki_runtime_graph", "Return runtime graph Mermaid.", "read"),
        ("wiki_limitations_report", "Return known/fixed/remaining limitations.", "read"),
        ("wiki_agentic_tool_manifest", "Return the agentic safe tool manifest.", "read"),
        ("wiki_safe_self_dispatch", "Map natural language to safe read-only self tool.", "read"),

        ("wiki_lint", "Lint broken links, orphans and short pages.", "read"),
        ("wiki_repair_lint", "Repair lint issues; dry-run unless apply requested.", "maintenance"),
        ("wiki_mermaid_graph", "Render full wiki graph.", "read"),
        ("wiki_mermaid_neighbourhood", "Render page-centred graph.", "read"),
        ("wiki_health_report", "Return health report.", "read"),
        ("wiki_status", "Return status summary.", "read"),

        ("wiki_config", "Read/update config depending on args.", "config"),
        ("wiki_config_load", "Load config.", "read"),
        ("wiki_config_set", "Set config key.", "config"),
        ("wiki_config_test", "Test configured Ollama connection.", "read"),

        ("wiki_ask_history", "Return recent ask history.", "read"),
        ("wiki_compress_history", "Compress ask history if needed.", "write_log"),
        ("wiki_recursive_context", "Build recursive memory context.", "read"),
        ("wiki_record_ask_turn_with_metrics", "Record ask turn and metrics.", "write_log"),

        ("wiki_ingest_directory", "Ingest supported files from a directory or one file.", "write_wiki"),
        ("wiki_reindex", "Rebuild SQLite index.", "maintenance"),
        ("wiki_export_llms_txt", "Export llms.txt context.", "read"),
    ]
    out=[]
    for name, desc, safety in tools:
        out.append({
            "name": name,
            "description": desc,
            "safety": safety,
            "inputSchema": {
                "type": "object",
                "properties": {
                    "vault_path": {"type": "string"},
                    "question": {"type": "string"},
                    "query": {"type": "string"},
                    "title": {"type": "string"},
                    "mode": {"type": "string"},
                    "limit": {"type": "integer"},
                    "dry_run": {"type": "boolean"},
                    "apply": {"type": "boolean"}
                },
                "required": []
            }
        })
    return out

def release_mcp_tool_catalog():
    """Override a small catalogue with the full MCP tool catalogue."""
    return release_full_tool_catalog()

def release_tool_defs_from_catalog():
    """Implement the tool defs from catalog operation for the local LLM Wiki workflow."""
    return [{"name": t["name"], "description": t["description"], "inputSchema": t["inputSchema"]} for t in release_full_tool_catalog()]

def release_tool_catalog_markdown():
    """Implement the tool catalog markdown operation for the local LLM Wiki workflow."""
    lines = ["# Release Full MCP Tool Catalogue", "", f"Tool count: {len(release_full_tool_catalog())}", ""]
    by_safety={}
    for t in release_full_tool_catalog():
        by_safety.setdefault(t["safety"], []).append(t)
    for safety, items in sorted(by_safety.items()):
        lines.append(f"## {safety}")
        lines.append("")
        for t in items:
            lines.append(f"- `{t['name']}` — {t['description']}")
        lines.append("")
    return "\n".join(lines)

def release_agentic_tool_manifest():
    """Agent-visible manifest now mirrors the full read/safe subset."""
    manifest={}
    for t in release_full_tool_catalog():
        if t["safety"] in {"read"}:
            manifest[t["name"]] = t["description"]
    return manifest

def release_tool_catalog_summary():
    """Implement the tool catalog summary operation for the local LLM Wiki workflow."""
    return {
        "tool_count": len(release_full_tool_catalog()),
        "read_only_count": len([t for t in release_full_tool_catalog() if t["safety"] == "read"]),
        "catalog": release_full_tool_catalog(),
    }


# ---------------------------------------------------------------------------
# Intent-aware ask for short/contextual prompts
# ---------------------------------------------------------------------------

import re as _release_intent_re

def release_recent_user_questions(vault_path="wiki_vault", limit=5):
    """Implement the recent user questions operation for the local LLM Wiki workflow."""
    hist=[]
    try:
        rows=release_load_history(vault_path, limit)
        for row in rows:
            if isinstance(row,dict):
                q=row.get("question") or row.get("prompt") or row.get("user")
                a=row.get("answer") or row.get("response")
                if q: hist.append({"question":q,"answer":a})
    except Exception:
        pass
    return hist

def release_resolve_tool_name_intent(question):
    """Implement the resolve tool name intent operation for the local LLM Wiki workflow."""
    q=str(question or "").strip().strip("`").strip()
    low=q.lower()
    if "wiki_mermaid_neighbourhood" in low:
        target=low.replace("wiki_mermaid_neighbourhood","").replace("for","").replace("use","").strip()
        if not target or "random" in target or "example" in target:
            target="Overview"
        if "agi" in target:
            target="Source - BOOK_SELF_EXTENDING_AGENTIC_SYSTEMS"
        return {"intent":"run_tool","tool":"wiki_mermaid_neighbourhood","target":target,"resolved_question":f"Generate a Mermaid neighbourhood graph for {target}."}
    if "wiki_mermaid_graph" in low or ("global graph" in low and "runtime" not in low):
        return {"intent":"run_tool","tool":"wiki_mermaid_graph","target":None,"resolved_question":"Generate the global Mermaid wiki graph and explain what it shows."}
    if "wiki_runtime_graph" in low or "runtime graph" in low:
        return {"intent":"run_tool","tool":"wiki_runtime_graph","target":None,"resolved_question":"Generate the runtime Mermaid graph and explain what it shows."}
    return None

def release_is_short_contextual_query(question):
    """Implement the is short contextual query operation for the local LLM Wiki workflow."""
    q=str(question or "").strip().lower().strip("` ")
    if len(q.split()) <= 3:
        return True
    return any(q == p or q.startswith(p+" ") for p in ["write more","continue","go on","more","expand","you do it","do it","use it","show it","hello","hi","use graph","use mermaid"])

def release_resolve_intent(vault_path, question):
    """Implement the resolve intent operation for the local LLM Wiki workflow."""
    original=str(question or "")
    q=original.strip()
    low=q.lower().strip()
    tool=release_resolve_tool_name_intent(q)
    if tool:
        tool["original_question"]=original
        return tool
    recent=release_recent_user_questions(vault_path,5)
    last_q=recent[-1]["question"] if recent else ""
    if low in {"hello","hi","hiya"}:
        return {"intent":"conversation","original_question":original,"resolved_question":"Greet the user briefly, introduce the LLM Wiki MCP, and offer useful next commands.","context_hint":"Greeting, not a document search."}
    if low in {"write more","more","continue","go on","expand"} or low.startswith("write more"):
        subject=""
        if "about" in low:
            subject=q.split("about",1)[1].strip()
        elif last_q:
            subject=last_q
        else:
            subject="the LLM Wiki MCP and its tools"
        return {"intent":"expand_previous","original_question":original,"resolved_question":f"Write a fuller continuation about {subject}, using recent context and wiki evidence.","context_hint":f"Previous question: {last_q}"}
    if low in {"you do it","do it","use it","show it"}:
        return {"intent":"act_on_previous","original_question":original,"resolved_question":f"Carry out the likely requested action from the previous exchange. Previous question: {last_q}","context_hint":"Short imperative referring to prior context."}
    if "your tools" in low or "tools are" in low:
        return {"intent":"explain_tools","original_question":original,"resolved_question":"Explain available MCP tools using the full catalogue, planner groups and command map.","context_hint":"Prefer tool catalogue and command map."}
    if "graph store" in low:
        return {"intent":"explain_graph","original_question":original,"resolved_question":"Explain the graph-related storage and Mermaid graph tools in the LLM Wiki MCP.","context_hint":"Graph/store means wiki link graph, notes graph, runtime graph, Mermaid exports."}
    if "mermaid" in low and any(x in low for x in ["plot","show","graph","connections","link"]):
        return {"intent":"graph_request","original_question":original,"resolved_question":"Use Mermaid graph tools to show relevant wiki/page connections.","context_hint":"Route to graph and notes graph tools."}
    if release_is_short_contextual_query(q) and last_q:
        return {"intent":"contextual_followup","original_question":original,"resolved_question":f"Interpret as a follow-up to previous topic. Previous question: {last_q}. Current follow-up: {q}","context_hint":"Short follow-up expanded using ask history."}
    return {"intent":"normal","original_question":original,"resolved_question":original,"context_hint":""}

def release_mermaid_neighbourhood_fallback(vault_path, target):
    """Create graph-oriented output for mermaid neighbourhood fallback."""
    links=release_find_candidate_links(vault_path,target,10)
    safe=lambda s: _release_intent_re.sub(r"[^A-Za-z0-9_]", "_", str(s))
    lines=["```mermaid","flowchart TD"]
    src=safe(target)
    lines.append(f'    {src}["{target}"]')
    for row in links:
        node=safe(row.get("title"))
        lines.append(f'    {src} --> {node}["{row.get("title")}"]')
    lines.append("```")
    return {"title":target,"candidate_links":links,"mermaid":"\n".join(lines)}

def release_execute_agent_tool_by_name(vault_path,name,args,question,state=None):
    """Implement the execute agent tool by name operation for the local LLM Wiki workflow."""
    args=args or {}; state=state or {}; q=args.get("query") or question
    title=args.get("title") or args.get("page") or state.get("focus_page")
    if name=="wiki_mermaid_graph":
        if "release_mermaid_graph" in globals():
            return {"mermaid":release_mermaid_graph(vault_path)}
        return {"mermaid":release_notes_graph_mermaid(vault_path)}
    if name=="wiki_mermaid_neighbourhood":
        return release_mermaid_neighbourhood_fallback(vault_path,title or "Overview")
    # broad implementation
    if name=="wiki_search_with_notes": return release_search_with_notes(vault_path,q,int(args.get("limit",8)))
    if name=="wiki_read_page": return release_read_page_tool(vault_path,title or q)
    if name=="wiki_list_pages": return release_list_pages_tool(vault_path,int(args.get("limit",300)))
    if name=="wiki_candidate_links": return {"title":title,"candidate_links":release_find_candidate_links(vault_path,title or q,int(args.get("limit",12)))}
    if name=="wiki_notes_status": return release_notes_status(vault_path)
    if name=="wiki_notes_graph": return {"mermaid":release_notes_graph_mermaid(vault_path)}
    if name=="wiki_usage_snapshot": return release_wiki_usage_snapshot(vault_path)
    if name=="wiki_self_access_snapshot": return release_self_access_snapshot(vault_path)
    if name=="wiki_llm_metrics": return release_metrics_summary(vault_path,int(args.get("limit",100)))
    if name=="wiki_runtime_journal": return {"events":release_load_runtime_journal(vault_path,int(args.get("limit",20)))}
    if name=="wiki_capabilities": return release_runtime_capabilities(vault_path)
    if name=="wiki_command_tool_map": return release_command_tool_map()
    if name=="wiki_reconcile_docs": return release_reconcile_docs(vault_path)
    if name=="wiki_tool_catalog": return {"catalog":release_mcp_tool_catalog() if "release_mcp_tool_catalog" in globals() else release_full_tool_catalog()}
    if name=="wiki_runtime_graph": return {"mermaid":release_runtime_graph_mermaid()}
    if name=="wiki_limitations_report": return release_limitations_report(vault_path)
    if name=="wiki_ask_history": return {"history":release_load_history(vault_path,int(args.get("limit",10)))}
    return {"ok":False,"error":"unknown tool "+str(name)}

def release_plan_tools(question,max_tool_calls=12,intent=None):
    """Implement the plan tools operation for the local LLM Wiki workflow."""
    intent=intent or {}; q=str(question or "").lower()
    tools=["wiki_search_with_notes","wiki_capabilities"]
    if intent.get("intent")=="conversation":
        return ["wiki_capabilities","wiki_tool_catalog"][:max_tool_calls]
    if intent.get("intent")=="explain_tools":
        tools+=["wiki_tool_catalog","wiki_command_tool_map","wiki_planner_tools"]
    if intent.get("intent")=="run_tool":
        if intent.get("tool"): tools.append(intent.get("tool"))
        if intent.get("tool")=="wiki_mermaid_neighbourhood": tools+=["wiki_list_pages","wiki_candidate_links"]
        return list(dict.fromkeys(tools))[:max_tool_calls]
    if intent.get("intent") in {"expand_previous","contextual_followup","act_on_previous"}:
        tools+=["wiki_ask_history","wiki_tool_catalog"]
    if intent.get("intent")=="graph_request":
        tools+=["wiki_notes_status","wiki_candidate_links","wiki_notes_graph","wiki_mermaid_graph"]
    if any(w in q for w in ["stats","diagnostic","health","usage","token","metrics"]):
        tools+=["wiki_usage_snapshot","wiki_self_access_snapshot","wiki_llm_metrics","wiki_runtime_journal"]
    if any(w in q for w in ["tool","tools","catalog","command","mapping","what can you"]):
        tools+=["wiki_command_tool_map","wiki_tool_catalog","wiki_planner_tools"]
    if any(w in q for w in ["history","recursive","memory","compress"]):
        tools+=["wiki_ask_history"]
    if any(w in q for w in ["notes","links","linkages","connect","related","page connections"]):
        tools+=["wiki_notes_status","wiki_candidate_links","wiki_notes_graph"]
    if any(w in q for w in ["graph","mermaid","visual","plot"]):
        tools+=["wiki_mermaid_graph","wiki_runtime_graph"]
    if any(w in q for w in ["limitation","weakness","missing"]):
        tools+=["wiki_limitations_report","wiki_reconcile_docs"]
    if any(w in q for w in ["architecture","docs","pages","page","document","source"]):
        tools+=["wiki_list_pages"]
    seen=[]
    for t in tools:
        if t not in seen: seen.append(t)
    return seen[:int(max_tool_calls)]

def release_agentic_ask(vault_path="wiki_vault",question="",max_tool_calls=12,dry_run=False,synthesize=True,protocol_mode=False):
    """Build or execute the local LLM ask workflow for agentic ask."""
    original=question
    intent=release_resolve_intent(vault_path,question)
    resolved=intent.get("resolved_question") or str(question or "")
    repaired=release_repair_query_text(resolved) if "release_repair_query_text" in globals() else resolved
    state={}; trace=[]; evidence=[]
    for tool in release_plan_tools(repaired,max_tool_calls,intent):
        args={"query":repaired}
        if intent.get("tool")==tool and intent.get("target"): args["title"]=intent.get("target")
        res=release_execute_agent_tool_by_name(vault_path,tool,args,repaired,state)
        trace.append({"tool":tool,"result_type":type(res).__name__})
        evidence.append({"tool":tool,"result":res})
        if tool=="wiki_search_with_notes" and isinstance(res,dict) and res.get("results"):
            state["focus_page"]=res["results"][0].get("title")
    if state.get("focus_page") and any(x in repaired.lower() for x in ["architecture","docs","pages","what pages","source","based upon","what is this wiki"]):
        res=release_execute_agent_tool_by_name(vault_path,"wiki_read_page",{"title":state["focus_page"]},repaired,state)
        trace.append({"tool":"wiki_read_page","result_type":type(res).__name__,"follow_up":True})
        evidence.append({"tool":"wiki_read_page","result":res})
    prompt=(release_agent_system_prompt(vault_path) if "release_agent_system_prompt" in globals() else "# LLM Wiki MCP Agent")+"\n\n"
    prompt+="# Release Intent-Aware Agentic Ask\n\nRespect resolved intent. Short prompts may refer to prior context. Tool-name prompts should use/explain tool output.\n\n"
    prompt+="Original question: "+str(original)+"\nResolved intent: "+_release_json.dumps(intent,indent=2,ensure_ascii=False)+"\nResolved question: "+str(repaired)+"\n\n"
    prompt+="## Tool Trace\n```json\n"+_release_json.dumps(trace,indent=2,ensure_ascii=False)+"\n```\n\n## Evidence\n"
    for item in evidence:
        prompt+="\n### "+item["tool"]+"\n```json\n"+_release_json.dumps(item["result"],indent=2,ensure_ascii=False)[:9000]+"\n```\n"
    if dry_run or not synthesize:
        return {"ok":True,"question":original,"resolved_intent":intent,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm_prompt":prompt,"answer":None}
    llm=release_ollama_generate(vault_path,prompt) if "release_ollama_generate" in globals() else {"ok":False,"error":"missing ollama","answer":""}
    ans=(llm.get("answer") or "").strip()
    if not ans:
        ans=release_summarise_evidence_fallback(original,trace,evidence,llm.get("error") or "empty LLM response")
    return {"ok":bool(llm.get("ok")),"question":original,"resolved_intent":intent,"answer":ans,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm":llm,"llm_prompt":prompt}

def release_ask(vault_path="wiki_vault",question="",mode=None,max_tool_calls=12,dry_run=False):
    """Build or execute the local LLM ask workflow for ask."""
    mode=(mode or release_current_ask_mode(vault_path) or "agentic").lower()
    if mode=="plain":
        return release_plain_ask(vault_path,question,top_k=5,dry_run=dry_run)
    return release_agentic_ask(vault_path,question,max_tool_calls=max_tool_calls,dry_run=dry_run,synthesize=not dry_run)


# ---------------------------------------------------------------------------
# Release ingest provenance, startup commands, GitHub history
# ---------------------------------------------------------------------------
import hashlib as _release_hashlib
import importlib.util as _release_importlib_util
import sqlite3 as _release_sqlite3
import re as _release_re
from datetime import datetime as _release_datetime, timezone as _release_timezone

def release_provenance_db_path(vault_path="wiki_vault"):
    """Implement the provenance db path operation for the local LLM Wiki workflow."""
    return _ProjectPath(vault_path) / "llm_wiki_provenance.sqlite"

def release_file_sha256(path, chunk_size=1024*1024):
    """Implement the file sha256 operation for the local LLM Wiki workflow."""
    h=_release_hashlib.sha256()
    with open(path,"rb") as f:
        for chunk in iter(lambda:f.read(chunk_size), b""):
            h.update(chunk)
    return h.hexdigest()

def release_file_stat_record(path):
    """Implement the file stat record operation for the local LLM Wiki workflow."""
    p=_ProjectPath(path).resolve(); st=p.stat()
    return {"filename":p.name,"path":str(p),"directory":str(p.parent),"size_bytes":int(st.st_size),"mtime_ns":int(st.st_mtime_ns),"mtime_iso":_release_datetime.fromtimestamp(st.st_mtime,tz=_release_timezone.utc).isoformat(),"sha256":release_file_sha256(p)}

def release_connect_provenance(vault_path="wiki_vault"):
    """Implement the connect provenance operation for the local LLM Wiki workflow."""
    db=release_provenance_db_path(vault_path); db.parent.mkdir(parents=True,exist_ok=True)
    con=_release_sqlite3.connect(str(db)); con.row_factory=_release_sqlite3.Row
    con.execute("PRAGMA foreign_keys=ON"); return con

def release_init_provenance_schema(vault_path="wiki_vault"):
    """Implement the init provenance schema operation for the local LLM Wiki workflow."""
    con=release_connect_provenance(vault_path)
    con.executescript("""
    CREATE TABLE IF NOT EXISTS directories(id INTEGER PRIMARY KEY AUTOINCREMENT,path TEXT NOT NULL UNIQUE,first_seen_utc TEXT NOT NULL,last_seen_utc TEXT NOT NULL);
    CREATE TABLE IF NOT EXISTS source_files(id INTEGER PRIMARY KEY AUTOINCREMENT,directory_id INTEGER NOT NULL,filename TEXT NOT NULL,full_path TEXT NOT NULL UNIQUE,current_sha256 TEXT NOT NULL,size_bytes INTEGER NOT NULL,mtime_ns INTEGER NOT NULL,mtime_iso TEXT NOT NULL,wiki_title TEXT,wiki_path TEXT,first_ingested_utc TEXT NOT NULL,last_ingested_utc TEXT NOT NULL,last_seen_utc TEXT NOT NULL,ingest_count INTEGER NOT NULL DEFAULT 1,status TEXT NOT NULL DEFAULT 'current',FOREIGN KEY(directory_id) REFERENCES directories(id));
    CREATE TABLE IF NOT EXISTS file_versions(id INTEGER PRIMARY KEY AUTOINCREMENT,source_file_id INTEGER NOT NULL,sha256 TEXT NOT NULL,size_bytes INTEGER NOT NULL,mtime_ns INTEGER NOT NULL,mtime_iso TEXT NOT NULL,ingested_utc TEXT NOT NULL,wiki_title TEXT,wiki_path TEXT,event TEXT NOT NULL,FOREIGN KEY(source_file_id) REFERENCES source_files(id));
    CREATE INDEX IF NOT EXISTS idx_source_files_filename ON source_files(filename);
    CREATE INDEX IF NOT EXISTS idx_source_files_hash ON source_files(current_sha256);
    CREATE INDEX IF NOT EXISTS idx_source_files_dir ON source_files(directory_id);
    """)
    con.commit(); con.close(); return str(release_provenance_db_path(vault_path))

def release_get_or_create_directory(con,directory):
    """Implement the get or create directory operation for the local LLM Wiki workflow."""
    now=_release_datetime.now(_release_timezone.utc).isoformat()
    row=con.execute("SELECT id FROM directories WHERE path=?",(directory,)).fetchone()
    if row:
        con.execute("UPDATE directories SET last_seen_utc=? WHERE id=?",(now,row["id"])); return row["id"]
    return con.execute("INSERT INTO directories(path,first_seen_utc,last_seen_utc) VALUES(?,?,?)",(directory,now,now)).lastrowid

def release_provenance_status_for_file(vault_path,path):
    """Implement the provenance status for file operation for the local LLM Wiki workflow."""
    release_init_provenance_schema(vault_path); rec=release_file_stat_record(path)
    con=release_connect_provenance(vault_path); row=con.execute("SELECT * FROM source_files WHERE full_path=?",(rec["path"],)).fetchone(); con.close()
    if not row: return {"action":"new","reason":"file path not seen before","record":rec}
    if int(row["mtime_ns"])==rec["mtime_ns"] and row["current_sha256"]==rec["sha256"]:
        return {"action":"skip","reason":"same datestamp and hash","record":rec,"source_file_id":row["id"],"wiki_title":row["wiki_title"]}
    return {"action":"update","reason":"file changed","record":rec,"source_file_id":row["id"],"wiki_title":row["wiki_title"],"previous_sha256":row["current_sha256"]}

def release_record_ingest_result(vault_path,path,wiki_title=None,wiki_path=None,action=None):
    """Ingest source content into Markdown wiki pages for record ingest result."""
    release_init_provenance_schema(vault_path); rec=release_file_stat_record(path); now=_release_datetime.now(_release_timezone.utc).isoformat()
    con=release_connect_provenance(vault_path); did=release_get_or_create_directory(con,rec["directory"])
    row=con.execute("SELECT * FROM source_files WHERE full_path=?",(rec["path"],)).fetchone()
    if row is None:
        sfid=con.execute("INSERT INTO source_files(directory_id,filename,full_path,current_sha256,size_bytes,mtime_ns,mtime_iso,wiki_title,wiki_path,first_ingested_utc,last_ingested_utc,last_seen_utc,ingest_count,status) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(did,rec["filename"],rec["path"],rec["sha256"],rec["size_bytes"],rec["mtime_ns"],rec["mtime_iso"],wiki_title,wiki_path,now,now,now,1,"current")).lastrowid
        event=action or "new"
    else:
        sfid=row["id"]; event=action or ("unchanged" if row["current_sha256"]==rec["sha256"] and int(row["mtime_ns"])==rec["mtime_ns"] else "updated")
        con.execute("UPDATE source_files SET directory_id=?,filename=?,current_sha256=?,size_bytes=?,mtime_ns=?,mtime_iso=?,wiki_title=COALESCE(?,wiki_title),wiki_path=COALESCE(?,wiki_path),last_ingested_utc=?,last_seen_utc=?,ingest_count=ingest_count+1,status='current' WHERE id=?",(did,rec["filename"],rec["sha256"],rec["size_bytes"],rec["mtime_ns"],rec["mtime_iso"],wiki_title,wiki_path,now,now,sfid))
    con.execute("INSERT INTO file_versions(source_file_id,sha256,size_bytes,mtime_ns,mtime_iso,ingested_utc,wiki_title,wiki_path,event) VALUES(?,?,?,?,?,?,?,?,?)",(sfid,rec["sha256"],rec["size_bytes"],rec["mtime_ns"],rec["mtime_iso"],now,wiki_title,wiki_path,event))
    con.commit(); con.close()
    return {"ok":True,"source_file_id":sfid,"event":event,"record":rec,"wiki_title":wiki_title,"wiki_path":wiki_path}

def release_provenance_summary(vault_path="wiki_vault",limit=20):
    """Implement the provenance summary operation for the local LLM Wiki workflow."""
    release_init_provenance_schema(vault_path); con=release_connect_provenance(vault_path)
    dirs=con.execute("SELECT COUNT(*) n FROM directories").fetchone()["n"]; files=con.execute("SELECT COUNT(*) n FROM source_files").fetchone()["n"]; versions=con.execute("SELECT COUNT(*) n FROM file_versions").fetchone()["n"]
    recent=[dict(r) for r in con.execute("SELECT sf.filename,d.path directory,sf.current_sha256,sf.mtime_iso,sf.wiki_title,sf.status FROM source_files sf JOIN directories d ON d.id=sf.directory_id ORDER BY sf.last_seen_utc DESC LIMIT ?",(int(limit),)).fetchall()]
    con.close(); return {"directories":dirs,"source_files":files,"file_versions":versions,"recent":recent,"db":str(release_provenance_db_path(vault_path))}

def release_startup_commands(vault_path="wiki_vault"):
    """Implement the startup commands operation for the local LLM Wiki workflow."""
    cfg=release_load_config(vault_path); return cfg.get("startup_commands",[]) or cfg.get("startup",{}).get("commands",[]) or []

def release_set_startup_commands(vault_path,commands):
    """Implement the set startup commands operation for the local LLM Wiki workflow."""
    if isinstance(commands,str): commands=[c.strip() for c in commands.split(";") if c.strip()]
    return release_set_config_value(vault_path,"startup_commands",commands)

def release_default_startup_commands():
    """Implement the default startup commands operation for the local LLM Wiki workflow."""
    return ["/ingest ./docs","/notes-all"]

def release_startup_commands_markdown(vault_path="wiki_vault"):
    """Implement the startup commands markdown operation for the local LLM Wiki workflow."""
    cmds=release_startup_commands(vault_path)
    return "# Startup Commands\n\n" + ("\n".join(f"{i}. `{c}`" for i,c in enumerate(cmds,1)) if cmds else "No startup commands configured.")

def release_markdown_public_sanitise(text):
    """Implement the markdown public sanitise operation for the local LLM Wiki workflow."""
    text=_release_re.sub(r"<local-project>`)]+/","<local-project>/",text)
    text=_release_re.sub(r"C:\\\\Users\\\\[^\s`)]+\\\\","<local-project>\\\\",text)
    text=_release_re.sub(r"/home/[^\s`)]+/","<home>/",text)
    return text

def release_build_development_history(project_root=None):
    """Inspect runtime history, token metrics, or usage state for build development history."""
    root=_ProjectPath(project_root) if project_root else _ProjectPath.cwd(); docs=root/"docs"; docs.mkdir(exist_ok=True)
    sections=["# Development History","","This document consolidates the project history for a generic public audience. Local machine-specific paths have been removed or replaced with placeholders.",""]
    for p in sorted(docs.glob("V*_*.md")):
        if "SOURCE" in p.name: continue
        txt=release_markdown_public_sanitise(p.read_text(encoding="utf-8",errors="replace"))
        sections += [f"## {p.stem.replace('_',' ')}","",txt[:6000].strip(),""]
    out=docs/"DEVELOPMENT_HISTORY.md"; out.write_text("\n".join(sections),encoding="utf-8")
    return {"ok":True,"path":str(out),"sections":len(sections)}


# ---------------------------------------------------------------------------
# Release provenance-aware directory/single-file ingest summary
# ---------------------------------------------------------------------------

def release_supported_ingest_suffixes():
    """Ingest source content into Markdown wiki pages for supported ingest suffixes."""
    return {".md", ".markdown", ".txt", ".rtf", ".pdf", ".docx", ".doc"}

def release_pdf_extraction_status():
    """Return the optional PDF extraction status for the current runtime.

    The release keeps PDF parsing optional so the base CLI stays lightweight.
    This check is safe to run at shell startup because it only asks Python
    whether the `pypdf` module is importable in the active environment.
    """
    available = _release_importlib_util.find_spec("pypdf") is not None
    return {
        "available": available,
        "module": "pypdf",
        "message": (
            "PDF text extraction is available."
            if available
            else "PDF text extraction is unavailable. Install with `pip install -e .[docs,mcp]` or `pip install pypdf`."
        ),
    }

def release_wiki_page_needs_pdf_reextract(vault_path, wiki_title):
    """Return True when an old PDF placeholder page should be refreshed.

    If a PDF was ingested before `pypdf` was installed, the generated page may
    contain a diagnostic placeholder instead of actual transcript text. The PDF
    source file can remain unchanged, so hash/mtime provenance would normally
    skip it. This helper lets the planner force a refresh once PDF extraction is
    available.
    """
    if not wiki_title:
        return False
    safe_title = _release_re.sub(r"[^A-Za-z0-9 _\\-]", "", str(wiki_title)).strip().replace("  ", " ")
    page_path = _ProjectPath(vault_path) / "wiki" / (safe_title + ".md")
    if not page_path.exists():
        return False
    body = page_path.read_text(encoding="utf-8", errors="replace")
    markers = [
        "PDF text extraction unavailable or failed",
        "No module named 'pypdf'",
        'No module named "pypdf"',
        "Install the document extras",
        "Install pypdf for better extraction",
    ]
    return any(marker in body for marker in markers)

def release_scan_ingest_directory(directory, recursive=True):
    """Return supported files for a directory or a single ingest target file.

    The interactive `/ingest` command accepts either a directory (`/ingest .`)
    or one explicit file (`/ingest ./notes/idea.rtf`).  Treating a file path as
    a one-item scan keeps provenance, summary output, and reindexing behaviour
    identical for both workflows.
    """
    root=_ProjectPath(directory).expanduser()
    if not root.exists():
        return []
    if root.is_file():
        return [root] if root.suffix.lower() in release_supported_ingest_suffixes() else []
    pattern="**/*" if recursive else "*"
    files=[]
    for p in root.glob(pattern):
        if p.is_file() and p.suffix.lower() in release_supported_ingest_suffixes():
            files.append(p)
    return sorted(files)

def release_wiki_title_for_source_file(path, base_dir=None):
    """Implement the wiki title for source file operation for the local LLM Wiki workflow."""
    p=_ProjectPath(path)
    if base_dir:
        try:
            rel=p.resolve().relative_to(_ProjectPath(base_dir).resolve())
            prefix=str(rel.parent).replace("/", " - ").replace("\\", " - ")
            stem=p.stem
            return f"Source - {prefix} - {stem}" if prefix and prefix!="." else f"Source - {stem}"
        except Exception:
            pass
    return f"Source - {p.stem}"

def release_strip_rtf_to_text(raw):
    """Convert simple Rich Text Format content into plain text without extras.

    This lightweight extractor is intentionally dependency-free. It handles the
    common RTF emitted by word processors well enough for search/ask ingestion:
    escaped hex bytes, Unicode escapes, paragraph markers, and ordinary control
    words are normalised into readable text. Complex embedded objects are
    ignored; users can convert very complex RTF files to Markdown/text first.
    """
    text=str(raw or "")

    def _hex(match):
        """Document the `_hex` function used by the LLM Wiki MCP release."""
        try:
            return bytes.fromhex(match.group(1)).decode("latin-1")
        except Exception:
            return " "

    def _unicode(match):
        """Document the `_unicode` function used by the LLM Wiki MCP release."""
        try:
            value=int(match.group(1))
            if value < 0:
                value += 65536
            return chr(value)
        except Exception:
            return " "

    text=_release_re.sub(r"\\'([0-9a-fA-F]{2})", _hex, text)
    text=_release_re.sub(r"\\u(-?\d+)\??", _unicode, text)
    text=_release_re.sub(r"\\(par|line|tab)\b ?", "\n", text)
    text=_release_re.sub(r"\\[{}\\]", lambda m: m.group(0)[1:], text)
    text=_release_re.sub(r"\\[a-zA-Z]+-?\d* ?", "", text)
    text=_release_re.sub(r"[{}]", "", text)
    text=_release_re.sub(r"\n{3,}", "\n\n", text)
    text=_release_re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def release_extract_text_from_source_file(path):
    """Extract readable text from source files used by provenance-aware ingest.

    The interactive shell command `ingest .` uses the provenance-aware ingest
    path. Keep rich document extraction here, rather than only in the top-level
    `ingest-dir` CLI path, so both commands can absorb TXT/RTF/PDF/DOCX content. PDF
    extraction reads the embedded text layer; scanned image-only PDFs still need
    OCR or pre-conversion to text/Markdown.
    """
    p=_ProjectPath(path)
    suffix=p.suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        return p.read_text(encoding="utf-8", errors="replace")
    if suffix == ".rtf":
        return release_strip_rtf_to_text(p.read_text(encoding="utf-8", errors="replace"))
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore
            reader=PdfReader(str(p))
            total=len(reader.pages)
            parts=[]
            for idx,page in enumerate(reader.pages,1):
                page_text=page.extract_text() or ""
                parts.append(f"--- Page {idx} of {total} ---\n\n{page_text}")
            text="\n\n".join(parts).strip()
            if text:
                return text
            return f"[No extractable PDF text layer found in {p.name}. Convert scanned/image-only PDFs with OCR before ingesting.]"
        except Exception as exc:
            return f"[PDF text extraction unavailable or failed for {p.name}: {exc}. Install the document extras with `pip install -e .[docs,mcp]` or install `pypdf`.]"
    if suffix == ".docx":
        try:
            import docx  # type: ignore
            document=docx.Document(str(p))
            text="\n".join(paragraph.text for paragraph in document.paragraphs).strip()
            if text:
                return text
            return f"[No extractable DOCX paragraph text found in {p.name}.]"
        except Exception as exc:
            return f"[DOCX text extraction unavailable or failed for {p.name}: {exc}. Install the document extras with `pip install -e .[docs,mcp]` or install `python-docx`.]"
    if suffix == ".doc":
        return f"[Legacy .doc extraction is not built in for {p.name}. Convert to .docx, .pdf, .md, .rtf, or .txt before ingesting.]"
    return p.read_text(encoding="utf-8", errors="replace")

def release_ingest_directory_plan(vault_path, directory, recursive=True):
    """Return new/updated/skipped lists using mtime+SHA256 provenance."""
    release_init_provenance_schema(vault_path)
    files=release_scan_ingest_directory(directory, recursive=recursive)
    pdf_status=release_pdf_extraction_status()
    out={"directory":str(_ProjectPath(directory).resolve()),"new":[],"updated":[],"skipped":[],"errors":[]}
    for p in files:
        try:
            status=release_provenance_status_for_file(vault_path, p)
            action=status.get("action")
            row={"path":str(p.resolve()),"filename":p.name,"sha256":status.get("record",{}).get("sha256"),"mtime_iso":status.get("record",{}).get("mtime_iso"),"reason":status.get("reason"),"wiki_title":status.get("wiki_title") or release_wiki_title_for_source_file(p, directory)}
            if action=="skip":
                if p.suffix.lower()==".pdf" and pdf_status.get("available") and release_wiki_page_needs_pdf_reextract(vault_path, row.get("wiki_title")):
                    row["reason"]="PDF extraction is now available; refreshing previous placeholder page"
                    row["previous_sha256"]=status.get("record",{}).get("sha256")
                    out["updated"].append(row)
                else:
                    out["skipped"].append(row)
            elif action=="update":
                row["previous_sha256"]=status.get("previous_sha256")
                out["updated"].append(row)
            else:
                out["new"].append(row)
        except Exception as exc:
            out["errors"].append({"path":str(p),"error":str(exc)})
    out["counts"]={"new":len(out["new"]),"updated":len(out["updated"]),"skipped":len(out["skipped"]),"errors":len(out["errors"]),"total":len(files)}
    return out

def release_markdown_to_wiki_page(vault_path, source_path, title=None, base_dir=None):
    """Write a source file into the Markdown wiki and record provenance.

    This helper is used by the interactive shell command `ingest .`. It extracts
    text from Markdown, TXT, RTF, PDF, and DOCX sources, then stores the extracted
    body in a generated wiki page so basic `search`, `find`, and `ask` can all
    use the same absorbed content.
    """
    p=_ProjectPath(source_path)
    title=title or release_wiki_title_for_source_file(p, base_dir)
    wiki_dir=_ProjectPath(vault_path)/"wiki"
    wiki_dir.mkdir(parents=True, exist_ok=True)
    safe_title=_release_re.sub(r"[^A-Za-z0-9 _\\-]", "", title).strip().replace("  "," ")
    wiki_path=wiki_dir/(safe_title+".md")
    body=release_extract_text_from_source_file(p)
    content=f"# {title}\n\nSource file: `{p}`\nSource type: `{p.suffix.lower() or 'file'}`\n\n---\n\n{body}\n"
    wiki_path.write_text(content, encoding="utf-8")
    release_record_ingest_result(vault_path, p, wiki_title=title, wiki_path=str(wiki_path), action="updated_or_new")
    return {"title":title,"wiki_path":str(wiki_path)}

def release_ingest_directory_provenance(vault_path, directory, recursive=True, apply=True):
    """Provenance-aware ingest.

    Second and subsequent runs report skipped files as already ingested and only
    process new/updated files.
    """
    plan=release_ingest_directory_plan(vault_path, directory, recursive=recursive)
    processed=[]
    if apply:
        for row in plan["new"]+plan["updated"]:
            try:
                res=release_markdown_to_wiki_page(vault_path, row["path"], row.get("wiki_title"), directory)
                row["wiki_title"]=res["title"]; row["wiki_path"]=res["wiki_path"]
                processed.append(row)
            except Exception as exc:
                plan["errors"].append({"path":row["path"],"error":str(exc)})
    plan["processed"]=processed
    plan["counts"]["processed"]=len(processed)
    plan["summary"]=release_ingest_directory_summary_text(plan)
    return plan

def release_ingest_directory_summary_text(plan):
    """Ingest source content into Markdown wiki pages for ingest directory summary text."""
    c=plan.get("counts",{})
    lines=[]
    lines.append(f"Ingest scan: {plan.get('directory')}")
    lines.append(f"New files: {c.get('new',0)} | Updated files: {c.get('updated',0)} | Already ingested: {c.get('skipped',0)} | Errors: {c.get('errors',0)}")
    if plan.get("updated"):
        lines.append("")
        lines.append("Updated files:")
        for row in plan["updated"]:
            lines.append(f"- {row.get('path')} ({str(row.get('previous_sha256',''))[:12]} -> {str(row.get('sha256',''))[:12]})")
    if plan.get("new"):
        lines.append("")
        lines.append("New files:")
        for row in plan["new"]:
            lines.append(f"- {row.get('path')} ({str(row.get('sha256',''))[:12]})")
    if plan.get("skipped"):
        lines.append("")
        lines.append(f"{len(plan['skipped'])} files were already ingested and unchanged.")
    if plan.get("errors"):
        lines.append("")
        lines.append("Errors:")
        for row in plan["errors"]:
            lines.append(f"- {row.get('path')}: {row.get('error')}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Release fix: robust re-ingest summary after processing
# ---------------------------------------------------------------------------

def release_ingest_directory_provenance(vault_path, directory, recursive=True, apply=True):
    """Ingest source content into Markdown wiki pages for an ingest target.

    `directory` is kept as the public argument name for compatibility, but it
    may now be either a directory or a single supported file path.
    """
    plan=release_ingest_directory_plan(vault_path, directory, recursive=recursive)
    processed=[]
    target_path=_ProjectPath(directory).expanduser()
    base_dir=target_path.parent if target_path.is_file() else target_path
    if apply:
        for row in list(plan["new"])+list(plan["updated"]):
            try:
                res=release_markdown_to_wiki_page(vault_path, row["path"], row.get("wiki_title"), base_dir)
                row["wiki_title"]=res["title"]; row["wiki_path"]=res["wiki_path"]
                processed.append(row)
            except Exception as exc:
                plan["errors"].append({"path":row["path"],"error":str(exc)})
    plan["processed"]=processed
    plan["counts"]={
        "new":len(plan.get("new",[])),
        "updated":len(plan.get("updated",[])),
        "skipped":len(plan.get("skipped",[])),
        "errors":len(plan.get("errors",[])),
        "total":len(release_scan_ingest_directory(directory, recursive=recursive)),
        "processed":len(processed),
    }
    # The provenance-aware ingest path writes Markdown pages directly rather than
    # going through WikiStore.create_page(). Rebuild the SQLite/FTS index here so
    # the basic CLI `search` / `find` commands see newly ingested pages
    # immediately after `ingest .`. Agentic ask scans the Markdown/wiki evidence
    # path more broadly, so this bug could otherwise appear only in simple search.
    if apply and processed:
        try:
            WikiStore(WikiConfig(_ProjectPath(vault_path))).reindex()
            plan["counts"]["indexed_pages"] = len(list((_ProjectPath(vault_path) / "wiki").glob("*.md")))
        except Exception as exc:
            plan["errors"].append({"path": str(vault_path), "error": f"reindex after ingest failed: {exc}"})
            plan["counts"]["errors"] = len(plan.get("errors", []))
    plan["summary"]=release_ingest_directory_summary_text(plan)
    return plan


# ---------------------------------------------------------------------------
# Evidence-first answers for research/reference/history/config questions
# ---------------------------------------------------------------------------

def release_unescape_newlines(text):
    """Older fallbacks accidentally printed literal \\n. Render them as lines."""
    return str(text or "").replace("\\n", "\n")

def release_research_priority_titles(question):
    """Search or retrieve wiki content for research priority titles."""
    q=str(question or "").lower()
    titles=[]
    if any(w in q for w in ["inspiration","reference","references","karpathy","paper","post","based upon","based on"]):
        titles += [
            "Source - BOOK_REFERENCES_AND_INSPIRATION",
            "Source - REFERENCES",
            "Source - KARPATHY_AND_KB_CLI_COMPARISON",
            "Source - 06_REFERENCES_AND_RESEARCH",
            "Source - BOOK_SELF_EXTENDING_AGENTIC_SYSTEMS",
        ]
    if any(w in q for w in ["history","project history","development history"]):
        titles += ["Source - DEVELOPMENT_HISTORY"]
    if any(w in q for w in ["config","configuration","startup","ollama","wsl","host"]):
        titles += [
            "Source - 02_USER_GUIDE",
            "Source - INSTALL",
            "Source - 02_USER_GUIDE",
            "Source - Configuration",
        ]
    # preserve order, unique
    seen=[]
    for t in titles:
        if t not in seen:
            seen.append(t)
    return seen

def release_should_force_page_reads(question):
    """Implement the should force page reads operation for the local LLM Wiki workflow."""
    q=str(question or "").lower()
    return any(w in q for w in [
        "inspiration","reference","references","karpathy","paper","post","based upon","based on",
        "history","project history","development history","config","configuration","startup",
        "github","docs","research"
    ])

def release_extract_page_summary(title, text, question="", max_chars=1800):
    """Crude but useful extractive summary without LLM."""
    text=str(text or "")
    lines=[ln.strip() for ln in text.splitlines() if ln.strip()]
    q_terms=[w.lower() for w in re.findall(r"[A-Za-z0-9_]+", question or "") if len(w)>3]
    picked=[]
    # Prefer headings and bullet/reference/link lines.
    for ln in lines:
        low=ln.lower()
        if ln.startswith("#") or any(term in low for term in q_terms) or "http" in low or ln.startswith("- "):
            if ln not in picked:
                picked.append(ln)
        if sum(len(x)+1 for x in picked) > max_chars:
            break
    if not picked:
        picked=lines[:12]
    snippet="\n".join(picked)
    if len(snippet)>max_chars:
        snippet=snippet[:max_chars].rstrip()+"..."
    return f"### {title}\n\n{snippet}"

def release_clean_evidence_line(line, max_len=260):
    """Return one readable evidence sentence/line for deterministic fallback answers."""
    line=str(line or "").strip()
    line=_release_re.sub(r"^#+\s*", "", line)
    line=_release_re.sub(r"^[>*-]+\s*", "", line)
    line=_release_re.sub(r"`{1,3}", "", line)
    line=_release_re.sub(r"\s+", " ", line).strip()
    if len(line)>max_len:
        line=line[:max_len].rstrip()+"..."
    return line

def release_evidence_terms(question):
    """Extract meaningful query terms used to score fallback evidence lines."""
    stop={"about","what","when","where","which","would","could","should","please","tell","answer","your","from","with","that","this","have","into","does","ready","feel","using","user","wiki","work"}
    return [w.lower() for w in _release_re.findall(r"[A-Za-z0-9_]+", str(question or "")) if len(w)>3 and w.lower() not in stop]

def release_evidence_candidate_lines(text, question="", limit=8):
    """Pick compact, answer-worthy lines from a page for non-LLM fallback synthesis."""
    terms=release_evidence_terms(question)
    raw=[]
    for ln in str(text or "").splitlines():
        clean=release_clean_evidence_line(ln)
        if not clean or clean in raw:
            continue
        if len(clean)<24 and not any(term in clean.lower() for term in terms):
            continue
        low=clean.lower()
        score=sum(3 for term in terms if term in low)
        if any(word in low for word in ["install", "github", "release", "readme", "license", "mcp", "cli", "ollama", "ingest", "architecture", "local-first", "documentation", "provenance"]):
            score+=2
        if ln.lstrip().startswith("#"):
            score+=1
        if clean.startswith("|"):
            score-=3
        raw.append((score, clean))
    raw.sort(key=lambda item: item[0], reverse=True)
    picked=[]
    for score, clean in raw:
        if score<=0 and len(picked)>=3:
            continue
        if clean not in picked:
            picked.append(clean)
        if len(picked)>=limit:
            break
    return picked

def release_compose_best_effort_direct_answer(question, read_pages=None, search_hits=None, caps=None):
    """Compose a natural-language answer before listing references when Ollama is blank."""
    q=str(question or "").strip()
    low=q.lower()
    read_pages=read_pages or []
    search_hits=search_hits or []
    titles=[str(p.get("title") or "") for p in read_pages if isinstance(p,dict) and p.get("title")]
    if any(word in low for word in ["ready", "release", "github"]):
        docs=[t for t in titles if any(k in t.lower() for k in ["readme", "install", "cli", "mcp", "architecture", "product", "books", "function"])]
        doc_phrase="the available release documentation" if not docs else ", ".join(f"`{t}`" for t in docs[:5])
        return (
            "Yes — based on the wiki evidence, this looks broadly ready for a GitHub release, with the normal caveat that a final clean checkout, "
            "test run, README/install walkthrough, and secret/path scrub should still be done immediately before publishing. "
            f"The evidence includes {doc_phrase}, which indicates that the project has user-facing documentation, installation guidance, MCP/CLI references, "
            "and a product/architecture explanation rather than only internal change notes. If the LLM synthesis is unavailable, that does not by itself mean the release is not ready; "
            "it only means this answer is being composed from retrieved wiki evidence rather than from the final model pass."
        )
    if any(w in low for w in ["karpathy", "inspiration", "reference", "references", "based upon", "based on", "paper", "post"]):
        return (
            "The evidence points to Andrej Karpathy’s LLM Wiki pattern as a key inspiration: a local, file-based knowledge base that an LLM can read and use as maintained context. "
            "This project extends that idea with an MCP tool layer, a human CLI, Ollama-backed answering, provenance-aware ingestion, sidecar notes, graph tooling, and release-oriented documentation. "
            "So the lineage is not just retrieval over files; it is a more agent-ready version of the same maintained-context idea."
        )
    if any(w in low for w in ["config", "configuration", "startup", "ollama", "wsl", "model"]):
        return (
            "Configuration is handled through the local wiki vault configuration file and the CLI config commands. "
            "For normal local use, set the Ollama host with `/config host http://localhost:11434`, choose a default model with `/config model <model-name>`, "
            "and then ask questions directly at the `wiki>` prompt or with `/ask`. Smaller machines can choose a lighter default model while GPU machines can use a larger one."
        )
    if "history" in low:
        return (
            "The project history is represented as consolidated development documentation rather than only scattered version notes. "
            "The evidence suggests the release keeps a curated public-facing narrative of how the tool evolved, while avoiding private/local details that do not belong in a GitHub release."
        )
    # Generic answer synthesis from the best page lines.
    candidates=[]
    for page in read_pages[:4]:
        title=page.get("title") or "a retrieved page"
        text=page.get("text") or page.get("content") or ""
        for line in release_evidence_candidate_lines(text, q, limit=3):
            candidates.append((title,line))
    if candidates:
        first_title, first_line = candidates[0]
        parts=[f"Based on the retrieved wiki evidence, the best answer is that `{first_title}` is the strongest source for this question: {first_line}"]
        if len(candidates)>1:
            support=[]
            for title,line in candidates[1:4]:
                support.append(f"`{title}` adds that {line[0].lower()+line[1:] if line else line}")
            parts.append(" ".join(support))
        parts.append("I would treat this as a best-effort answer from the local wiki: useful, but still worth checking the cited pages if the question is release-critical or requires exact wording.")
        return " ".join(parts)
    if search_hits:
        hit_titles=[str(r.get("title") or "") for r in search_hits[:5] if isinstance(r,dict)]
        return "I found matching wiki pages but did not get enough full-page text to make a strong synthesis. The likely relevant pages are " + ", ".join(f"`{t}`" for t in hit_titles if t) + ". I would read those pages next before making a firm claim."
    return "I found runtime evidence but not enough page content to make a confident detailed answer. The safest answer is that the wiki has some relevant context, but the exact conclusion should be checked against the underlying pages."

def release_evidence_extract_answer(question, trace=None, evidence=None, llm_error=None):
    """Produce an answer-first fallback from evidence when Ollama returns blank or fails."""
    trace=trace or []
    evidence=evidence or []
    q=str(question or "")
    read_pages=[]
    search_hits=[]
    caps=None
    for item in evidence:
        tool=item.get("tool") if isinstance(item,dict) else None
        result=item.get("result") if isinstance(item,dict) else None
        if not isinstance(result,dict):
            continue
        if tool=="wiki_read_page" and result.get("ok", True):
            read_pages.append(result)
        elif tool=="wiki_search_with_notes":
            search_hits.extend(result.get("results",[])[:8])
        elif tool=="wiki_capabilities":
            caps=result
    lines=["## Direct answer", "", release_compose_best_effort_direct_answer(q, read_pages, search_hits, caps), ""]
    if llm_error:
        lines += ["> Note: Ollama synthesis returned `"+str(llm_error)+"`, so this answer was composed deterministically from retrieved wiki evidence. It is still an answer, not just a reference dump.", ""]
    # Put supporting evidence after the answer so the user gets a reply first.
    if read_pages:
        lines.append("## Evidence used")
        seen=[]
        for page in read_pages[:6]:
            title=page.get("title")
            if title and title not in seen:
                lines.append(f"- `{title}`")
                seen.append(title)
        lines.append("")
        lines.append("## Supporting notes")
        for page in read_pages[:3]:
            title=page.get("title") or "Retrieved page"
            text=page.get("text") or page.get("content") or ""
            picked=release_evidence_candidate_lines(text, q, limit=4)
            if picked:
                lines.append(f"### {title}")
                for item in picked[:4]:
                    lines.append(f"- {item}")
                lines.append("")
    elif search_hits:
        lines.append("## Search evidence")
        for r in search_hits[:6]:
            title=r.get("title") or "Untitled"
            matched=", ".join(r.get("matched_terms",[])[:8]) if isinstance(r.get("matched_terms"),list) else ""
            lines.append(f"- `{title}`" + (f" — matched: {matched}" if matched else ""))
        lines.append("")
    if caps:
        llm=caps.get("llm",{}) if isinstance(caps,dict) else {}
        retrieval=caps.get("retrieval",{}) if isinstance(caps,dict) else {}
        if llm or retrieval:
            lines.append("## Runtime context")
            if llm:
                lines.append(f"- LLM: {llm.get('provider')} / {llm.get('model')}")
            if retrieval:
                lines.append(f"- Retrieval: SQLite FTS={retrieval.get('sqlite_fts')}, markdown source of truth={retrieval.get('markdown_source_of_truth')}")
            lines.append("")
    used=[t.get("tool") for t in trace if isinstance(t,dict) and t.get("tool")]
    if used:
        lines.append("Tools used: "+", ".join(used))
    return "\n".join(lines).strip()+"\n"

def release_agentic_ask(vault_path="wiki_vault",question="",max_tool_calls=12,dry_run=False,synthesize=True,protocol_mode=False):
    """Build or execute the local LLM ask workflow for agentic ask."""
    original=question
    intent=release_resolve_intent(vault_path,question) if "release_resolve_intent" in globals() else {"intent":"normal","resolved_question":str(question or "")}
    resolved=intent.get("resolved_question") or str(question or "")
    repaired=release_repair_query_text(resolved) if "release_repair_query_text" in globals() else resolved
    state={}; trace=[]; evidence=[]
    planned=release_plan_tools(repaired,max_tool_calls,intent) if "release_plan_tools" in globals() else ["wiki_search_with_notes","wiki_capabilities"]
    # Force search/caps, then priority reads for research/reference/history/config questions.
    if "wiki_search_with_notes" not in planned:
        planned.insert(0,"wiki_search_with_notes")
    if "wiki_capabilities" not in planned:
        planned.append("wiki_capabilities")
    for title in reversed(release_research_priority_titles(original)):
        # insert page reads after search/capability base tools
        planned.append("wiki_read_page:"+title)
    # de-dupe simple tools but keep read page pseudo tools
    clean=[]
    for t in planned:
        if t.startswith("wiki_read_page:") or t not in clean:
            clean.append(t)
    planned=clean[:max_tool_calls+5]
    for tool in planned:
        if tool.startswith("wiki_read_page:"):
            title=tool.split(":",1)[1]
            res=release_execute_agent_tool_by_name(vault_path,"wiki_read_page",{"title":title},repaired,state)
            trace.append({"tool":"wiki_read_page","arguments":{"title":title},"result_type":type(res).__name__,"forced_read":True})
            evidence.append({"tool":"wiki_read_page","result":res})
            continue
        res=release_execute_agent_tool_by_name(vault_path,tool,{"query":repaired},repaired,state)
        trace.append({"tool":tool,"result_type":type(res).__name__})
        evidence.append({"tool":tool,"result":res})
        if tool=="wiki_search_with_notes" and isinstance(res,dict) and res.get("results"):
            state["focus_page"]=res["results"][0].get("title")
    # Always read top search result for research-ish questions if not already read.
    already={e.get("result",{}).get("title") for e in evidence if e.get("tool")=="wiki_read_page" and isinstance(e.get("result"),dict)}
    if release_should_force_page_reads(original) and state.get("focus_page") and state.get("focus_page") not in already:
        res=release_execute_agent_tool_by_name(vault_path,"wiki_read_page",{"title":state["focus_page"]},repaired,state)
        trace.append({"tool":"wiki_read_page","arguments":{"title":state["focus_page"]},"result_type":type(res).__name__,"follow_up":True})
        evidence.append({"tool":"wiki_read_page","result":res})
    # Build compact prompt to avoid empty Ollama responses caused by very large evidence.
    prompt=(release_agent_system_prompt(vault_path) if "release_agent_system_prompt" in globals() else "# LLM Wiki MCP Agent")+"\n\n"
    prompt+="# Release Evidence-First Ask\n\nUse the read pages and evidence. If evidence contains page text, answer from it. Be concise.\n\n"
    prompt+="Original question: "+str(original)+"\nResolved question: "+str(repaired)+"\nResolved intent: "+_release_json.dumps(intent,ensure_ascii=False)[:1200]+"\n\n"
    prompt+="## Evidence\n"
    for item in evidence:
        tool=item.get("tool")
        result=item.get("result")
        prompt+="\n### "+str(tool)+"\n```json\n"+_release_json.dumps(result,indent=2,ensure_ascii=False)[:4500]+"\n```\n"
    if dry_run or not synthesize:
        return {"ok":True,"question":original,"resolved_intent":intent,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm_prompt":prompt,"answer":None}
    llm=release_ollama_generate(vault_path,prompt) if "release_ollama_generate" in globals() else {"ok":False,"error":"missing ollama","answer":""}
    ans=release_unescape_newlines((llm.get("answer") or "").strip())
    if not ans:
        ans=release_evidence_extract_answer(original,trace,evidence,llm.get("error") or "empty LLM response")
    return {"ok":bool(llm.get("ok")),"question":original,"resolved_intent":intent,"answer":ans,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm":llm,"llm_prompt":prompt}

def release_ask(vault_path="wiki_vault",question="",mode=None,max_tool_calls=12,dry_run=False):
    """Build or execute the local LLM ask workflow for ask."""
    mode=(mode or release_current_ask_mode(vault_path) or "agentic").lower()
    if mode=="plain":
        out=release_plain_ask(vault_path,question,top_k=5,dry_run=dry_run)
        if not dry_run and not (out.get("answer") or "").strip():
            out["answer"]=release_evidence_extract_answer(question,[],[{"tool":"plain_search","result":out.get("search",{})}],out.get("llm",{}).get("error"))
        return out
    return release_agentic_ask(vault_path,question,max_tool_calls=max_tool_calls,dry_run=dry_run,synthesize=not dry_run)


# ---------------------------------------------------------------------------
# Answer-first fallback and discussion-aware follow-up context
# ---------------------------------------------------------------------------

def release_normalise_answer_text(text):
    """Implement the normalise answer text operation for the local LLM Wiki workflow."""
    return str(text or "").replace("\\n", "\n").strip()

def release_basic_question_intent(question):
    """Implement the basic question intent operation for the local LLM Wiki workflow."""
    q=str(question or "").strip().lower()
    q=_release_re.sub(r"[^a-z0-9\s:_\-\.]", "", q)
    if q in {"what are you","what r u","who are you","tell me what you are","tell me about yourself","tell me about yourself","what is this","what is this tool","describe yourself"}:
        return "identity"
    if q in {"what can you do","what tools do you have","tell me your tools","show tools","help me","how do i use you","how to use this"} or "what can you do" in q:
        return "capabilities"
    if "config" in q or "startup" in q or "ollama host" in q:
        return "config"
    if "followup" in q or "follow-up" in q or q in {"continue","write more","more","you do it","do it"}:
        return "followup"
    return None

def release_get_runtime_bits_from_evidence(evidence):
    """Implement the get runtime bits from evidence operation for the local LLM Wiki workflow."""
    caps={}
    catalog={}
    mapping={}
    history=[]
    for item in evidence or []:
        if not isinstance(item,dict):
            continue
        tool=item.get("tool")
        result=item.get("result")
        if not isinstance(result,dict):
            continue
        if tool=="wiki_capabilities":
            caps=result
        elif tool=="wiki_tool_catalog":
            catalog=result
        elif tool=="wiki_command_tool_map":
            mapping=result
        elif tool=="wiki_ask_history":
            history=result.get("history",[]) if isinstance(result.get("history",[]),list) else []
    return caps,catalog,mapping,history

def release_identity_answer(question,evidence=None,llm_error=None):
    """Implement the identity answer operation for the local LLM Wiki workflow."""
    caps,catalog,mapping,history=release_get_runtime_bits_from_evidence(evidence or [])
    llm=caps.get("llm",{}) if isinstance(caps,dict) else {}
    retrieval=caps.get("retrieval",{}) if isinstance(caps,dict) else {}
    runtime=caps.get("runtime",{}) if isinstance(caps,dict) else {}
    maintenance=caps.get("maintenance",{}) if isinstance(caps,dict) else {}
    host=llm.get("host") or "the configured Ollama host"
    model=llm.get("model") or "the configured local model"
    provider=llm.get("provider") or "Ollama/local LLM"
    lines=[
        "I am the **LLM Wiki MCP Agent** running over this local Markdown wiki vault.",
        "",
        "In practical terms, I am a command-line and MCP-accessible knowledge tool that can:",
        "",
        "- search and read maintained wiki pages;",
        "- retrieve context for questions;",
        "- use an Ollama LLM to synthesize answers;",
        "- inspect my own runtime/configuration/tool catalogue;",
        "- maintain AI sidecar notes and graph links;",
        "- lint, graph, ingest, and track source-file provenance.",
        "",
        f"Current LLM runtime: **{provider} / {model}** at `{host}`.",
    ]
    if retrieval:
        lines.append(f"Retrieval: SQLite FTS={retrieval.get('sqlite_fts')}, markdown source of truth={retrieval.get('markdown_source_of_truth')}.")
    if runtime:
        lines.append(f"Runtime: journaling={runtime.get('runtime_journaling')}, token metrics={runtime.get('token_metrics')}.")
    if maintenance:
        lines.append(f"Maintenance: lint/repair/graph/reconcile tools are exposed via the CLI/MCP.")
    if history:
        last=history[-1] if history else {}
        q=last.get("question") or last.get("prompt")
        if q:
            lines += ["", f"Recent discussion context: the last recorded question was `{q}`."]
    if llm_error:
        lines += ["", f"Note: the LLM returned `{llm_error}`, so this answer was generated from runtime/tool evidence rather than final model synthesis."]
    return "\n".join(lines)

def release_capabilities_answer(question,evidence=None,llm_error=None):
    """Implement the capabilities answer operation for the local LLM Wiki workflow."""
    caps,catalog,mapping,history=release_get_runtime_bits_from_evidence(evidence or [])
    lines=["Here is what I can do in this LLM Wiki MCP:"]
    lines += [
        "",
        "## Core use",
        "- `search <query>` — find relevant wiki snippets.",
        "- `read <title>` — read a full page.",
        "- `retrieve <query>` / `articles <query>` — build context packs.",
        "- `ask <question>` — agentic ask using wiki tools and Ollama.",
        "- `ask-agentic-dry <question>` — inspect the evidence pack without calling Ollama.",
        "",
        "## Maintenance",
        "- `ingest <directory>` — ingest files with provenance-aware skip/update logic.",
        "- `notes-all` — generate/update AI sidecar notes.",
        "- `lint`, `repair`, `graph`, `mermaid`, `map` — maintain and visualise the wiki.",
        "",
        "## Self-access",
        "- `stats`, `usage`, `tokens`, `capabilities`, `tool-count`, `tool-catalog`.",
        "- `provenance` — inspect the source-file provenance database.",
        "- `history-build` — rebuild consolidated public development history.",
    ]
    if caps:
        llm=caps.get("llm",{})
        lines += ["", f"Configured LLM: `{llm.get('model')}` via `{llm.get('host')}`."]
    if llm_error:
        lines += ["", f"Note: LLM synthesis returned `{llm_error}`; this is a built-in capabilities answer."]
    return "\n".join(lines)

def release_config_answer(question,evidence=None,llm_error=None):
    """Read, write, or update local wiki configuration for config answer."""
    caps,catalog,mapping,history=release_get_runtime_bits_from_evidence(evidence or [])
    llm=caps.get("llm",{}) if isinstance(caps,dict) else {}
    host=llm.get("host") or "http://localhost:11434"
    model=llm.get("model") or "llama3.2:3b"
    return "\n".join([
        "Configuration is stored in `wiki_vault/llm_wiki_config.json`.",
        "",
        "Useful commands:",
        "",
        "```text",
        "config",
        f"config host {host}",
        f"config model {model}",
        "config set <key> <value>",
        "config test",
        "config edit",
        "ask-mode agentic",
        "ask-mode plain",
        "/startup default",
        "/startup set /ingest ./docs; /notes-all",
        "```",
        "",
        "For your WSL/Ollama setup, the host reminder is:",
        "",
        "```text",
        "/config host http://localhost:11434",
        "```",
    ])

def release_discussion_followup_answer(question,evidence=None,llm_error=None):
    """Implement the discussion followup answer operation for the local LLM Wiki workflow."""
    caps,catalog,mapping,history=release_get_runtime_bits_from_evidence(evidence or [])
    lines=["I can treat this as a follow-up to the recent discussion."]
    if history:
        lines += ["", "Recent recorded questions:"]
        for row in history[-5:]:
            q=row.get("question") or row.get("prompt")
            if q:
                lines.append(f"- {q}")
        lines += ["", "I would use those recent turns plus the wiki search/read tools to continue the same topic rather than treating the short prompt as a standalone search."]
    else:
        lines += ["", "I did not find usable ask-history entries, so I would fall back to the wiki pages, tool catalogue and runtime capability snapshot."]
    return "\n".join(lines)

def release_fallback_answer_router(question, trace=None, evidence=None, llm_error=None):
    """Implement the fallback answer router operation for the local LLM Wiki workflow."""
    intent=release_basic_question_intent(question)
    if intent=="identity":
        return release_identity_answer(question,evidence,llm_error)
    if intent=="capabilities":
        return release_capabilities_answer(question,evidence,llm_error)
    if intent=="config":
        return release_config_answer(question,evidence,llm_error)
    if intent=="followup":
        return release_discussion_followup_answer(question,evidence,llm_error)
    # If release evidence fallback exists, use it for evidence-rich pages.
    if "release_evidence_extract_answer" in globals():
        try:
            return release_evidence_extract_answer(question, trace or [], evidence or [], llm_error)
        except Exception:
            pass
    # Last resort should still answer the user's basic question in natural language.
    return "\n".join([
        "I can answer, but the LLM did not return a final synthesis.",
        "",
        "Based on the available runtime evidence, this is an LLM Wiki MCP: a local Markdown wiki system with search, page reads, Ollama-backed asking, MCP tools, provenance-aware ingest, notes, graphs and maintenance commands.",
        "",
        "For more detail, try:",
        "",
        "```text",
        "/ask what are you",
        "capabilities",
        "tool-catalog",
        "ask-agentic-dry <question>",
        "```",
    ])

def release_plan_tools(question,max_tool_calls=14,intent=None):
    """Wider planner: basic questions still get useful runtime/history tools."""
    base=[]
    basic=release_basic_question_intent(question)
    if basic in {"identity","capabilities","config","followup"}:
        base=["wiki_capabilities","wiki_tool_catalog","wiki_command_tool_map","wiki_ask_history"]
        if basic=="config":
            base += ["wiki_search_with_notes","wiki_read_page:Source - 02_USER_GUIDE","wiki_read_page:Source - INSTALL"]
        return base[:max_tool_calls]
    try:
        base=release_plan_tools_base(question,max_tool_calls,intent)
    except Exception:
        base=["wiki_search_with_notes","wiki_capabilities"]
    # Pull in ask history for vague short questions.
    if len(str(question).split()) <= 4 and "wiki_ask_history" not in base:
        base.append("wiki_ask_history")
    return base[:max_tool_calls]

def release_agentic_ask(vault_path="wiki_vault",question="",max_tool_calls=14,dry_run=False,synthesize=True,protocol_mode=False):
    """Build or execute the local LLM ask workflow for agentic ask."""
    original=question
    try:
        intent=release_resolve_intent(vault_path,question)
    except Exception:
        intent={"intent":"normal","resolved_question":str(question or "")}
    resolved=intent.get("resolved_question") or str(question or "")
    repaired=release_repair_query_text(resolved) if "release_repair_query_text" in globals() else resolved
    state={}; trace=[]; evidence=[]
    planned=release_plan_tools(repaired,max_tool_calls,intent)
    if "wiki_capabilities" not in planned:
        planned.insert(0,"wiki_capabilities")
    # For non-basic questions, search remains useful.
    if not release_basic_question_intent(original) and "wiki_search_with_notes" not in planned:
        planned.insert(0,"wiki_search_with_notes")
    # Execute, supporting pseudo read tools.
    clean=[]
    for t in planned:
        if t not in clean or str(t).startswith("wiki_read_page:"):
            clean.append(t)
    for tool in clean[:max_tool_calls+4]:
        if str(tool).startswith("wiki_read_page:"):
            title=str(tool).split(":",1)[1]
            res=release_execute_agent_tool_by_name(vault_path,"wiki_read_page",{"title":title},repaired,state)
            trace.append({"tool":"wiki_read_page","arguments":{"title":title},"result_type":type(res).__name__,"forced_read":True})
            evidence.append({"tool":"wiki_read_page","result":res})
            continue
        res=release_execute_agent_tool_by_name(vault_path,tool,{"query":repaired},repaired,state)
        trace.append({"tool":tool,"result_type":type(res).__name__})
        evidence.append({"tool":tool,"result":res})
        if tool=="wiki_search_with_notes" and isinstance(res,dict) and res.get("results"):
            # Read top hit for any non-basic query with evidence.
            top=res["results"][0].get("title")
            if top:
                state["focus_page"]=top
    # Read top hit if available and not already read.
    if state.get("focus_page") and not release_basic_question_intent(original):
        already={e.get("result",{}).get("title") for e in evidence if e.get("tool")=="wiki_read_page" and isinstance(e.get("result"),dict)}
        if state["focus_page"] not in already:
            res=release_execute_agent_tool_by_name(vault_path,"wiki_read_page",{"title":state["focus_page"]},repaired,state)
            trace.append({"tool":"wiki_read_page","arguments":{"title":state["focus_page"]},"result_type":type(res).__name__,"follow_up":True})
            evidence.append({"tool":"wiki_read_page","result":res})
    # Basic questions should not depend on Ollama at all if it is flaky.
    if release_basic_question_intent(original):
        answer=release_fallback_answer_router(original,trace,evidence,None)
        return {"ok":True,"question":original,"resolved_intent":intent,"answer":answer,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm":{"ok":None,"skipped_for_basic_answer":True},"llm_prompt":""}
    prompt=(release_agent_system_prompt(vault_path) if "release_agent_system_prompt" in globals() else "# LLM Wiki MCP Agent")+"\n\n"
    prompt+="# Release Answer-First Ask\n\nAnswer the user directly. If the wiki has no exact page, answer from runtime capabilities, ask history, and general tool context. Never output only a generic failure message.\n\n"
    prompt+="Original question: "+str(original)+"\nResolved question: "+str(repaired)+"\nResolved intent: "+_release_json.dumps(intent,ensure_ascii=False)[:1200]+"\n\n"
    prompt+="## Evidence\n"
    for item in evidence:
        prompt+="\n### "+str(item.get("tool"))+"\n```json\n"+_release_json.dumps(item.get("result"),indent=2,ensure_ascii=False)[:4200]+"\n```\n"
    if dry_run or not synthesize:
        return {"ok":True,"question":original,"resolved_intent":intent,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm_prompt":prompt,"answer":None}
    llm=release_ollama_generate(vault_path,prompt) if "release_ollama_generate" in globals() else {"ok":False,"error":"missing ollama","answer":""}
    ans=release_normalise_answer_text((llm.get("answer") or "").strip())
    if not ans:
        ans=release_fallback_answer_router(original,trace,evidence,llm.get("error") or "empty LLM response")
    return {"ok":bool(llm.get("ok")),"question":original,"resolved_intent":intent,"answer":ans,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm":llm,"llm_prompt":prompt}

def release_ask(vault_path="wiki_vault",question="",mode=None,max_tool_calls=14,dry_run=False):
    """Build or execute the local LLM ask workflow for ask."""
    mode=(mode or release_current_ask_mode(vault_path) or "agentic").lower()
    if mode=="plain":
        out=release_plain_ask(vault_path,question,top_k=5,dry_run=dry_run)
        if not dry_run and not (out.get("answer") or "").strip():
            out["answer"]=release_fallback_answer_router(question,[],[{"tool":"plain_search","result":out.get("search",{})}],out.get("llm",{}).get("error"))
        return out
    return release_agentic_ask(vault_path,question,max_tool_calls=max_tool_calls,dry_run=dry_run,synthesize=not dry_run)

# Release compatibility alias for late-added functions
re = _release_re

# Release regex compatibility alias
re = _release_re


# ---------------------------------------------------------------------------
# Multi-pass blank synthesis retry and notes/wiki expansion
# ---------------------------------------------------------------------------

def release_read_page_text_for_title(vault_path, title):
    """Implement the read page text for title operation for the local LLM Wiki workflow."""
    try:
        res=release_execute_agent_tool_by_name(vault_path,"wiki_read_page",{"title":title},title,{})
        if isinstance(res,dict):
            return res.get("text") or res.get("content") or ""
    except Exception:
        pass
    return ""

def release_second_pass_titles_from_evidence(evidence, limit=6):
    """Implement the second pass titles from evidence operation for the local LLM Wiki workflow."""
    titles=[]
    for item in evidence or []:
        if not isinstance(item,dict):
            continue
        if item.get("tool")=="wiki_search_with_notes":
            result=item.get("result") or {}
            for r in result.get("results",[])[:limit*2]:
                title=r.get("title")
                if title and title not in titles:
                    titles.append(title)
        if len(titles)>=limit:
            break
    return titles[:limit]

def release_collect_second_pass_evidence(vault_path, question, first_evidence, limit=6):
    """Build a richer evidence pack from the first search hits plus notes."""
    extra=[]
    titles=release_second_pass_titles_from_evidence(first_evidence, limit=limit)
    for title in titles:
        try:
            page=release_execute_agent_tool_by_name(vault_path,"wiki_read_page",{"title":title},question,{})
            extra.append({"tool":"wiki_read_page","result":page,"second_pass":True})
        except Exception as exc:
            extra.append({"tool":"wiki_read_page","result":{"ok":False,"title":title,"error":str(exc)},"second_pass":True})
        try:
            notes=release_execute_agent_tool_by_name(vault_path,"wiki_candidate_links",{"title":title},question,{})
            extra.append({"tool":"wiki_candidate_links","result":notes,"second_pass":True})
        except Exception:
            pass
    # Search notes with the original question too.
    try:
        notes_search=release_execute_agent_tool_by_name(vault_path,"wiki_search_with_notes",{"query":question,"limit":10},question,{})
        extra.append({"tool":"wiki_search_with_notes","result":notes_search,"second_pass":True,"query":"original"})
    except Exception:
        pass
    return extra

def release_build_second_pass_prompt(vault_path, question, first_trace, first_evidence, extra_evidence):
    """Implement the build second pass prompt operation for the local LLM Wiki workflow."""
    prompt=(release_agent_system_prompt(vault_path) if "release_agent_system_prompt" in globals() else "# LLM Wiki MCP Agent")+"\n\n"
    prompt+="# Release Second-Pass Synthesis\n\n"
    prompt+="The previous LLM synthesis returned blank. You must answer the original question using the expanded evidence below. If evidence is partial, give a best-effort answer and say what evidence you used.\n\n"
    prompt+="Original question: "+str(question)+"\n\n"
    prompt+="## First tool trace\n```json\n"+_release_json.dumps(first_trace,indent=2,ensure_ascii=False)[:4000]+"\n```\n\n"
    prompt+="## First-pass evidence\n"
    for item in first_evidence or []:
        prompt+="\n### "+str(item.get("tool"))+"\n```json\n"+_release_json.dumps(item.get("result"),indent=2,ensure_ascii=False)[:3000]+"\n```\n"
    prompt+="\n## Second-pass expanded evidence\n"
    for item in extra_evidence or []:
        prompt+="\n### "+str(item.get("tool"))+"\n```json\n"+_release_json.dumps(item.get("result"),indent=2,ensure_ascii=False)[:4500]+"\n```\n"
    prompt+="\n## Instructions\nAnswer directly. Do not return blank. Use the page reads, search hits, notes and candidate links to build the answer.\n"
    return prompt

def release_try_second_pass_synthesis(vault_path, question, trace, evidence):
    """Implement the try second pass synthesis operation for the local LLM Wiki workflow."""
    extra=release_collect_second_pass_evidence(vault_path, question, evidence, limit=6)
    prompt=release_build_second_pass_prompt(vault_path, question, trace, evidence, extra)
    llm=release_ollama_generate(vault_path,prompt) if "release_ollama_generate" in globals() else {"ok":False,"error":"missing ollama","answer":""}
    answer=release_normalise_answer_text((llm.get("answer") or "").strip()) if "release_normalise_answer_text" in globals() else (llm.get("answer") or "").strip()
    if not answer:
        answer=release_fallback_answer_router(question, trace, (evidence or [])+extra, llm.get("error") or "empty LLM response after second pass")
    return {"answer":answer,"llm":llm,"extra_evidence":extra,"prompt":prompt}

def release_agentic_ask(vault_path="wiki_vault",question="",max_tool_calls=14,dry_run=False,synthesize=True,protocol_mode=False):
    """Release ask wraps Release behaviour and adds second-pass synthesis on blank LLM output.

    For basic questions, keep direct runtime answers. For non-basic questions,
    call the LLM once; if blank, expand using top wiki links + notes and call once again.
    """
    original=question
    # Reuse the existing implementation body by calling the previous-named function if available
    # is impossible after replacement, so rebuild via broad evidence path.
    try:
        intent=release_resolve_intent(vault_path,question)
    except Exception:
        intent={"intent":"normal","resolved_question":str(question or "")}
    resolved=intent.get("resolved_question") or str(question or "")
    repaired=release_repair_query_text(resolved) if "release_repair_query_text" in globals() else resolved
    state={}; trace=[]; evidence=[]
    # Basic questions use direct answers.
    if "release_basic_question_intent" in globals() and release_basic_question_intent(original):
        for tool in ["wiki_capabilities","wiki_tool_catalog","wiki_command_tool_map","wiki_ask_history"]:
            res=release_execute_agent_tool_by_name(vault_path,tool,{"query":repaired},repaired,state)
            trace.append({"tool":tool,"result_type":type(res).__name__})
            evidence.append({"tool":tool,"result":res})
        answer=release_fallback_answer_router(original,trace,evidence,None)
        return {"ok":True,"question":original,"resolved_intent":intent,"answer":answer,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm":{"ok":None,"skipped_for_basic_answer":True},"llm_prompt":""}
    # Normal evidence path.
    planned=["wiki_search_with_notes","wiki_capabilities"]
    try:
        p=release_plan_tools(repaired,max_tool_calls,intent)
        for tool in p:
            if tool not in planned:
                planned.append(tool)
    except Exception:
        pass
    for tool in planned[:max_tool_calls]:
        if str(tool).startswith("wiki_read_page:"):
            title=str(tool).split(":",1)[1]
            res=release_execute_agent_tool_by_name(vault_path,"wiki_read_page",{"title":title},repaired,state)
            trace.append({"tool":"wiki_read_page","arguments":{"title":title},"result_type":type(res).__name__})
            evidence.append({"tool":"wiki_read_page","result":res})
            continue
        res=release_execute_agent_tool_by_name(vault_path,tool,{"query":repaired},repaired,state)
        trace.append({"tool":tool,"result_type":type(res).__name__})
        evidence.append({"tool":tool,"result":res})
        if tool=="wiki_search_with_notes" and isinstance(res,dict) and res.get("results"):
            state["focus_page"]=res["results"][0].get("title")
    if state.get("focus_page"):
        res=release_execute_agent_tool_by_name(vault_path,"wiki_read_page",{"title":state["focus_page"]},repaired,state)
        trace.append({"tool":"wiki_read_page","arguments":{"title":state["focus_page"]},"result_type":type(res).__name__,"follow_up":True})
        evidence.append({"tool":"wiki_read_page","result":res})
    prompt=(release_agent_system_prompt(vault_path) if "release_agent_system_prompt" in globals() else "# LLM Wiki MCP Agent")+"\n\n"
    prompt+="# Release First-Pass Ask\n\nAnswer the question from evidence. If evidence is partial, give a useful best-effort answer.\n\n"
    prompt+="Original question: "+str(original)+"\nResolved question: "+str(repaired)+"\n\n"
    for item in evidence:
        prompt+="\n### "+str(item.get("tool"))+"\n```json\n"+_release_json.dumps(item.get("result"),indent=2,ensure_ascii=False)[:4000]+"\n```\n"
    if dry_run or not synthesize:
        return {"ok":True,"question":original,"resolved_intent":intent,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm_prompt":prompt,"answer":None}
    llm=release_ollama_generate(vault_path,prompt) if "release_ollama_generate" in globals() else {"ok":False,"error":"missing ollama","answer":""}
    ans=release_normalise_answer_text((llm.get("answer") or "").strip()) if "release_normalise_answer_text" in globals() else (llm.get("answer") or "").strip()
    second=None
    if not ans:
        second=release_try_second_pass_synthesis(vault_path, original, trace, evidence)
        ans=second.get("answer")
        evidence=evidence+second.get("extra_evidence",[])
    if not ans:
        ans=release_fallback_answer_router(original,trace,evidence,llm.get("error") or "empty LLM response")
    return {"ok":bool(llm.get("ok")) or bool(second and second.get("llm",{}).get("ok")),"question":original,"resolved_intent":intent,"answer":ans,"tools_called":[t["tool"] for t in trace],"trace":trace,"evidence":evidence,"llm":llm,"second_pass":second,"llm_prompt":prompt}

def release_ask(vault_path="wiki_vault",question="",mode=None,max_tool_calls=14,dry_run=False):
    """Build or execute the local LLM ask workflow for ask."""
    mode=(mode or release_current_ask_mode(vault_path) or "agentic").lower()
    if mode=="plain":
        out=release_plain_ask(vault_path,question,top_k=5,dry_run=dry_run)
        if not dry_run and not (out.get("answer") or "").strip():
            second=release_try_second_pass_synthesis(vault_path,question,[],[{"tool":"plain_search","result":out.get("search",{})}])
            out["answer"]=second.get("answer")
            out["second_pass"]=second
        return out
    return release_agentic_ask(vault_path,question,max_tool_calls=max_tool_calls,dry_run=dry_run,synthesize=not dry_run)

# ---------------------------------------------------------------------------
# Documentation consolidation helpers
# ---------------------------------------------------------------------------

def release_consolidate_docs(project_root=None, archive=True):
    """Consolidate scattered Markdown notes into the public documentation set."""
    root=_ProjectPath(project_root) if project_root else _ProjectPath.cwd()
    docs=root/"docs"
    docs.mkdir(exist_ok=True)
    archive_dir=docs/"archive_internal_notes"
    if archive:
        archive_dir.mkdir(exist_ok=True)
    categories={
        "01_OVERVIEW_AND_ARCHITECTURE.md":["ARCHITECTURE.md","METHODOLOGY.md","PYTHON_CONTEXT_API.md"],
        "02_USER_GUIDE.md":["USER_CLI_GUIDE.md","CLI_USER_EXPERIENCE.md","DIRECTORY_INGESTION.md","MAINTENANCE_COMMANDS.md"],
        "03_FUNCTIONS_AND_TOOLS.md":["FUNCTION_LIST.md","COMMAND_TO_TOOL_MAPPING.md","TOOL_ONTOLOGY.md","TOOL_CALLING_ASK_WORKFLOW.md"],
        "04_AGENTIC_ASK_AND_MEMORY.md":["AGENTIC_MCP_REASONING.md","QUERY_USAGE_AND_TOKEN_STATS.md","RECURSIVE_HISTORY_COMPRESSION.md","RECURSIVE_COGNITION_LAYER.md"],
        "05_NOTES_GRAPHS_AND_MAINTENANCE.md":["MERMAID_GRAPH_EXPORT.md","RUNTIME_GRAPH_VISUALISATION.md","SELF_HEALING_DOCUMENTATION.md","RUNTIME_EXECUTION_JOURNAL.md","CAPABILITIES_COMMAND.md"],
        "06_REFERENCES_AND_RESEARCH.md":["REFERENCES.md","BOOK_REFERENCES_AND_INSPIRATION.md","BOOK_KARPATHY_LLM_WIKI_PATTERN.md","LLM_WIKI_IMPLEMENTATIONS_AND_CREDITS.md","KARPATHY_AND_KB_CLI_COMPARISON.md"],
    }
    written=[]
    consumed=set()
    for outname, files in categories.items():
        parts=[f"# {outname.replace('_',' ').replace('.md','').title()}","","> Consolidated to reduce docs sprawl while preserving content.",""]
        for fn in files:
            src=docs/fn
            if src.exists():
                txt=src.read_text(encoding="utf-8",errors="replace")
                parts += [f"\n---\n\n## Source: {fn}\n", txt.strip(), ""]
                consumed.add(fn)
        target=docs/outname
        target.write_text("\n".join(parts),encoding="utf-8")
        written.append(str(target))
    if archive:
        keep=set(categories.keys()) | {"INDEX.md","SKILLS.md","CLI_DEMO_TUTORIAL.md","PRODUCT_ARCHITECTURE_OVERVIEW.md","RESEARCH_BLOG_ARTICLE_FROM_DEMO.md"}
        for item in docs.glob("*.md"):
            if item.name in keep:
                continue
            if item.name in consumed:
                dest=archive_dir/item.name
                if dest.exists():
                    dest.unlink()
                shutil.move(str(item),str(dest))
    index=["# Documentation Index","","The docs folder has been consolidated. Start here:",""]
    for item in sorted([x.name for x in docs.glob("*.md")]):
        index.append(f"- [{item}]({item})")
    (docs/"INDEX.md").write_text("\n".join(index),encoding="utf-8")
    return {"ok":True,"written":written,"archived":len(list(archive_dir.glob('*.md'))) if archive_dir.exists() else 0}


# Versionless DB note:
# The package no longer ships a prebuilt wiki_vault/index.sqlite3. It is rebuilt
# locally by WikiStore on first use. If a user's previous DB is malformed, the
# server layer deletes/recreates it.


# ---------------------------------------------------------------------------
# Versionless compact ingest output
# ---------------------------------------------------------------------------
def estimate_words_in_text_file(path):
    """Implement the estimate words in text file operation for the local LLM Wiki workflow."""
    try:
        text = Path(path).read_text(encoding="utf-8", errors="replace")
        return len(re.findall(r"\w+", text)), len(text)
    except Exception:
        return 0, 0

def compact_ingest_summary_text(plan, verbose=False):
    """Ingest source content into Markdown wiki pages for compact ingest summary text."""
    c = plan.get("counts", {})
    files_read = c.get("total", 0)
    processed = c.get("processed", 0)
    lines = []
    lines.append(f"Ingest scan: {plan.get('directory')}")
    lines.append(f"Files read: {files_read}")
    lines.append(
        f"New: {c.get('new', 0)} | Updated: {c.get('updated', 0)} | "
        f"Already ingested: {c.get('skipped', 0)} | Errors: {c.get('errors', 0)}"
    )
    lines.append(f"Pages written/updated: {processed}")

    total_words = 0
    total_chars = 0
    for row in (plan.get("new", []) + plan.get("updated", []) + plan.get("skipped", [])):
        path = row.get("path")
        if path and str(path).lower().endswith((".md", ".markdown", ".txt", ".rtf")):
            w, ch = estimate_words_in_text_file(path)
            total_words += w
            total_chars += ch
    if total_words or total_chars:
        lines.append(f"Approx document words: {total_words}")
        lines.append(f"Approx document characters: {total_chars}")

    if plan.get("skipped"):
        lines.append(f"{len(plan['skipped'])} files were already ingested and unchanged.")

    if verbose:
        if plan.get("updated"):
            lines.append("")
            lines.append("Updated files:")
            for row in plan["updated"]:
                lines.append(f"- {row.get('path')} ({str(row.get('previous_sha256',''))[:12]} -> {str(row.get('sha256',''))[:12]})")
        if plan.get("new"):
            lines.append("")
            lines.append("New files:")
            for row in plan["new"]:
                lines.append(f"- {row.get('path')} ({str(row.get('sha256',''))[:12]})")
        if plan.get("skipped"):
            lines.append("")
            lines.append("Already ingested:")
            for row in plan["skipped"]:
                lines.append(f"- {row.get('path')} ({str(row.get('sha256',''))[:12]})")
    if plan.get("errors"):
        lines.append("")
        lines.append("Errors:")
        for row in plan["errors"]:
            lines.append(f"- {row.get('path')}: {row.get('error')}")
    return "\n".join(lines)

# Override existing ingest summary function if present.
def release_ingest_directory_summary_text(plan):
    """Ingest source content into Markdown wiki pages for ingest directory summary text."""
    return compact_ingest_summary_text(plan, verbose=bool(plan.get("verbose")))

def release_ingest_directory_summary_text(plan):
    """Ingest source content into Markdown wiki pages for ingest directory summary text."""
    return compact_ingest_summary_text(plan, verbose=bool(plan.get("verbose")))

def release_ingest_directory_summary_text(plan):
    """Ingest source content into Markdown wiki pages for ingest directory summary text."""
    return compact_ingest_summary_text(plan, verbose=bool(plan.get("verbose")))

# ---------------------------------------------------------------------------
# Release self-history and public-answer sanitisation fixes
# ---------------------------------------------------------------------------

def release_sanitize_public_text(text):
    """Redact local-only paths and private workspace names.

    The wiki may ingest source documents from arbitrary user machines. Runtime
    service URLs/IPs are intentionally visible to the local user, because they
    are needed for Ollama configuration and diagnostics. Public release docs use
    generic localhost examples and should not include developer-specific private
    addresses.
    """
    s = str(text or "")
    try:
        # Windows drive paths and common user/profile paths.
        s = _release_re.sub(r"[A-Za-z]:\\(?:[^\s`'\")\]>|]+\\?)+", "[local path redacted]", s)
        s = _release_re.sub(r"[A-Za-z]:/(?:[^\s`'\")\]>|]+/?)+", "[local path redacted]", s)
        # WSL/Linux mounted drive paths and home/sandbox/tmp project paths.
        s = _release_re.sub(r"/(?:mnt|home|Users|tmp|var/tmp|private/tmp|sandbox)/(?:[^\s`'\")\]>|]+/?)+", "[local path redacted]", s)
        private_name = "".join(["Qyber", "netics"])
        s = s.replace(private_name, "local project")
        s = s.replace(private_name.lower(), "local project")
    except Exception:
        pass
    return s

def release_sanitize_public_obj(obj):
    """Recursively sanitise strings inside a JSON-like object for prompts/answers."""
    if isinstance(obj, str):
        return release_sanitize_public_text(obj)
    if isinstance(obj, list):
        return [release_sanitize_public_obj(x) for x in obj]
    if isinstance(obj, tuple):
        return tuple(release_sanitize_public_obj(x) for x in obj)
    if isinstance(obj, dict):
        cleaned = {}
        for k, v in obj.items():
            # Keep local raw provenance out of answer prompts. Filename/wiki title remain.
            if str(k).lower() in {"path", "full_path", "wiki_path", "source_path", "db", "vault"}:
                cleaned[k] = "[local path redacted]"
            elif str(k).lower() in {"host", "ollama_host"}:
                cleaned[k] = release_sanitize_public_text(v)
            else:
                cleaned[k] = release_sanitize_public_obj(v)
        return cleaned
    return obj


def release_recent_history_markdown(vault_path="wiki_vault", limit=8, answer_chars=1200):
    """Return recent prompt/response turns as bounded Markdown for recursive ask."""
    rows = release_load_history(vault_path, int(limit)) if "release_load_history" in globals() else []
    if not rows:
        return ""
    lines = ["# Recent Ask/Reply History", "", "These are recent local CLI ask turns. Use them for follow-up context, not as immutable truth.", ""]
    for i, row in enumerate(rows[-int(limit):], 1):
        q = release_sanitize_public_text(row.get("question", "")).strip()
        a = release_sanitize_public_text(row.get("answer", "")).strip()
        if len(a) > answer_chars:
            a = a[:answer_chars].rstrip() + " ..."
        if q:
            lines.append(f"## Turn {i} - User")
            lines.append(q)
            lines.append("")
        if a:
            lines.append(f"## Turn {i} - Assistant")
            lines.append(a)
            lines.append("")
    return "\n".join(lines).strip()


def release_history_answer(vault_path="wiki_vault", limit=8):
    """Natural-language answer for questions about prompt/response self-history."""
    rows = release_load_history(vault_path, int(limit)) if "release_load_history" in globals() else []
    if not rows:
        return (
            "I do not yet have useful recorded ask-history turns in this vault. "
            "From this patched release onward, every completed `/ask` turn is stored locally in `ask_history.jsonl` with the user question, the synthesis prompt, the final answer, sources/tools and token estimates. "
            "After one or two `/ask` turns, `/history` and `wiki_ask_history` should show the recent prompt/response trail."
        )
    lines = [
        "Yes. My local self-history now includes recent `/ask` prompt/response turns from this vault.",
        "",
        "The stored record contains the user question, the prompt/context sent to the local model, the final assistant answer, sources/tools used and token estimates. The agent can retrieve this through `wiki_ask_history`, and recent turns are included in the agentic prompt as bounded follow-up context.",
        "",
        "Recent turns:",
    ]
    for row in rows[-int(limit):]:
        q = release_sanitize_public_text(row.get("question", "")).strip()
        a = " ".join(release_sanitize_public_text(row.get("answer", "")).split())[:220]
        if q:
            lines.append(f"- User: {q}")
        if a:
            lines.append(f"  Assistant: {a}" + (" ..." if len(a) >= 220 else ""))
    return "\n".join(lines)


def release_context_window_answer(vault_path="wiki_vault"):
    """Answer context-window questions from available configuration without hallucinating."""
    cfg = release_load_config(vault_path) if "release_load_config" in globals() else {}
    ask = cfg.get("ask", {}) if isinstance(cfg, dict) else {}
    wiki = cfg.get("wiki", {}) if isinstance(cfg, dict) else {}
    recent = release_load_history(vault_path, int(ask.get("keep_recent_turns", 8))) if "release_load_history" in globals() else []
    history_tokens = release_history_token_total(recent) if recent and "release_history_token_total" in globals() else 0
    return "\n".join([
        "I do not know the exact native context-window length of the configured Ollama model from the wiki alone. That depends on the model and Ollama runtime settings.",
        "",
        "What I can say from my own configuration is:",
        f"- Retrieved wiki context target: about `{wiki.get('context_char_limit', 12000)}` characters.",
        f"- Recursive ask/history budget: about `{ask.get('max_history_tokens', 4000)}` estimated tokens.",
        f"- Recent exact turns kept before compression: `{ask.get('keep_recent_turns', 8)}`.",
        f"- Current recent-history estimate: about `{history_tokens}` tokens.",
        "",
        "For an exact runtime/model context length, check the Ollama model configuration or run a dry prompt inspection with `/ask-agentic-dry <question>` and compare it with your model's configured context limit.",
    ])


def release_basic_question_intent(question):
    """Recognise conversational/runtime questions that should get deterministic answers."""
    q = str(question or "").strip().lower()
    q_norm = _release_re.sub(r"[^a-z0-9\s:_\-\.]", "", q)
    if "context window" in q_norm or "context length" in q_norm or "context-window" in q_norm:
        return "context_window"
    if any(x in q_norm for x in ["history of your replies", "history of replies", "prompts and responses", "prompt and response", "ask history", "self history"]):
        return "history"
    if "does" in q_norm and "context" in q_norm and "history" in q_norm:
        return "history"
    if q_norm in {"what are you", "what r u", "who are you", "tell me what you are", "tell me about yourself", "what is this", "what is this tool", "describe yourself", "are you ok", "are you okay"}:
        return "identity"
    if q_norm in {"what can you do", "what tools do you have", "tell me your tools", "show tools", "help me", "how do i use you", "how to use this"} or "what can you do" in q_norm:
        return "capabilities"
    if "config" in q_norm or "startup" in q_norm or "ollama host" in q_norm:
        return "config"
    if "followup" in q_norm or "follow-up" in q_norm or q_norm in {"continue", "write more", "more", "you do it", "do it"}:
        return "followup"
    return None


# Keep the prior implementation available for debugging, but override the public router.
_release_previous_fallback_answer_router = release_fallback_answer_router if "release_fallback_answer_router" in globals() else None

def release_fallback_answer_router(question, trace=None, evidence=None, llm_error=None, vault_path="wiki_vault"):
    """Answer-first fallback that understands self-history/context questions."""
    intent = release_basic_question_intent(question)
    if intent == "history":
        return release_history_answer(vault_path)
    if intent == "context_window":
        return release_context_window_answer(vault_path)
    if intent == "identity":
        return release_sanitize_public_text(release_identity_answer(question, evidence, llm_error))
    if intent == "capabilities":
        return release_sanitize_public_text(release_capabilities_answer(question, evidence, llm_error))
    if intent == "config":
        return release_sanitize_public_text(release_config_answer(question, evidence, llm_error))
    if intent == "followup":
        return release_sanitize_public_text(release_discussion_followup_answer(question, evidence, llm_error))
    if "release_evidence_extract_answer" in globals():
        try:
            return release_sanitize_public_text(release_evidence_extract_answer(question, trace or [], release_sanitize_public_obj(evidence or []), llm_error))
        except Exception:
            pass
    return release_sanitize_public_text(
        "I can answer, but the LLM did not return a final synthesis. Based on the available runtime evidence, this is an LLM Wiki MCP: a local Markdown wiki system with search, page reads, Ollama-backed asking, MCP tools, provenance-aware ingest, notes, graphs and maintenance commands."
    )


# Wrap tool execution so self-history is useful and raw local paths/IPs are not fed into prompts.
_release_previous_execute_agent_tool_by_name = release_execute_agent_tool_by_name if "release_execute_agent_tool_by_name" in globals() else None

def release_execute_agent_tool_by_name(vault_path, name, args, question, state=None):
    """Execute read-only agent tools with sanitised public outputs."""
    args = args or {}
    if name == "wiki_ask_history":
        limit = int(args.get("limit", 10))
        rows = release_load_history(vault_path, limit) if "release_load_history" in globals() else []
        return {
            "ok": True,
            "history": release_sanitize_public_obj(rows),
            "count": len(rows),
            "note": "Each row may include question, prompt/context, answer, sources/tools and estimated token counts.",
        }
    if _release_previous_execute_agent_tool_by_name:
        return release_sanitize_public_obj(_release_previous_execute_agent_tool_by_name(vault_path, name, args, question, state))
    return {"ok": False, "error": "no tool executor available"}


def release_agentic_ask(vault_path="wiki_vault", question="", max_tool_calls=14, dry_run=False, synthesize=True, protocol_mode=False):
    """Agentic ask with recorded prompt/response self-history and safe public evidence.

    This final override fixes the release behaviour where the public `/ask` path
    could answer without appending to `ask_history.jsonl`. Recent ask/reply
    history is now included in the model prompt, and every completed answer is
    recorded for future self-history queries.
    """
    original = str(question or "")
    try:
        intent = release_resolve_intent(vault_path, original)
    except Exception:
        intent = {"intent": "normal", "resolved_question": original}
    resolved = intent.get("resolved_question") or original
    repaired = release_repair_query_text(resolved) if "release_repair_query_text" in globals() else resolved
    state, trace, evidence = {}, [], []
    basic = release_basic_question_intent(original)

    if basic in {"identity", "capabilities", "config", "history", "context_window", "followup"}:
        planned = ["wiki_capabilities", "wiki_tool_catalog", "wiki_command_tool_map", "wiki_ask_history"]
    else:
        planned = ["wiki_search_with_notes", "wiki_capabilities"]
        try:
            for tool in release_plan_tools(repaired, max_tool_calls, intent):
                if tool not in planned:
                    planned.append(tool)
        except Exception:
            pass
        if any(x in original.lower() for x in ["history", "previous", "recent", "context", "reply", "responses", "prompt"]):
            if "wiki_ask_history" not in planned:
                planned.append("wiki_ask_history")

    clean = []
    for t in planned:
        if str(t).startswith("wiki_read_page:") or t not in clean:
            clean.append(t)

    for tool in clean[: int(max_tool_calls) + 4]:
        if str(tool).startswith("wiki_read_page:"):
            title = str(tool).split(":", 1)[1]
            res = release_execute_agent_tool_by_name(vault_path, "wiki_read_page", {"title": title}, repaired, state)
            trace.append({"tool": "wiki_read_page", "arguments": {"title": title}, "result_type": type(res).__name__, "forced_read": True})
            evidence.append({"tool": "wiki_read_page", "result": res})
            continue
        res = release_execute_agent_tool_by_name(vault_path, tool, {"query": repaired, "limit": 10}, repaired, state)
        trace.append({"tool": tool, "result_type": type(res).__name__})
        evidence.append({"tool": tool, "result": res})
        if tool == "wiki_search_with_notes" and isinstance(res, dict) and res.get("results"):
            top = res["results"][0].get("title")
            if top:
                state["focus_page"] = top

    if state.get("focus_page") and not basic:
        already = {e.get("result", {}).get("title") for e in evidence if e.get("tool") == "wiki_read_page" and isinstance(e.get("result"), dict)}
        if state["focus_page"] not in already:
            res = release_execute_agent_tool_by_name(vault_path, "wiki_read_page", {"title": state["focus_page"]}, repaired, state)
            trace.append({"tool": "wiki_read_page", "arguments": {"title": state["focus_page"]}, "result_type": type(res).__name__, "follow_up": True})
            evidence.append({"tool": "wiki_read_page", "result": res})

    evidence = release_sanitize_public_obj(evidence)
    history_md = release_recent_history_markdown(vault_path, limit=8)
    prompt_parts = [
        release_agent_system_prompt(vault_path) if "release_agent_system_prompt" in globals() else "# LLM Wiki MCP Agent",
        "",
        "# Release Answer-First Ask",
        "",
        "Answer the user directly. Use wiki evidence, runtime state, and recent ask/reply history when relevant. If exact evidence is missing, say so and give a cautious best-effort answer. Do not output only references.",
        "",
    ]
    if history_md:
        prompt_parts += [history_md, ""]
    prompt_parts += [
        "# Current Question",
        f"Original question: {release_sanitize_public_text(original)}",
        f"Resolved question: {release_sanitize_public_text(repaired)}",
        f"Resolved intent: {_release_json.dumps(release_sanitize_public_obj(intent), ensure_ascii=False)[:1200]}",
        "",
        "# Evidence",
    ]
    for item in evidence:
        prompt_parts.append("\n### " + str(item.get("tool")))
        prompt_parts.append("```json\n" + _release_json.dumps(item.get("result"), indent=2, ensure_ascii=False)[:4200] + "\n```")
    prompt = "\n".join(prompt_parts)

    if dry_run or not synthesize:
        return {"ok": True, "question": original, "resolved_intent": intent, "tools_called": [t["tool"] for t in trace], "trace": trace, "evidence": evidence, "llm_prompt": prompt, "answer": None}

    if basic:
        answer = release_fallback_answer_router(original, trace, evidence, None, vault_path=vault_path)
        llm = {"ok": None, "skipped_for_basic_answer": True}
    else:
        llm = release_ollama_generate(vault_path, prompt) if "release_ollama_generate" in globals() else {"ok": False, "error": "missing ollama", "answer": ""}
        answer = release_normalise_answer_text((llm.get("answer") or "").strip()) if "release_normalise_answer_text" in globals() else (llm.get("answer") or "").strip()
        if not answer and "release_try_second_pass_synthesis" in globals():
            second = release_try_second_pass_synthesis(vault_path, original, trace, evidence)
            answer = second.get("answer") or ""
            llm = {**llm, "second_pass": second.get("llm")}
            evidence = evidence + release_sanitize_public_obj(second.get("extra_evidence", []))
        if not answer:
            answer = release_fallback_answer_router(original, trace, evidence, llm.get("error") or "empty LLM response", vault_path=vault_path)

    answer = release_sanitize_public_text(answer)
    metrics = {}
    if "release_record_ask_turn_with_metrics" in globals() and not dry_run:
        try:
            metrics = release_record_ask_turn_with_metrics(
                vault_path,
                original,
                prompt,
                answer,
                [t.get("tool") for t in trace],
                llm.get("elapsed_seconds") if isinstance(llm, dict) else None,
                {"agentic": True, "protocol_mode": protocol_mode, "ollama": release_sanitize_public_obj(llm)},
            )
        except Exception as exc:
            metrics = {"ok": False, "error": str(exc)}
    return {
        "ok": bool(llm.get("ok")) if isinstance(llm, dict) and llm.get("ok") is not None else True,
        "question": original,
        "resolved_intent": intent,
        "answer": answer,
        "tools_called": [t["tool"] for t in trace],
        "trace": trace,
        "evidence": evidence,
        "llm": release_sanitize_public_obj(llm),
        "metrics": metrics,
        "llm_prompt": prompt,
    }


def release_ask(vault_path="wiki_vault", question="", mode=None, max_tool_calls=14, dry_run=False):
    """Unified ask entry point with history recording for agentic mode."""
    mode = (mode or release_current_ask_mode(vault_path) or "agentic").lower()
    if mode == "plain":
        out = release_plain_ask(vault_path, question, top_k=5, dry_run=dry_run)
        if not dry_run and not (out.get("answer") or "").strip():
            out["answer"] = release_fallback_answer_router(question, [], [{"tool": "plain_search", "result": out.get("search", {})}], out.get("llm", {}).get("error"), vault_path=vault_path)
        if not dry_run and out.get("answer") and "release_record_ask_turn_with_metrics" in globals():
            try:
                release_record_ask_turn_with_metrics(vault_path, question, out.get("llm_prompt", ""), release_sanitize_public_text(out.get("answer", "")), ["plain_search"], None, {"agentic": False})
            except Exception:
                pass
        if out.get("answer"):
            out["answer"] = release_sanitize_public_text(out["answer"])
        return out
    return release_agentic_ask(vault_path, question, max_tool_calls=max_tool_calls, dry_run=dry_run, synthesize=not dry_run)


# ---------------------------------------------------------------------------
# Release hardening: public URL/path sanitisation final override
# ---------------------------------------------------------------------------

_RELEASE_LOCAL_URL_RE = _release_re.compile(
    r"https?://(?:localhost|127\.0\.0\.1|0\.0\.0\.0|10\.\d{1,3}\.\d{1,3}\.\d{1,3}|192\.168\.\d{1,3}\.\d{1,3}|172\.(?:1[6-9]|2\d|3[0-1])\.\d{1,3}\.\d{1,3})(?::\d+)?",
    _release_re.IGNORECASE,
)
_RELEASE_PRIVATE_IP_RE = _release_re.compile(
    r"\b(?:127\.0\.0\.1|0\.0\.0\.0|10\.\d{1,3}\.\d{1,3}\.\d{1,3}|192\.168\.\d{1,3}\.\d{1,3}|172\.(?:1[6-9]|2\d|3[0-1])\.\d{1,3}\.\d{1,3})\b"
)
_RELEASE_WINDOWS_PATH_RE = _release_re.compile(r"(?<![A-Za-z])\b[A-Za-z]:[\\/](?:[^\s`'\")\]>|]+[\\/]?)+")
_RELEASE_POSIX_PATH_RE = _release_re.compile(r"/(?:mnt|home|Users|tmp|var/tmp|private/tmp|sandbox)/(?:[^\s`'\")\]>|]+/?)+")


def release_sanitize_public_text(text):
    """Final public-text sanitiser used by answer prompts and CLI-visible output.

    Local filesystem paths and internal project names are redacted because they
    can leak private machine details. Runtime service URLs/IPs are intentionally
    left visible to the local user so configuration and diagnostics remain
    understandable. Public release documentation still uses generic localhost
    examples and should not contain developer-specific private IP addresses.
    """
    s = str(text or "")
    try:
        s = _RELEASE_WINDOWS_PATH_RE.sub("[local path redacted]", s)
        s = _RELEASE_POSIX_PATH_RE.sub("[local path redacted]", s)
        private_name = "".join(["Qyber", "netics"])
        s = s.replace(private_name, "local project")
        s = s.replace(private_name.lower(), "local project")
    except Exception:
        pass
    return s

# ---------------------------------------------------------------------------
# Release context-budget tuning and debug-context final override
# ---------------------------------------------------------------------------

# Extend the mutable default config used by release_load_config. Existing vaults
# are deep-merged at load time, so these settings become available without
# rewriting the user's config file until they choose to save changes.
try:
    DEFAULT_CONFIG.setdefault("ask", {})
    DEFAULT_CONFIG["ask"].setdefault("context_budget_tokens", 24000)
    DEFAULT_CONFIG["ask"].setdefault("history_budget_tokens", 3000)
    DEFAULT_CONFIG["ask"].setdefault("source_budget_tokens", 6000)
    DEFAULT_CONFIG["ask"].setdefault("max_sources", 8)
    DEFAULT_CONFIG["ask"].setdefault("full_page_threshold", 3)
    DEFAULT_CONFIG["ask"].setdefault("debug_context_chars", 12000)
except Exception:
    pass


def release_context_settings(vault_path="wiki_vault"):
    """Return the active ask-context budget settings with safe defaults."""
    cfg = release_load_config(vault_path) if "release_load_config" in globals() else {}
    ask = cfg.get("ask", {}) if isinstance(cfg, dict) else {}
    def as_int(key, default, minimum=0):
        """Document the `as_int` function used by the LLM Wiki MCP release."""
        try:
            value = int(ask.get(key, default))
        except Exception:
            value = default
        return max(minimum, value)
    return {
        "context_budget_tokens": as_int("context_budget_tokens", 24000, 2000),
        "history_budget_tokens": as_int("history_budget_tokens", 3000, 0),
        "source_budget_tokens": as_int("source_budget_tokens", 6000, 500),
        "max_sources": as_int("max_sources", 8, 1),
        "full_page_threshold": as_int("full_page_threshold", 3, 0),
        "debug_context_chars": as_int("debug_context_chars", 12000, 1000),
    }


def release_trim_to_token_budget(text, token_budget, suffix="\n\n[truncated to context budget]"):
    """Trim text using the release token estimator while preserving whole words."""
    s = str(text or "")
    try:
        budget = int(token_budget)
    except Exception:
        budget = 0
    if budget <= 0:
        return ""
    if estimate_tokens_from_text(s) <= budget:
        return s
    max_chars = max(0, budget * 4 - len(suffix))
    if max_chars <= 0:
        return suffix.strip()
    cut = s[:max_chars]
    # Prefer clean paragraph/word boundaries.
    para = cut.rfind("\n\n")
    if para > max_chars * 0.65:
        cut = cut[:para]
    else:
        space = cut.rfind(" ")
        if space > max_chars * 0.80:
            cut = cut[:space]
    return cut.rstrip() + suffix


def release_read_page_tool(vault_path, title, max_chars=None):
    """Read a wiki page with a caller-controlled character budget.

    Earlier release builds capped `wiki_read_page` at 8,000 characters. That was
    too small for large-context models and made broad questions feel shallow.
    This final override lets the ask pipeline allocate fuller per-source text.
    """
    path = release_page_path(vault_path, title) if "release_page_path" in globals() else (_ProjectPath(vault_path) / "wiki" / (str(title) if str(title).endswith(".md") else str(title) + ".md"))
    if not path.exists():
        return {"ok": False, "title": title, "error": "page not found"}
    text = path.read_text(encoding="utf-8", errors="replace")
    if max_chars is None:
        max_chars = 24000
    try:
        max_chars = int(max_chars)
    except Exception:
        max_chars = 24000
    truncated = len(text) > max_chars > 0
    if max_chars > 0:
        text = text[:max_chars]
    return {
        "ok": True,
        "title": path.stem,
        "path": str(path),
        "text": text,
        "chars": len(text),
        "truncated": truncated,
    }


def release_execute_agent_tool_by_name(vault_path, name, args, question, state=None):
    """Execute read-only agent tools with sanitisation and richer page reads."""
    args = args or {}
    if name == "wiki_ask_history":
        limit = int(args.get("limit", 10))
        rows = release_load_history(vault_path, limit) if "release_load_history" in globals() else []
        return {
            "ok": True,
            "history": release_sanitize_public_obj(rows),
            "count": len(rows),
            "note": "Each row may include question, prompt/context, answer, sources/tools and estimated token counts.",
        }
    if name == "wiki_read_page":
        return release_sanitize_public_obj(release_read_page_tool(vault_path, args.get("title", ""), args.get("max_chars")))
    if name == "wiki_search_with_notes":
        return release_sanitize_public_obj(release_search_with_notes(vault_path, args.get("query") or args.get("question") or question, int(args.get("limit", 8)), int(args.get("context_chars", 1200))))
    if name == "wiki_search_following_pages":
        return release_sanitize_public_obj(release_search_following_pages(vault_path, args.get("query") or args.get("question") or question, int(args.get("limit", 5)), int(args.get("following_pages", args.get("pages", 2))), int(args.get("context_chars", 4000)), int(args.get("max_chars", 16000))))
    if _release_previous_execute_agent_tool_by_name:
        return release_sanitize_public_obj(_release_previous_execute_agent_tool_by_name(vault_path, name, args, question, state))
    return {"ok": False, "error": "no tool executor available"}


def release_self_context_markdown(vault_path, question=""):
    """Build priority operational/self-context for identity and runtime questions."""
    caps = release_sanitize_public_obj(release_runtime_capabilities(vault_path)) if "release_runtime_capabilities" in globals() else {}
    settings = release_context_settings(vault_path)
    mode = release_current_ask_mode(vault_path) if "release_current_ask_mode" in globals() else "agentic"
    lines = [
        "# Operational Self-Context",
        "",
        "Use this section before normal wiki search for questions about identity, name, context, runtime, history, capabilities, or how this CLI/MCP works.",
        "",
        "- Name: LLM Wiki MCP Agent.",
        "- Role: local-first Markdown wiki assistant with CLI and MCP tools.",
        "- Default ask mode: " + str(mode),
        "- Source of truth: local Markdown wiki pages plus AI sidecar notes and recorded ask history.",
        "- Important distinction: `/history` shows ask prompt/response history; plain `history` is treated as a semantic `/ask history` unless the user clearly asks for ask history.",
        "",
        "## Active Context Budgets",
    ]
    for k, v in settings.items():
        lines.append(f"- {k}: {v}")
    if isinstance(caps, dict):
        lines += ["", "## Runtime Capabilities Snapshot", "```json", _release_json.dumps(caps, indent=2, ensure_ascii=False)[:6000], "```"]
    return "\n".join(lines)


def release_collect_search_titles(evidence, max_sources=8):
    """Collect de-duplicated page titles from wiki_search_with_notes evidence."""
    titles = []
    for item in evidence or []:
        if not isinstance(item, dict) or item.get("tool") != "wiki_search_with_notes":
            continue
        result = item.get("result") or {}
        for row in result.get("results", []) or []:
            title = row.get("title")
            if title and title not in titles:
                titles.append(title)
            if len(titles) >= int(max_sources):
                return titles
    return titles


def release_expand_page_evidence(vault_path, titles, settings):
    """Read relevant pages using the configured per-source budget."""
    expanded = []
    max_chars = int(settings.get("source_budget_tokens", 6000)) * 4
    for title in titles[: int(settings.get("max_sources", 8))]:
        page = release_execute_agent_tool_by_name(vault_path, "wiki_read_page", {"title": title, "max_chars": max_chars}, title, {})
        expanded.append({"tool": "wiki_read_page", "result": page})
    return expanded


def release_build_agent_context_pack(vault_path="wiki_vault", question="", intent=None, trace=None, evidence=None):
    """Assemble the exact ask context pack using explicit token budgets.

    The result is used by `/ask`, `/ask-agentic-dry`, and `/debug-context` so the
    user can inspect exactly what the model is likely to see.
    """
    settings = release_context_settings(vault_path)
    original = str(question or "")
    intent = intent or {"intent": "normal", "resolved_question": original}
    trace = trace or []
    evidence = release_sanitize_public_obj(evidence or [])

    system = release_agent_system_prompt(vault_path) if "release_agent_system_prompt" in globals() else "# LLM Wiki MCP Agent"
    self_context = release_self_context_markdown(vault_path, original)
    history_md = release_recent_history_markdown(vault_path, limit=8, answer_chars=900)
    history_md = release_trim_to_token_budget(history_md, settings["history_budget_tokens"])

    # Build direct text sections instead of dumping mostly JSON. This gives the
    # LLM more usable source material while preserving traceability.
    source_sections = []
    source_titles = []
    for item in evidence:
        tool = item.get("tool")
        result = item.get("result")
        if tool == "wiki_read_page" and isinstance(result, dict) and result.get("ok"):
            title = result.get("title") or "Untitled"
            text = release_sanitize_public_text(result.get("text", ""))
            source_titles.append(title)
            source_sections.append(f"## Source: {title}\n\n{text}")
        elif tool == "wiki_search_with_notes" and isinstance(result, dict):
            rows = []
            for row in result.get("results", [])[: int(settings.get("max_sources", 8))]:
                rows.append({
                    "title": row.get("title"),
                    "score": row.get("score"),
                    "matched_terms": row.get("matched_terms"),
                    "snippet": row.get("snippet"),
                })
                if row.get("title") and row.get("title") not in source_titles:
                    source_titles.append(row.get("title"))
            source_sections.append("## Search Results\n\n```json\n" + _release_json.dumps(rows, indent=2, ensure_ascii=False) + "\n```")
        elif tool in {"wiki_capabilities", "wiki_tool_catalog", "wiki_command_tool_map", "wiki_ask_history"}:
            # Keep runtime/tool evidence, but compact it.
            source_sections.append(f"## Tool Evidence: {tool}\n\n```json\n" + _release_json.dumps(result, indent=2, ensure_ascii=False)[:6000] + "\n```")

    source_text = "\n\n".join(source_sections).strip()
    # Leave room for system, question, and answer instructions.
    used = estimate_tokens_from_text(system + self_context + history_md + original) + 1200
    remaining = max(1000, settings["context_budget_tokens"] - used)
    source_text = release_trim_to_token_budget(source_text, remaining)

    prompt_parts = [
        system,
        "",
        "# Release Answer-First Ask",
        "",
        "Answer the user directly first. Use the operational self-context, ask/reply history, and wiki sources below. If evidence is incomplete, say what is uncertain and still give the best supported answer. Do not output only references or raw extracts.",
        "",
        self_context,
    ]
    if history_md:
        prompt_parts += ["", history_md]
    prompt_parts += [
        "",
        "# Current Question",
        f"Original question: {release_sanitize_public_text(original)}",
        "Resolved intent:",
        "```json",
        _release_json.dumps(release_sanitize_public_obj(intent), indent=2, ensure_ascii=False)[:2400],
        "```",
        "",
        "# Retrieved Wiki / Tool Evidence",
        source_text or "No wiki source evidence was retrieved.",
    ]
    prompt = release_sanitize_public_text("\n".join(prompt_parts))
    prompt = release_trim_to_token_budget(prompt, settings["context_budget_tokens"], suffix="\n\n[context pack truncated to configured context_budget_tokens]")
    return {
        "ok": True,
        "question": original,
        "settings": settings,
        "prompt": prompt,
        "prompt_tokens_estimate": estimate_tokens_from_text(prompt),
        "prompt_chars": len(prompt),
        "source_titles": list(dict.fromkeys([t for t in source_titles if t])),
        "tools_called": [t.get("tool") for t in trace],
        "history_tokens_estimate": estimate_tokens_from_text(history_md) if history_md else 0,
    }


def release_basic_question_intent(question):
    """Recognise conversational/runtime questions that need self-context first."""
    q = str(question or "").strip().lower()
    q_norm = _release_re.sub(r"[^a-z0-9\s:_\-\.]", "", q)
    if "context window" in q_norm or "context length" in q_norm or "context-window" in q_norm or "context too short" in q_norm:
        return "context_window"
    if any(x in q_norm for x in ["history of your replies", "history of replies", "prompts and responses", "prompt and response", "ask history", "self history"]):
        return "history"
    if "does" in q_norm and "context" in q_norm and "history" in q_norm:
        return "history"
    if q_norm in {"history", "show history", "what history", "what history do you see"}:
        return "history"
    if q_norm in {"what are you", "what r u", "who are you", "tell me what you are", "tell me about yourself", "what is this", "what is this tool", "describe yourself", "are you ok", "are you okay", "what is your name", "whats your name", "your name", "name"}:
        return "identity"
    if any(x in q_norm for x in ["what context", "what is your context", "do you see all", "do you see ask history", "what about context", "is ask context truncated"]):
        return "context_window"
    if q_norm in {"what can you do", "what tools do you have", "tell me your tools", "show tools", "help me", "how do i use you", "how to use this"} or "what can you do" in q_norm:
        return "capabilities"
    if "config" in q_norm or "startup" in q_norm or "ollama host" in q_norm:
        return "config"
    if "followup" in q_norm or "follow-up" in q_norm or q_norm in {"continue", "write more", "more", "you do it", "do it"}:
        return "followup"
    return None


def release_context_window_answer(vault_path="wiki_vault"):
    """Answer context-window questions from active settings and history metrics."""
    settings = release_context_settings(vault_path)
    cfg = release_load_config(vault_path) if "release_load_config" in globals() else {}
    ask = cfg.get("ask", {}) if isinstance(cfg, dict) else {}
    recent = release_load_history(vault_path, int(ask.get("keep_recent_turns", 8))) if "release_load_history" in globals() else []
    history_tokens = release_history_token_total(recent) if recent and "release_history_token_total" in globals() else 0
    return "\n".join([
        "I do not know the exact native context-window length of the configured Ollama model from the wiki alone. The ask context is now configurable and less conservative than the earlier release.",
        "",
        "Retrieved wiki context and self-history are now governed by these active context settings:",
        f"- context_budget_tokens: `{settings['context_budget_tokens']}` estimated tokens for the whole ask prompt.",
        f"- history_budget_tokens: `{settings['history_budget_tokens']}` estimated tokens reserved for recent ask/reply history.",
        f"- source_budget_tokens: `{settings['source_budget_tokens']}` estimated tokens available per full source page read.",
        f"- max_sources: `{settings['max_sources']}` retrieved source pages considered.",
        f"- full_page_threshold: `{settings['full_page_threshold']}`. If search finds only this many strong sources, the pipeline favours fuller page reads.",
        f"- Current recent-history estimate: about `{history_tokens}` tokens.",
        "",
        "Use `/debug-context <question>` to see the assembled prompt and `/context-settings` to inspect or change these values. For example:",
        "",
        "```text",
        "/context-settings context-budget 24000",
        "/context-settings history-budget 3000",
        "/context-settings source-budget 6000",
        "/debug-context tell me about vishva",
        "```",
    ])


def release_agentic_ask(vault_path="wiki_vault", question="", max_tool_calls=14, dry_run=False, synthesize=True, protocol_mode=False):
    """Agentic ask with larger configurable context budgets and debug prompt pack."""
    original = str(question or "")
    settings = release_context_settings(vault_path)
    try:
        intent = release_resolve_intent(vault_path, original)
    except Exception:
        intent = {"intent": "normal", "resolved_question": original}
    resolved = intent.get("resolved_question") or original
    repaired = release_repair_query_text(resolved) if "release_repair_query_text" in globals() else resolved
    state, trace, evidence = {}, [], []
    basic = release_basic_question_intent(original)

    if basic in {"identity", "capabilities", "config", "history", "context_window", "followup"}:
        planned = ["wiki_capabilities", "wiki_tool_catalog", "wiki_command_tool_map", "wiki_ask_history"]
        # Runtime questions should not be polluted by random document hits, but
        # can still use search if the user mentions a specific wiki concept.
        if any(term in original.lower() for term in ["vishva", "bindu", "bootstrap", "eurorack", "synth"]):
            planned.insert(0, "wiki_search_with_notes")
    else:
        planned = ["wiki_search_with_notes", "wiki_capabilities", "wiki_ask_history"]
        try:
            for tool in release_plan_tools(repaired, max_tool_calls, intent):
                if tool not in planned:
                    planned.append(tool)
        except Exception:
            pass

    clean = []
    for t in planned:
        if str(t).startswith("wiki_read_page:") or t not in clean:
            clean.append(t)

    for tool in clean[: int(max_tool_calls) + 4]:
        if str(tool).startswith("wiki_read_page:"):
            title = str(tool).split(":", 1)[1]
            res = release_execute_agent_tool_by_name(vault_path, "wiki_read_page", {"title": title, "max_chars": settings["source_budget_tokens"] * 4}, repaired, state)
            trace.append({"tool": "wiki_read_page", "arguments": {"title": title}, "result_type": type(res).__name__, "forced_read": True})
            evidence.append({"tool": "wiki_read_page", "result": res})
            continue
        res = release_execute_agent_tool_by_name(vault_path, tool, {"query": repaired, "limit": settings["max_sources"]}, repaired, state)
        trace.append({"tool": tool, "result_type": type(res).__name__})
        evidence.append({"tool": tool, "result": res})
        if tool == "wiki_search_with_notes" and isinstance(res, dict) and res.get("results"):
            top = res["results"][0].get("title")
            if top:
                state["focus_page"] = top

    # Decide how much full text to read. For a few strong matches, read all top
    # matches up to full_page_threshold; otherwise read max_sources with a
    # per-source budget and let the final prompt budget trim if necessary.
    search_titles = release_collect_search_titles(evidence, settings["max_sources"])
    if search_titles and not basic:
        count_to_read = min(len(search_titles), settings["max_sources"])
        if len(search_titles) <= settings["full_page_threshold"]:
            count_to_read = len(search_titles)
        expanded = release_expand_page_evidence(vault_path, search_titles[:count_to_read], settings)
        already = {e.get("result", {}).get("title") for e in evidence if e.get("tool") == "wiki_read_page" and isinstance(e.get("result"), dict)}
        for item in expanded:
            title = item.get("result", {}).get("title") if isinstance(item.get("result"), dict) else None
            if title and title not in already:
                trace.append({"tool": "wiki_read_page", "arguments": {"title": title}, "result_type": "dict", "expanded_context": True})
                evidence.append(item)
                already.add(title)
    elif state.get("focus_page") and not basic:
        res = release_execute_agent_tool_by_name(vault_path, "wiki_read_page", {"title": state["focus_page"], "max_chars": settings["source_budget_tokens"] * 4}, repaired, state)
        trace.append({"tool": "wiki_read_page", "arguments": {"title": state["focus_page"]}, "result_type": type(res).__name__, "follow_up": True})
        evidence.append({"tool": "wiki_read_page", "result": res})

    evidence = release_sanitize_public_obj(evidence)
    context_pack = release_build_agent_context_pack(vault_path, original, intent, trace, evidence)
    prompt = context_pack["prompt"]

    if dry_run or not synthesize:
        return {
            "ok": True,
            "question": original,
            "resolved_intent": intent,
            "tools_called": [t["tool"] for t in trace],
            "trace": trace,
            "evidence": evidence,
            "context_pack": context_pack,
            "llm_prompt": prompt,
            "answer": None,
            "dry_run": True,
        }

    # Deterministic direct answers for operational self questions avoid random
    # document matches and blank LLM outputs.
    if basic:
        answer = release_fallback_answer_router(original, trace, evidence, None, vault_path=vault_path)
        llm = {"ok": None, "skipped_for_basic_answer": True}
    else:
        llm = release_ollama_generate(vault_path, prompt) if "release_ollama_generate" in globals() else {"ok": False, "error": "missing ollama", "answer": ""}
        answer = release_normalise_answer_text((llm.get("answer") or "").strip()) if "release_normalise_answer_text" in globals() else (llm.get("answer") or "").strip()
        if not answer and "release_try_second_pass_synthesis" in globals():
            second = release_try_second_pass_synthesis(vault_path, original, trace, evidence)
            answer = second.get("answer") or ""
            llm = {**llm, "second_pass": second.get("llm")}
            evidence = evidence + release_sanitize_public_obj(second.get("extra_evidence", []))
        if not answer:
            answer = release_fallback_answer_router(original, trace, evidence, llm.get("error") or "empty LLM response", vault_path=vault_path)

    answer = release_sanitize_public_text(answer)
    metrics = {}
    if "release_record_ask_turn_with_metrics" in globals() and not dry_run:
        try:
            metrics = release_record_ask_turn_with_metrics(
                vault_path,
                original,
                prompt,
                answer,
                [t.get("tool") for t in trace],
                llm.get("elapsed_seconds") if isinstance(llm, dict) else None,
                {"agentic": True, "protocol_mode": protocol_mode, "ollama": release_sanitize_public_obj(llm), "context_pack": {k: v for k, v in context_pack.items() if k != "prompt"}},
            )
        except Exception as exc:
            metrics = {"ok": False, "error": str(exc)}
    return {
        "ok": bool(llm.get("ok")) if isinstance(llm, dict) and llm.get("ok") is not None else True,
        "question": original,
        "resolved_intent": intent,
        "answer": answer,
        "tools_called": [t["tool"] for t in trace],
        "trace": trace,
        "evidence": evidence,
        "llm": release_sanitize_public_obj(llm),
        "metrics": metrics,
        "context_pack": context_pack,
        "llm_prompt": prompt,
    }


def release_ask(vault_path="wiki_vault", question="", mode=None, max_tool_calls=14, dry_run=False):
    """Unified ask entry point with configurable context budgets."""
    mode = (mode or release_current_ask_mode(vault_path) or "agentic").lower()
    if mode == "plain":
        out = release_plain_ask(vault_path, question, top_k=5, dry_run=dry_run)
        if not dry_run and not (out.get("answer") or "").strip():
            out["answer"] = release_fallback_answer_router(question, [], [{"tool": "plain_search", "result": out.get("search", {})}], out.get("llm", {}).get("error"), vault_path=vault_path)
        if not dry_run and out.get("answer") and "release_record_ask_turn_with_metrics" in globals():
            try:
                release_record_ask_turn_with_metrics(vault_path, question, out.get("llm_prompt", ""), release_sanitize_public_text(out.get("answer", "")), ["plain_search"], None, {"agentic": False})
            except Exception:
                pass
        if out.get("answer"):
            out["answer"] = release_sanitize_public_text(out["answer"])
        return out
    return release_agentic_ask(vault_path, question, max_tool_calls=max_tool_calls, dry_run=dry_run, synthesize=not dry_run)


def release_debug_context(vault_path="wiki_vault", question=""):
    """Build the current ask context without calling Ollama."""
    return release_agentic_ask(vault_path, question, dry_run=True, synthesize=False)


# ---------------------------------------------------------------------------
# Release progress callback overrides for long-running file/page operations
# ---------------------------------------------------------------------------

def release_ingest_directory_provenance(vault_path, directory, recursive=True, apply=True, progress_callback=None):
    """Ingest a directory or single file with optional one-line progress.

    The callback receives ``(current, total, path, action)`` before each file is
    processed. Existing callers can keep using the four-argument form; the CLI
    uses the callback to display a same-line counter while large ingests run.
    """
    plan=release_ingest_directory_plan(vault_path, directory, recursive=recursive)
    processed=[]
    target_path=_ProjectPath(directory).expanduser()
    base_dir=target_path.parent if target_path.is_file() else target_path
    work=list(plan.get("new",[]))+list(plan.get("updated",[]))
    total=len(work)
    if apply:
        for idx, row in enumerate(work, 1):
            try:
                if progress_callback:
                    try:
                        progress_callback(idx, total, row.get("path"), row.get("reason") or "ingest")
                    except TypeError:
                        progress_callback(idx, total, row.get("path"))
                res=release_markdown_to_wiki_page(vault_path, row["path"], row.get("wiki_title"), base_dir)
                row["wiki_title"]=res["title"]; row["wiki_path"]=res["wiki_path"]
                processed.append(row)
            except Exception as exc:
                plan.setdefault("errors",[]).append({"path":row.get("path"),"error":str(exc)})
    if progress_callback and apply:
        try:
            progress_callback(total, total, None, "done")
        except TypeError:
            progress_callback(total, total, None)
    plan["processed"]=processed
    plan["counts"]={
        "new":len(plan.get("new",[])),
        "updated":len(plan.get("updated",[])),
        "skipped":len(plan.get("skipped",[])),
        "errors":len(plan.get("errors",[])),
        "total":len(release_scan_ingest_directory(directory, recursive=recursive)),
        "processed":len(processed),
    }
    if apply and processed:
        try:
            if progress_callback:
                try:
                    progress_callback(len(processed), len(processed), None, "reindex")
                except TypeError:
                    progress_callback(len(processed), len(processed), None)
            WikiStore(WikiConfig(_ProjectPath(vault_path))).reindex()
            plan["counts"]["indexed_pages"] = len(list((_ProjectPath(vault_path) / "wiki").glob("*.md")))
        except Exception as exc:
            plan.setdefault("errors",[]).append({"path": str(vault_path), "error": f"reindex after ingest failed: {exc}"})
            plan["counts"]["errors"] = len(plan.get("errors", []))
    plan["summary"]=release_ingest_directory_summary_text(plan)
    return plan


def release_iterate_all_notes(vault_path, limit=None, progress_callback=None):
    """Build AI sidecar notes for every page with optional same-line progress."""
    wiki = release_wiki_dir(vault_path)
    pages=list(wiki.glob("*.md")) if wiki.exists() else []
    if limit:
        pages=pages[:int(limit)]
    results=[]
    total=len(pages)
    for idx, p in enumerate(pages, 1):
        if progress_callback:
            try:
                progress_callback(idx, total, p.name, "notes")
            except TypeError:
                progress_callback(idx, total, p.name)
        results.append(release_build_page_note(vault_path,p.stem))
    if progress_callback:
        try:
            progress_callback(total, total, None, "done")
        except TypeError:
            progress_callback(total, total, None)
    return {"ok":True,"count":len(pages),"results":results}


# ---------------------------------------------------------------------------
# Release polish: DOCX status, reindex progress, and vault-safe ingest scans
# ---------------------------------------------------------------------------

def release_docx_extraction_status():
    """Return the optional DOCX extraction status for the current runtime.

    DOCX parsing depends on the lightweight `python-docx` package, imported as
    `docx`. The check runs at CLI startup so users know whether `.docx` files
    will be converted to searchable text or placeholder diagnostics.
    """
    available = _release_importlib_util.find_spec("docx") is not None
    return {
        "available": available,
        "module": "python-docx",
        "message": (
            "DOCX text extraction is available."
            if available
            else "DOCX text extraction is unavailable. Install with `pip install -e .[docs,mcp]` or `pip install python-docx`."
        ),
    }


def release_is_inside_path(path, parent):
    """Return True when `path` is contained by `parent` after resolution."""
    try:
        _ProjectPath(path).expanduser().resolve().relative_to(_ProjectPath(parent).expanduser().resolve())
        return True
    except Exception:
        return False


def release_scan_ingest_directory_for_vault(vault_path, directory, recursive=True, progress_callback=None):
    """Scan an ingest target while excluding the active wiki vault.

    This prevents `/ingest .` from absorbing generated files under `wiki_vault/`
    when the vault lives inside the project directory being ingested. The optional
    progress callback is intentionally called during the walk, before a final
    candidate list exists, so very large directories show immediate activity.
    """
    root = _ProjectPath(directory).expanduser()
    vault_root = _ProjectPath(vault_path).expanduser().resolve()
    suffixes = release_supported_ingest_suffixes()
    files = []
    seen = set()

    def _add(candidate):
        """Document the `_add` function used by the LLM Wiki MCP release."""
        if not candidate.is_file() or candidate.suffix.lower() not in suffixes:
            return
        try:
            resolved = candidate.resolve()
        except Exception:
            resolved = candidate
        if release_is_inside_path(resolved, vault_root):
            return
        key = str(resolved)
        if key in seen:
            return
        seen.add(key)
        files.append(candidate)
        if progress_callback:
            try:
                progress_callback(len(files), 0, candidate, "scan")
            except TypeError:
                progress_callback(len(files), 0, candidate)

    if not root.exists():
        return []
    if root.is_file():
        _add(root)
    else:
        pattern = "**/*" if recursive else "*"
        for candidate in root.glob(pattern):
            _add(candidate)
    return sorted(files)


def release_reindex_with_progress(vault_path, progress_callback=None):
    """Rebuild the SQLite/FTS index with optional per-page progress updates."""
    store = WikiStore(WikiConfig(_ProjectPath(vault_path)))
    pages = sorted(store.cfg.wiki_dir.glob("*.md")) if store.cfg.wiki_dir.exists() else []
    total = len(pages)
    with store._connect() as con:
        con.execute("DELETE FROM pages")
        con.execute("DELETE FROM pages_fts")
        con.commit()
    for idx, page in enumerate(pages, 1):
        if progress_callback:
            try:
                progress_callback(idx, total, page.name, "reindex")
            except TypeError:
                progress_callback(idx, total, page.name)
        store.index_page(page.stem)
    if progress_callback:
        try:
            progress_callback(total, total, None, "done")
        except TypeError:
            progress_callback(total, total, None)
    return {"indexed_pages": total, "db_path": str(store.cfg.db_path)}


def release_ingest_directory_plan(vault_path, directory, recursive=True, progress_callback=None):
    """Return new/updated/skipped lists, excluding the active vault itself."""
    release_init_provenance_schema(vault_path)
    files = release_scan_ingest_directory_for_vault(vault_path, directory, recursive=recursive, progress_callback=progress_callback)
    pdf_status = release_pdf_extraction_status()
    out = {"directory": str(_ProjectPath(directory).resolve()), "new": [], "updated": [], "skipped": [], "errors": []}
    for p in files:
        try:
            status = release_provenance_status_for_file(vault_path, p)
            action = status.get("action")
            row = {
                "path": str(p.resolve()),
                "filename": p.name,
                "sha256": status.get("record", {}).get("sha256"),
                "mtime_iso": status.get("record", {}).get("mtime_iso"),
                "reason": status.get("reason"),
                "wiki_title": status.get("wiki_title") or release_wiki_title_for_source_file(p, directory),
            }
            if action == "skip":
                if p.suffix.lower() == ".pdf" and pdf_status.get("available") and release_wiki_page_needs_pdf_reextract(vault_path, row.get("wiki_title")):
                    row["reason"] = "PDF extraction is now available; refreshing previous placeholder page"
                    row["previous_sha256"] = status.get("record", {}).get("sha256")
                    out["updated"].append(row)
                else:
                    out["skipped"].append(row)
            elif action == "update":
                row["previous_sha256"] = status.get("previous_sha256")
                out["updated"].append(row)
            else:
                out["new"].append(row)
        except Exception as exc:
            out["errors"].append({"path": str(p), "error": str(exc)})
    out["counts"] = {"new": len(out["new"]), "updated": len(out["updated"]), "skipped": len(out["skipped"]), "errors": len(out["errors"]), "total": len(files)}
    return out


def release_ingest_directory_provenance(vault_path, directory, recursive=True, apply=True, progress_callback=None):
    """Ingest a directory/single file, skip the active vault, then reindex with progress."""
    plan = release_ingest_directory_plan(vault_path, directory, recursive=recursive, progress_callback=progress_callback)
    processed = []
    target_path = _ProjectPath(directory).expanduser()
    base_dir = target_path.parent if target_path.is_file() else target_path
    work = list(plan.get("new", [])) + list(plan.get("updated", []))
    total = len(work)
    if apply:
        for idx, row in enumerate(work, 1):
            try:
                if progress_callback:
                    try:
                        progress_callback(idx, total, row.get("path"), row.get("reason") or "ingest")
                    except TypeError:
                        progress_callback(idx, total, row.get("path"))
                res = release_markdown_to_wiki_page(vault_path, row["path"], row.get("wiki_title"), base_dir)
                row["wiki_title"] = res["title"]
                row["wiki_path"] = res["wiki_path"]
                processed.append(row)
            except Exception as exc:
                plan.setdefault("errors", []).append({"path": row.get("path"), "error": str(exc)})
    plan["processed"] = processed
    plan["counts"] = {
        "new": len(plan.get("new", [])),
        "updated": len(plan.get("updated", [])),
        "skipped": len(plan.get("skipped", [])),
        "errors": len(plan.get("errors", [])),
        "total": plan.get("counts", {}).get("total", len(plan.get("new", [])) + len(plan.get("updated", [])) + len(plan.get("skipped", []))),
        "processed": len(processed),
    }
    if apply and processed:
        try:
            reindexed = release_reindex_with_progress(vault_path, progress_callback=progress_callback)
            plan["counts"]["indexed_pages"] = reindexed.get("indexed_pages")
        except Exception as exc:
            plan.setdefault("errors", []).append({"path": str(vault_path), "error": f"reindex after ingest failed: {exc}"})
            plan["counts"]["errors"] = len(plan.get("errors", []))
    plan["summary"] = release_ingest_directory_summary_text(plan)
    return plan


# ---------------------------------------------------------------------------
# Release long-document search windows and following-page retrieval
# ---------------------------------------------------------------------------

def release_search_terms_for_window(query, max_terms=20):
    """Return normalized search terms used for centred long-document snippets."""
    terms = []
    for term in release_extract_terms(query, max_terms) if "release_extract_terms" in globals() else []:
        term = str(term or "").strip().lower()
        if len(term) >= 2 and term not in terms:
            terms.append(term)
    if not terms:
        for term in _release_re.findall(r"[A-Za-z0-9][A-Za-z0-9_\-]{1,}", str(query or "").lower()):
            if term not in terms:
                terms.append(term)
    return terms[:int(max_terms)]


def release_find_first_match_offset(text, terms):
    """Find the earliest useful term match offset in a long string."""
    lower = str(text or "").lower()
    best = None
    for term in sorted([t for t in terms if t], key=len, reverse=True):
        idx = lower.find(term.lower())
        if idx >= 0 and (best is None or idx < best):
            best = idx
    return 0 if best is None else int(best)


def release_collapse_excerpt(text):
    """Normalise whitespace in a search excerpt while preserving readable gaps."""
    excerpt = str(text or "")
    excerpt = _release_re.sub(r"[ \t]+", " ", excerpt)
    excerpt = _release_re.sub(r"\n{3,}", "\n\n", excerpt)
    return excerpt.strip()


def release_centered_match_excerpt(text, terms, context_chars=1200):
    """Return text around the actual hit instead of the start of a long page.

    Long ingested books often exceed the LLM context budget. A first-600-chars
    snippet is not useful when the match occurs on page 90. This helper centres
    the snippet near the first matched term and marks whether the start/end of
    the original page were omitted.
    """
    text = str(text or "")
    context_chars = max(200, int(context_chars or 1200))
    if not text:
        return {"offset": 0, "snippet": "", "truncated_before": False, "truncated_after": False}
    offset = release_find_first_match_offset(text, terms)
    half = max(80, context_chars // 2)
    start = max(0, offset - half)
    end = min(len(text), start + context_chars)
    if end - start < context_chars and start > 0:
        start = max(0, end - context_chars)
    snippet = text[start:end]
    if start > 0:
        snippet = "[… earlier text omitted …]\n" + snippet
    if end < len(text):
        snippet = snippet + "\n[… later text omitted …]"
    return {
        "offset": offset,
        "start": start,
        "end": end,
        "snippet": release_collapse_excerpt(snippet),
        "truncated_before": start > 0,
        "truncated_after": end < len(text),
    }


def release_page_markers(text):
    """Detect common page markers inside extracted PDF/book text."""
    markers = []
    patterns = [
        r"(?im)^\s*-{0,3}\s*Page\s+(\d+)\s+(?:of|/)\s+(\d+)\s*-{0,3}\s*$",
        r"(?im)^\s*-{0,3}\s*Page\s+(\d+)\s*-{0,3}\s*$",
        r"(?i)\bPage\s+(\d+)\s+(?:of|/)\s+(\d+)\b",
    ]
    seen = set()
    for pattern in patterns:
        for m in _release_re.finditer(pattern, str(text or "")):
            key = (m.start(), m.group(1))
            if key in seen:
                continue
            seen.add(key)
            markers.append({"offset": m.start(), "end": m.end(), "page": int(m.group(1)), "raw": m.group(0).strip()})
    markers.sort(key=lambda row: row["offset"])
    # Drop duplicate inline matches that sit inside our canonical marker line
    # (for example both `--- Page 4 of 7 ---` and `Page 4 of 7`).
    deduped = []
    for marker in markers:
        if deduped and marker.get("page") == deduped[-1].get("page") and marker["offset"] - deduped[-1]["offset"] < 40:
            continue
        deduped.append(marker)
    return deduped


def release_following_page_excerpt(text, terms, following_pages=2, context_chars=4000, max_chars=16000):
    """Return the matched page plus N following pages, or a character window.

    When explicit page markers exist, the result starts at the marker preceding
    the match and ends after the requested number of following pages. If no page
    markers are available, it falls back to a larger character window beginning
    near the match.
    """
    text = str(text or "")
    following_pages = max(0, int(following_pages or 0))
    context_chars = max(500, int(context_chars or 4000))
    max_chars = max(context_chars, int(max_chars or 16000))
    offset = release_find_first_match_offset(text, terms)
    markers = release_page_markers(text)
    page_start = None
    page_end = None
    start_page = None
    end_page = None
    if markers:
        idx = 0
        for i, marker in enumerate(markers):
            if marker["offset"] <= offset:
                idx = i
            else:
                break
        start_marker = markers[idx]
        end_idx = idx + following_pages + 1
        page_start = start_marker["offset"]
        if end_idx < len(markers):
            page_end = markers[end_idx]["offset"]
            end_page = markers[end_idx - 1].get("page")
        else:
            page_end = len(text)
            end_page = markers[-1].get("page")
        start_page = start_marker.get("page")
        excerpt = text[page_start:page_end]
        if len(excerpt) > max_chars:
            # Keep the actual hit and as much following material as possible.
            rel = max(0, offset - page_start)
            local_start = max(0, rel - 400)
            excerpt = excerpt[local_start:local_start + max_chars]
            if local_start > 0:
                excerpt = "[… earlier text on matched page omitted …]\n" + excerpt
            if local_start + max_chars < (page_end - page_start):
                excerpt = excerpt + "\n[… following page excerpt truncated by max_chars …]"
    else:
        start = max(0, offset - max(250, context_chars // 3))
        end = min(len(text), offset + context_chars * (following_pages + 1))
        excerpt = text[start:end]
        if start > 0:
            excerpt = "[… earlier text omitted …]\n" + excerpt
        if end < len(text):
            excerpt = excerpt + "\n[… later text omitted …]"
    return {
        "offset": offset,
        "start_page": start_page,
        "end_page": end_page,
        "page_markers_found": len(markers),
        "following_pages": following_pages,
        "chars": len(excerpt),
        "excerpt": release_collapse_excerpt(excerpt),
    }


def release_search_with_notes(vault_path, query, limit=8, context_chars=1200):
    """Search source pages plus AI notes with snippets centred around hits.

    This overrides the earlier release helper that returned the beginning of the
    page/note. For long books this is critical: MCP search results now show text
    around the match offset, so the LLM receives the relevant part even when the
    full document is too long for the context window.
    """
    terms = set(release_search_terms_for_window(query, 20))
    rows = []
    wiki = release_wiki_dir(vault_path)
    if not wiki.exists():
        return {"query": query, "count": 0, "results": []}
    for p in wiki.glob("*.md"):
        source = p.read_text(encoding="utf-8", errors="replace")
        np = release_note_path(vault_path, p.stem)
        note = np.read_text(encoding="utf-8", errors="replace") if np.exists() else ""
        blob = (p.stem + "\n" + source + "\n" + note).lower()
        score = 0
        matched = []
        for term in terms:
            if term in blob:
                bonus = 3 if term in p.stem.lower() else 1
                if term in note.lower():
                    bonus += 1
                # Late-book hits still matter; count modestly without allowing
                # giant books to dominate purely by repetition.
                bonus += min(4, blob.count(term) // 3)
                score += bonus
                matched.append(term)
        if score:
            excerpt_source = note if note and any(t in note.lower() for t in terms) else source
            win = release_centered_match_excerpt(excerpt_source, matched or terms, context_chars=context_chars)
            rows.append({
                "title": p.stem,
                "score": score,
                "matched_terms": matched[:12],
                "has_note": bool(note),
                "match_offset": win.get("offset"),
                "snippet": win.get("snippet", ""),
                "truncated_before": win.get("truncated_before", False),
                "truncated_after": win.get("truncated_after", False),
            })
    return {"query": query, "count": len(rows), "results": sorted(rows, key=lambda r: (-r["score"], r["title"]))[:int(limit)]}


def release_search_following_pages(vault_path, query, limit=5, following_pages=2, context_chars=4000, max_chars=16000):
    """Search and return the matched page plus N following pages/excerpt.

    This is designed for long books where a hit near the end is otherwise lost
    when the complete source page is truncated by the LLM context budget.
    """
    base = release_search_with_notes(vault_path, query, limit=limit, context_chars=min(1600, int(context_chars or 4000)))
    terms = release_search_terms_for_window(query, 20)
    results = []
    for row in base.get("results", []):
        title = row.get("title")
        page = release_read_page_tool(vault_path, title) if "release_read_page_tool" in globals() else {"ok": False}
        text = page.get("text") or ""
        # release_read_page_tool intentionally limits the text for ordinary use;
        # read the full Markdown here because this tool's purpose is navigation
        # inside large source pages.
        path = release_page_path(vault_path, title) if "release_page_path" in globals() else (release_wiki_dir(vault_path) / (str(title) + ".md"))
        if path.exists():
            text = path.read_text(encoding="utf-8", errors="replace")
        follow = release_following_page_excerpt(text, terms, following_pages=following_pages, context_chars=context_chars, max_chars=max_chars)
        item = dict(row)
        item.update({
            "start_page": follow.get("start_page"),
            "end_page": follow.get("end_page"),
            "page_markers_found": follow.get("page_markers_found"),
            "following_pages": follow.get("following_pages"),
            "excerpt_chars": follow.get("chars"),
            "excerpt": follow.get("excerpt"),
        })
        results.append(item)
    return {"query": query, "limit": int(limit), "following_pages": int(following_pages), "count": len(results), "results": results}


def _release_wikistore_search_contextual(self, query: str, limit: int = 10):
    """Contextual WikiStore.search replacement with centred snippets."""
    key = stable_key('search-release-contextual', str(self.cfg.vault), query, limit)
    cached = search_cache.get(key)
    if cached is not None:
        return cached
    words = plain_words(query)
    expanded = expand_search_terms(query)
    q = ' '.join(words) or query.strip()
    rows = []
    search_mode = 'fts'
    with self._connect() as con:
        fts_candidates = []
        if q.strip():
            fts_candidates.append(q)
        if words:
            fts_candidates.append(' '.join(w + '*' for w in words))
        if expanded:
            fts_candidates.append(' OR '.join(t + '*' for t in expanded))
        for candidate in fts_candidates:
            if rows or not candidate.strip():
                continue
            try:
                for row in con.execute('SELECT title, body, rank FROM pages_fts WHERE pages_fts MATCH ? ORDER BY rank LIMIT ?', (candidate, limit)):
                    d = dict(row); d['match_mode'] = 'fts'; rows.append(d)
            except Exception:
                pass
        if not rows:
            search_mode = 'fallback'
            tokens = set(expanded or words)
            compact_query = _release_re.sub(r'[^a-z0-9]+', '', query.lower())
            for row in con.execute('SELECT title, body FROM pages_fts LIMIT 1000'):
                body = row['body']
                body_words = plain_words(body)
                title_l = row['title'].lower()
                title_compact = _release_re.sub(r'[^a-z0-9]+', '', title_l)
                score = 0
                for token in tokens:
                    score += sum(1 for w in body_words if w == token or w.startswith(token) or token.startswith(w))
                    if token in title_l or title_l.startswith(token):
                        score += 8
                    if token in title_compact:
                        score += 4
                if compact_query and compact_query in title_compact:
                    score += 12
                for token in tokens:
                    try:
                        import difflib as _difflib
                        if _difflib.SequenceMatcher(None, token, title_compact).ratio() > 0.72:
                            score += 5
                    except Exception:
                        pass
                if score:
                    rows.append({'title': row['title'], 'body': body, 'rank': -score, 'match_mode': 'fallback'})
            rows.sort(key=lambda r: r['rank'])
            rows = rows[:limit]
    terms = release_search_terms_for_window(query, 20)
    results = []
    for r in rows:
        body = r.get('body', '')
        win = release_centered_match_excerpt(body, terms, context_chars=900)
        results.append({
            'title': r['title'],
            'score': float(r.get('rank', 0.0)),
            'snippet': win.get('snippet', ''),
            'match_offset': win.get('offset'),
            'truncated_before': win.get('truncated_before', False),
            'truncated_after': win.get('truncated_after', False),
            'links': wiki_links(body)[:10],
            'match_mode': r.get('match_mode', search_mode),
        })
    payload = {'query': query, 'results': results, 'count': len(results), 'diagnostics': search_query_diagnostics(query), 'search_mode': search_mode}
    search_cache.set(key, payload)
    return payload

try:
    WikiStore.search = _release_wikistore_search_contextual
except Exception:
    pass
